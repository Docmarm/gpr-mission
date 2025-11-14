"""
Module pour la génération de PDF professionnels pour les procès-verbaux
Utilise reportlab pour créer des PDF structurés et formatés
"""

# Essayer d'importer reportlab avec gestion d'erreur
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from datetime import datetime
import io

# Essayer d'importer python-docx avec gestion d'erreur
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def create_pv_pdf(content, title="Procès-verbal de Mission", author="Responsable Mission"):
    """
    Génère un PDF professionnel pour un procès-verbal
    
    Args:
        content (str): Contenu du procès-verbal en texte
        title (str): Titre du document
        author (str): Nom de l'auteur/responsable
    
    Returns:
        bytes: Contenu du PDF généré
    """
    
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab n'est pas installé. Installez-le avec: pip install reportlab")
    
    # Créer un buffer en mémoire
    buffer = io.BytesIO()
    
    # Configuration du document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Style pour le titre principal
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Style pour les en-têtes de section
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        fontName='Helvetica-Bold',
        textColor=colors.black
    )
    
    # Style pour le texte normal
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        fontName='Helvetica'
    )
    
    # Style pour la signature
    signature_style = ParagraphStyle(
        'Signature',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=6
    )
    
    # Contenu du document
    story = []
    
    # En-tête du document
    story.append(Paragraph(title.upper(), title_style))
    story.append(Spacer(1, 12))
    
    # Date de génération
    date_str = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    story.append(Paragraph(date_str, signature_style))
    story.append(Spacer(1, 20))
    
    # Traitement du contenu
    lines = content.split('\n')
    current_section = ""
    
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
            
        # Détection des titres (commencent par I., II., III., etc. ou A., B., C.)
        if (line.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.')) or 
            line.startswith(('A.', 'B.', 'C.', 'D.', 'E.', 'F.'))):
            story.append(Paragraph(line, heading_style))
        # Détection des sous-titres (numérotés 1., 2., 3.)
        elif line.startswith(('1.', '2.', '3.', '4.', '5.')):
            sub_heading_style = ParagraphStyle(
                'SubHeading',
                parent=normal_style,
                fontSize=12,
                fontName='Helvetica-Bold',
                spaceBefore=10,
                spaceAfter=6
            )
            story.append(Paragraph(line, sub_heading_style))
        # Texte normal
        else:
            story.append(Paragraph(line, normal_style))
    
    # Signature
    story.append(Spacer(1, 30))
    story.append(Paragraph(f"Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}", signature_style))
    story.append(Spacer(1, 40))
    story.append(Paragraph("_" * 30, signature_style))
    story.append(Paragraph(f"<b>{author}</b>", signature_style))
    
    # Construction du PDF
    doc.build(story)
    
    # Récupération du contenu
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data

def create_word_document(content, title="Procès-verbal de Mission"):
    """
    Génère un document Word simple (format RTF) pour compatibilité
    
    Args:
        content (str): Contenu du document
        title (str): Titre du document
    
    Returns:
        str: Contenu RTF du document
    """
    
    rtf_content = f"""{{\\rtf1\\ansi\\deff0
{{\\fonttbl{{\\f0 Times New Roman;}}}}
{{\\colortbl;\\red0\\green0\\blue0;}}
\\f0\\fs24
\\qc\\b {title.upper()}\\b0\\par
\\par
\\qc Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}\\par
\\par
\\ql
{content.replace(chr(10), '\\par ')}
\\par
\\par
\\qc Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}\\par
\\par
\\par
\\qc ________________________________\\par
\\qc\\b Responsable Mission\\b0\\par
}}"""
    
    return rtf_content

def create_docx_document(content, title="Procès-verbal de Mission"):
    """
    Génère un document Word au format .docx
    
    Args:
        content (str): Contenu du document
        title (str): Titre du document
    
    Returns:
        bytes: Contenu du document Word en bytes
    """
    
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx n'est pas installé. Installez-le avec: pip install python-docx")
    
    # Créer un nouveau document
    doc = Document()
    
    # Ajouter le titre
    title_paragraph = doc.add_heading(title.upper(), 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajouter la date de génération
    date_paragraph = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajouter un saut de ligne
    doc.add_paragraph("")
    
    # Ajouter le contenu principal
    # Diviser le contenu en paragraphes
    paragraphs = content.split('\n')
    for paragraph_text in paragraphs:
        if paragraph_text.strip():  # Ignorer les lignes vides
            doc.add_paragraph(paragraph_text.strip())
    
    # Ajouter la signature
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    signature_paragraph = doc.add_paragraph(f"Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}")
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    signature_line = doc.add_paragraph("________________________________")
    signature_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    signature_name = doc.add_paragraph("Responsable Mission")
    signature_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature_name.runs[0].bold = True
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_mission_pdf(html_content, mission_title="Mission Terrain"):
    """
    Génère un PDF professionnel pour une mission à partir du contenu HTML
    
    Args:
        html_content (str): Contenu HTML de la mission
        mission_title (str): Titre de la mission
    
    Returns:
        bytes: Contenu du PDF généré
    """
    
    # Créer un buffer en mémoire
    buffer = io.BytesIO()
    
    # Configuration du document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Style pour le titre principal
    title_style = ParagraphStyle(
        'MissionTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        textColor=colors.darkblue
    )
    
    # Style pour les en-têtes de section
    heading_style = ParagraphStyle(
        'MissionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        fontName='Helvetica-Bold',
        textColor=colors.darkblue
    )
    
    # Style pour le texte normal
    normal_style = ParagraphStyle(
        'MissionNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )
    
    # Style pour les informations importantes
    info_style = ParagraphStyle(
        'MissionInfo',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        fontName='Helvetica-Bold',
        textColor=colors.darkgreen
    )
    
    # Contenu du document
    story = []
    
    # En-tête du document
    story.append(Paragraph(f"🗺️ {mission_title.upper()}", title_style))
    story.append(Spacer(1, 12))
    
    # Date de génération
    date_str = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    story.append(Paragraph(date_str, normal_style))
    story.append(Spacer(1, 20))
    
    # Extraction et traitement du contenu HTML
    import re
    
    # Supprimer les balises HTML et extraire le texte
    text_content = re.sub(r'<[^>]+>', '', html_content)
    text_content = text_content.replace('&nbsp;', ' ')
    text_content = text_content.replace('&amp;', '&')
    text_content = text_content.replace('&lt;', '<')
    text_content = text_content.replace('&gt;', '>')
    
    # Diviser en lignes et traiter
    lines = text_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Détection des titres de section
        if any(keyword in line.lower() for keyword in ['informations', 'itinéraire', 'statistiques', 'planning']):
            story.append(Paragraph(line, heading_style))
        # Détection des informations importantes (avec émojis ou mots-clés)
        elif any(char in line for char in ['📅', '🗓️', '⏰', '🚗', '📍', '⛽']) or any(keyword in line.lower() for keyword in ['jour', 'début', 'durée', 'distance', 'carburant']):
            story.append(Paragraph(line, info_style))
        # Texte normal
        else:
            story.append(Paragraph(line, normal_style))
    
    # Pied de page
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.grey
    )
    story.append(Paragraph("Planificateur de mission terrain - Développé par @Moctar", footer_style))
    
    # Construction du PDF
    doc.build(story)
    
    # Récupération du contenu
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data