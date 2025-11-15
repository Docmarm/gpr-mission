import os
import json
from datetime import datetime, timedelta, time
import time as time_module
from itertools import permutations
import requests
import toml
import re
import unicodedata

import streamlit as st
import pandas as pd

# --------------------------
# CONFIG APP (DOIT ÊTRE EN PREMIER)
# --------------------------
st.set_page_config(
    page_title="Planificateur de mission terrain", 
    layout="wide",
    page_icon="🗺️",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    .main-header h1 { font-size: 2.2rem; margin-bottom: 0.4rem; font-weight: bold; }
    .main-header p { font-size: 1rem; opacity: 0.9; }
    div[data-testid="stSidebarNav"] { display: none; }
    section[data-testid="stSidebarHeader"] { display: none; }
    .stSidebar { padding-top: 8px; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>🚗 Gestion des Missions</h1>
    <p>Plateforme digitale de gestion des véhicules et missions GPR/EMI</p>
</div>
""", unsafe_allow_html=True)

if st.button("⬅️ Retour à l'accueil"):
    st.switch_page("mission_home_page.py")

# Navigation personnalisée dans la sidebar
if st.sidebar.button("🏠 Accueil"):
    st.switch_page("mission_home_page.py")
if st.sidebar.button("📝 Demande de mission"):
    st.session_state.app_mode = "demande"
    st.switch_page("mission_home_page.py")
if st.sidebar.button("🗺️ Planification"):
    pass
st.sidebar.markdown("---")
if st.sidebar.button("⚙️ Admin"):
    st.switch_page("pages/admin.py")

# Import des modules pour l'export PDF et Word
PDF_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    # Vérifier d'abord si reportlab est installé
    import importlib
    importlib.import_module('reportlab')
    REPORTLAB_INSTALLED = True
except ImportError:
    REPORTLAB_INSTALLED = False

try:
    # Vérifier si python-docx est installé
    importlib.import_module('docx')
    DOCX_INSTALLED = True
except ImportError:
    DOCX_INSTALLED = False

try:
    from pdf_generator import create_pv_pdf, create_word_document, create_mission_pdf, create_docx_document
    # Vérifier si les fonctions sont disponibles
    PDF_AVAILABLE = REPORTLAB_INSTALLED
    DOCX_AVAILABLE = DOCX_INSTALLED
    if not PDF_AVAILABLE:
        st.warning("⚠️ Module reportlab non installé. Installez reportlab pour activer l'export PDF.")
    if not DOCX_AVAILABLE:
        st.warning("⚠️ Module python-docx non installé. Installez python-docx pour l'export Word.")
except ImportError as e:
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False
    st.warning(f"⚠️ Module PDF/Word non disponible: {e}. Installez: pip install reportlab python-docx")

from geopy.geocoders import Nominatim
from geopy.extra.rate_limiter import RateLimiter

import folium
from streamlit_folium import st_folium

# --------------------------
# AUTHENTIFICATION
# --------------------------
# INITIALISATION DES VARIABLES DE SESSION
# --------------------------

def setup_keyboard_shortcuts():
    st.markdown(
        "<script>document.addEventListener('keydown', function(e) {if (e.ctrlKey && e.key === 'k') { e.preventDefault(); const inputs = Array.from(document.querySelectorAll('input')); const s = inputs.find(i => i.placeholder && i.placeholder.toLowerCase().includes('recherche')); if (s) s.focus(); } if (e.ctrlKey && e.key === 's') { e.preventDefault(); const btn = document.querySelector('button[kind=\"primary\"]'); if (btn) btn.click(); }});</script>",
        unsafe_allow_html=True
    )
def modern_alert(message, alert_type="info", icon=None):
    cfg = {"info": {"color":"#2196F3","bg":"#E3F2FD","icon":"ℹ️"},"success":{"color":"#4caf50","bg":"#E8F5E9","icon":"✅"},"warning":{"color":"#ff9800","bg":"#FFF3E0","icon":"⚠️"},"error":{"color":"#f44336","bg":"#FFEBEE","icon":"❌"}}
    c = cfg.get(alert_type, cfg["info"])
    ic = icon if icon else c["icon"]
    st.markdown(f"<div style='background:{c['bg']};border-left:4px solid {c['color']};padding:16px 20px;border-radius:12px;margin:12px 0;box-shadow:0 4px 12px rgba(0,0,0,0.1);'><div style='display:flex;align-items:center;gap:12px;'><span style='font-size:24px;'>{ic}</span><span style='color:#2c3e50;font-weight:500;'>{message}</span></div></div>", unsafe_allow_html=True)
def modern_progress(label, value, max_value=100, color="#667eea"):
    pct = (value / max_value) * 100 if max_value > 0 else 0
    st.markdown(f"<div style='margin:16px 0;'><div style='display:flex;justify-content:space-between;margin-bottom:8px;'><span style='font-weight:600;color:#2c3e50;'>{label}</span><span style='font-weight:700;color:{color};'>{value}/{max_value}</span></div><div style='background:#e0e0e0;height:12px;border-radius:10px;overflow:hidden;'><div style='background:linear-gradient(90deg,{color} 0%, #764ba2 100%);width:{pct}%;height:100%;border-radius:10px;transition:width 0.5s ease;'></div></div><div style='text-align:right;font-size:12px;color:#6c757d;margin-top:4px;'>{pct:.1f}%</div></div>", unsafe_allow_html=True)
setup_keyboard_shortcuts()
st.title("🗺️ Planificateur de mission (Moctar)")
st.caption("Optimisation d'itinéraire + planning journalier + carte interactive + édition de rapport")

# --------------------------
# SIDEBAR: KEYS & OPTIONS
# --------------------------
st.sidebar.header("⚙️ Configuration")

# Chargement sécurisé des clés API (Secrets → ENV → config.toml)
graphhopper_api_key = (
    st.secrets.get("api_keys", {}).get("graphhopper")
    or os.getenv("GRAPHOPPER_API_KEY")
)
deepseek_api_key = (
    st.secrets.get("api_keys", {}).get("deepseek")
    or os.getenv("DEEPSEEK_API_KEY")
)

if not graphhopper_api_key or not deepseek_api_key:
    try:
        config = toml.load('config.toml')
        graphhopper_api_key = graphhopper_api_key or config.get('api_keys', {}).get('graphhopper', '')
        deepseek_api_key = deepseek_api_key or config.get('api_keys', {}).get('deepseek', '')
        st.sidebar.success("✅ Clés API chargées depuis config.toml")
    except Exception:
        # Pas de config.toml, on garde les valeurs actuelles (Secrets/ENV)
        pass

if graphhopper_api_key:
    st.sidebar.caption("🔑 Maps prêt")
else:
    st.sidebar.warning("⚠️ Clé GraphHopper absente — fallback activé si besoin")

if deepseek_api_key:
    st.sidebar.caption("🤖 Adja prêt")
else:
    st.sidebar.warning("⚠️ Clé DeepSeek absente — fonctionnalités IA limitées")

st.sidebar.subheader("Calcul des distances")
distance_method = st.sidebar.radio(
    "Méthode de calcul",
    [
        "Auto (OSRM → Automatique → Maps → Géométrique)",
        "Automatique uniquement",
        "OSRM uniquement (rapide)",
        "Géométrique uniquement",
        "Maps uniquement (précis)"
    ],
    index=0
)

use_deepseek_fallback = st.sidebar.checkbox(
    "Utiliser Maps si Automatique échoue", 
    value=True,
    help="Appeler le service de routage Maps si l'estimation automatique échoue"
)

with st.sidebar.expander("Options avancées"):
    # Charger les paramètres par défaut depuis la configuration (sécurisé)
    try:
        secrets_settings = st.secrets.get("settings", {})
    except Exception:
        secrets_settings = {}
    try:
        local_config = toml.load('config.toml') if os.path.exists('config.toml') else {}
    except Exception:
        local_config = {}

    config_speed = secrets_settings.get("default_speed_kmh")
    config_cache = secrets_settings.get("use_cache")
    config_debug = secrets_settings.get("debug_mode")
    config_osrm = secrets_settings.get("osrm_base_url")

    if config_speed is None:
        config_speed = local_config.get('settings', {}).get('default_speed_kmh', 95)
    if config_cache is None:
        config_cache = local_config.get('settings', {}).get('use_cache', True)
    if config_debug is None:
        config_debug = local_config.get('settings', {}).get('debug_mode', False)
    if config_osrm is None:
        config_osrm = local_config.get('settings', {}).get('osrm_base_url', "https://router.project-osrm.org")
    
    default_speed_kmh = st.number_input(
        "Vitesse moyenne (km/h) pour estimations", 
        min_value=20, max_value=120, value=config_speed
    )
    use_cache = st.checkbox("Utiliser le cache pour géocodage", value=config_cache)
    prefer_offline_geocoding = st.checkbox(
        "Prioriser coordonnées locales pour grandes villes",
        value=True,
        key="prefer_offline_geocoding",
        help="Utiliser des coordonnées vérifiées pour grandes villes du Sénégal (ex. Dakar, Louga, Touba)."
    )
    debug_mode = st.checkbox("Mode debug (afficher détails calculs)", value=config_debug)
    osrm_base_url = st.text_input(
        "OSRM base URL",
        value=config_osrm,
        help="Exemple: http://localhost:5000 ou https://router.project-osrm.org"
    )

# --------------------------
# ÉTAT DE SESSION
# --------------------------
if 'planning_results' not in st.session_state:
    st.session_state.planning_results = None

if 'editing_event' not in st.session_state:
    st.session_state.editing_event = None

if 'edit_mode' not in st.session_state:
    st.session_state.edit_mode = False

if 'manual_itinerary' not in st.session_state:
    st.session_state.manual_itinerary = None

# --------------------------
# FONCTIONS CARBURANT ET EMPREINTE CARBONE
# --------------------------

def get_vehicle_types():
    """Retourne les types de véhicules disponibles avec leurs consommations"""
    return {
        "Station-Wagon": {"consumption": 13.0, "fuel_type": "Essence", "co2_factor": 2.31},
        "Berline": {"consumption": 9.5, "fuel_type": "Essence", "co2_factor": 2.31},
        "SUV": {"consumption": 15.0, "fuel_type": "Essence", "co2_factor": 2.31},
        "4x4": {"consumption": 18.0, "fuel_type": "Diesel", "co2_factor": 2.68},
        "Utilitaire": {"consumption": 12.0, "fuel_type": "Diesel", "co2_factor": 2.68},
        "Minibus": {"consumption": 20.0, "fuel_type": "Diesel", "co2_factor": 2.68},
     }

def calculate_fuel_consumption(total_distance_km, vehicle_type):
    """Calcule la consommation de carburant pour un véhicule donné"""
    vehicles = get_vehicle_types()
    if vehicle_type not in vehicles:
        return None
    
    consumption_per_100km = vehicles[vehicle_type]["consumption"]
    fuel_needed = (total_distance_km * consumption_per_100km) / 100
    
    return {
        "fuel_needed_liters": fuel_needed,
        "consumption_per_100km": consumption_per_100km,
        "fuel_type": vehicles[vehicle_type]["fuel_type"]
    }

def calculate_carbon_footprint(fuel_consumption_data, total_distance_km, vehicle_type):
    """Calcule l'empreinte carbone de la mission"""
    vehicles = get_vehicle_types()
    if vehicle_type not in vehicles or not fuel_consumption_data:
        return None
    
    # Facteur d'émission CO2 (kg CO2 par litre de carburant)
    co2_factor = vehicles[vehicle_type]["co2_factor"]
    
    # Calcul des émissions CO2
    co2_emissions_kg = fuel_consumption_data["fuel_needed_liters"] * co2_factor
    co2_emissions_tons = co2_emissions_kg / 1000
    
    # Équivalences pour contextualiser
    trees_needed = co2_emissions_kg / 22  # Un arbre absorbe ~22kg CO2/an
    
    return {
        "co2_emissions_kg": co2_emissions_kg,
        "co2_emissions_tons": co2_emissions_tons,
        "trees_equivalent": trees_needed,
        "fuel_type": fuel_consumption_data["fuel_type"],
        "distance_km": total_distance_km
    }

def estimate_fuel_cost(fuel_consumption_data, fuel_price_per_liter=None):
    """Estime le coût du carburant"""
    if not fuel_consumption_data:
        return None
    
    # Charger les prix depuis la configuration (Streamlit Secrets ou config.toml)
    default_prices = None
    try:
        secrets_settings = st.secrets.get('settings', {})
        fuel_prices = secrets_settings.get('fuel_prices')
    except Exception:
        fuel_prices = None
    if fuel_prices is None:
        try:
            config = toml.load('config.toml') if os.path.exists('config.toml') else {}
            fuel_prices = config.get('settings', {}).get('fuel_prices')
        except Exception:
            fuel_prices = None
    if fuel_prices:
        default_prices = {
            "Essence": fuel_prices.get('essence'),
            "Diesel": fuel_prices.get('diesel')
        }
    else:
        # Prix par défaut au Sénégal (en FCFA) si aucune config disponible
        default_prices = {
            "Essence": 1350,  # FCFA par litre
            "Diesel": 1200    # FCFA par litre
        }
    
    fuel_type = fuel_consumption_data["fuel_type"]
    price = fuel_price_per_liter if fuel_price_per_liter else default_prices.get(fuel_type, 1300)
    
    total_cost = fuel_consumption_data["fuel_needed_liters"] * price
    
    return {
        "total_cost_fcfa": total_cost,
        "price_per_liter": price,
        "fuel_type": fuel_type,
        "liters": fuel_consumption_data["fuel_needed_liters"]
    }

# --------------------------
# FONCTIONS RAPPORT IA ADJA
# --------------------------
def collect_mission_data_for_ai():
    """Collecte toutes les données de mission pour l'IA Adja"""
    if not st.session_state.planning_results:
        return None
    
    results = st.session_state.planning_results
    itinerary = st.session_state.manual_itinerary or results['itinerary']
    
    # Données de base
    mission_data = {
        'sites': results['sites_ordered'],
        'stats': results['stats'],
        'itinerary': itinerary,
        'calculation_method': results.get('calculation_method', 'Non spécifié'),
        'base_location': results.get('base_location', ''),
        'segments_summary': results.get('segments_summary', [])
    }
    
    # Analyse détaillée des activités
    activities = {}
    detailed_activities = []
    
    for day, sdt, edt, desc in itinerary:
        activity_type = "Autre"
        if "Visite" in desc or "Réunion" in desc:
            activity_type = "Visite/Réunion"
        elif "Trajet" in desc or "km" in desc:
            activity_type = "Déplacement"
        elif "Pause" in desc or "Repos" in desc:
            activity_type = "Pause"
        elif "Nuitée" in desc:
            activity_type = "Hébergement"
        
        duration_hours = (edt - sdt).total_seconds() / 3600
        
        if activity_type not in activities:
            activities[activity_type] = 0
        activities[activity_type] += duration_hours
        
        # Détails de chaque activité
        detailed_activities.append({
            'day': day,
            'start_time': sdt.strftime('%H:%M'),
            'end_time': edt.strftime('%H:%M'),
            'duration': duration_hours,
            'type': activity_type,
            'description': desc
        })
    
    mission_data['activities_breakdown'] = activities
    mission_data['detailed_activities'] = detailed_activities
    
    # Ajouter les données enrichies si disponibles
    if hasattr(st.session_state, 'mission_notes'):
        mission_data['mission_notes'] = st.session_state.mission_notes
    if hasattr(st.session_state, 'activity_details'):
        mission_data['activity_details'] = st.session_state.activity_details
    if hasattr(st.session_state, 'mission_context'):
        mission_data['mission_context'] = st.session_state.mission_context
    
    return mission_data

def collect_construction_report_data():
    """Interface pour collecter des données spécifiques au procès-verbal de chantier"""
    st.markdown("### 🏗️ Données pour Procès-Verbal de Chantier")
    
    # Informations générales du chantier
    col1, col2 = st.columns(2)
    
    with col1:
        project_name = st.text_input(
            "🏗️ Nom du projet/chantier",
            placeholder="Ex: Travaux d'entretien PA DAL zone SUD",
            key="project_name"
        )
        
        report_date = st.date_input(
            "📅 Date de la visite",
            value=datetime.now().date(),
            key="report_date"
        )
        
        site_location = st.text_input(
            "📍 Localisation du site",
            placeholder="Ex: Vélingara et Kolda",
            key="site_location"
        )
    
    with col2:
        report_type = st.selectbox(
            "📋 Type de rapport",
            ["Procès-verbal de visite de chantier", "Rapport d'avancement", "Rapport de fin de travaux", "Rapport d'incident"],
            key="construction_report_type"
        )
        
        weather_conditions = st.text_input(
            "🌤️ Conditions météorologiques",
            placeholder="Ex: Ensoleillé, pluvieux, venteux...",
            key="weather_conditions"
        )
    
    # Liste de présence
    st.markdown("### 👥 Liste de Présence")
    
    if 'attendees' not in st.session_state:
        st.session_state.attendees = []
    
    col_add, col_clear = st.columns([3, 1])
    with col_add:
        new_attendee_name = st.text_input("Nom", key="new_attendee_name")
        new_attendee_structure = st.text_input("Structure/Entreprise", key="new_attendee_structure")
        new_attendee_function = st.text_input("Fonction", key="new_attendee_function")
    
    with col_clear:
        st.write("")  # Espacement
        st.write("")  # Espacement
        if st.button("➕ Ajouter"):
            if new_attendee_name and new_attendee_structure:
                st.session_state.attendees.append({
                    'nom': new_attendee_name,
                    'structure': new_attendee_structure,
                    'fonction': new_attendee_function
                })
                st.rerun()
        
        if st.button("🗑️ Vider"):
            st.session_state.attendees = []
            st.rerun()
    
    # Affichage de la liste
    if st.session_state.attendees:
        st.markdown("**Participants enregistrés :**")
        for i, attendee in enumerate(st.session_state.attendees):
            st.write(f"{i+1}. **{attendee['nom']}** - {attendee['structure']} ({attendee['fonction']})")
    
    # Intervenants dans le projet
    st.markdown("### 🏢 Différents Intervenants dans le Projet")
    
    col1, col2 = st.columns(2)
    with col1:
        master_contractor = st.text_input(
            "🏗️ Maître d'ouvrage",
            placeholder="Ex: Sonatel",
            key="master_contractor"
        )
        
        main_contractor = st.text_input(
            "🔧 Entreprise principale",
            placeholder="Ex: Koné Construction",
            key="main_contractor"
        )
    
    with col2:
        project_manager = st.text_input(
            "👨‍💼 Maître d'œuvre",
            placeholder="Ex: Sonatel",
            key="project_manager"
        )
        
        supervisor = st.text_input(
            "👷‍♂️ Superviseur/Contrôleur",
            placeholder="Ex: SECK CONS",
            key="supervisor"
        )
    
    # Documents contractuels
    st.markdown("### 📄 Documents Contractuels")
    
    if 'contract_documents' not in st.session_state:
        st.session_state.contract_documents = []
    
    col_doc1, col_doc2, col_doc3, col_add_doc = st.columns([2, 2, 2, 1])
    
    with col_doc1:
        doc_name = st.text_input("Document", key="doc_name")
    with col_doc2:
        doc_holder = st.text_input("Porteur", key="doc_holder")
    with col_doc3:
        doc_comments = st.text_input("Commentaires", key="doc_comments")
    with col_add_doc:
        st.write("")  # Espacement
        if st.button("➕", key="add_doc"):
            if doc_name and doc_holder:
                st.session_state.contract_documents.append({
                    'document': doc_name,
                    'porteur': doc_holder,
                    'commentaires': doc_comments
                })
                st.rerun()
    
    if st.session_state.contract_documents:
        st.markdown("**Documents enregistrés :**")
        for i, doc in enumerate(st.session_state.contract_documents):
            st.write(f"• **{doc['document']}** - Porteur: {doc['porteur']} - {doc['commentaires']}")
    
    # Respect du planning
    st.markdown("### ⏰ Respect du Planning")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        start_notification = st.date_input(
            "📅 Notification démarrage",
            key="start_notification"
        )
        
        contractual_delay = st.number_input(
            "⏱️ Délai contractuel (jours)",
            min_value=0,
            value=40,
            key="contractual_delay"
        )
    
    with col2:
        remaining_delay = st.number_input(
            "⏳ Délai restant (jours)",
            min_value=0,
            value=0,
            key="remaining_delay"
        )
        
        progress_percentage = st.slider(
            "📊 Avancement global (%)",
            min_value=0,
            max_value=100,
            value=50,
            key="progress_percentage"
        )
    
    with col3:
        planning_status = st.selectbox(
            "📈 État du planning",
            ["En avance", "Dans les temps", "En retard", "Critique"],
            index=2,
            key="planning_status"
        )
    
    # Observations détaillées par site
    st.markdown("### 🔍 Observations Détaillées par Site")
    
    if st.session_state.planning_results:
        sites = st.session_state.planning_results['sites_ordered']
        
        for i, site in enumerate(sites):
            st.markdown(f"#### 📍 Site de {site['Ville']}")
            
            # Observations par catégorie
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**🏢 Agence commerciale :**")
                agency_work = st.text_area(
                    "Travaux réalisés",
                    placeholder="Ex: Aucun des travaux prévus n'a été réalisé...",
                    height=80,
                    key=f"agency_work_{i}"
                )
                
                st.markdown("**🏗️ Bâtiment technique :**")
                technical_work = st.text_area(
                    "État des travaux techniques",
                    placeholder="Ex: Travaux de carrelage de façade et réhabilitation des toilettes...",
                    height=80,
                    key=f"technical_work_{i}"
                )
            
            with col2:
                st.markdown("**🏠 Logement du gardien :**")
                guard_housing = st.text_area(
                    "État du logement",
                    placeholder="Ex: Mécanisme de la chasse anglaise installé mais non fonctionnel...",
                    height=80,
                    key=f"guard_housing_{i}"
                )
                
                st.markdown("**🚪 Façade de l'agence :**")
                facade_work = st.text_area(
                    "Travaux de façade",
                    placeholder="Ex: Corriger les portes qui ne se ferment pas...",
                    height=80,
                    key=f"facade_work_{i}"
                )
            
            # Poste de garde
            st.markdown("**🛡️ Poste de garde :**")
            guard_post = st.text_area(
                "État du poste de garde",
                placeholder="Ex: Peinture du poste de garde non conforme...",
                height=68,
                key=f"guard_post_{i}"
            )
    
    # Observations générales et recommandations
    st.markdown("### 📝 Observations Générales et Recommandations")
    
    general_observations = st.text_area(
        "🔍 Constat général",
        placeholder="Ex: Lors des visites de chantier, plusieurs constats majeurs ont été relevés concernant la qualité d'exécution...",
        height=120,
        key="general_observations"
    )
    
    recommendations = st.text_area(
        "💡 Recommandations",
        placeholder="Ex: Il est impératif que KONE CONSTRUCTION mette en place un dispositif correctif immédiat...",
        height=120,
        key="construction_recommendations"
    )
    
    # Informations du rapporteur
    st.markdown("### ✍️ Informations du Rapporteur")
    
    col1, col2 = st.columns(2)
    
    with col1:
        reporter_name = st.text_input(
            "👤 Nom du rapporteur",
            placeholder="Ex: Moctar TALL",
            key="reporter_name"
        )
        
        report_location = st.text_input(
            "📍 Lieu de rédaction",
            placeholder="Ex: Dakar",
            key="report_location"
        )
    
    with col2:
        reporter_function = st.text_input(
            "💼 Fonction",
            placeholder="Ex: Ingénieur Projet",
            key="reporter_function"
        )
        
        report_completion_date = st.date_input(
            "📅 Date de finalisation",
            value=datetime.now().date(),
            key="report_completion_date"
        )
    
    return {
        'project_info': {
            'project_name': project_name,
            'report_date': report_date,
            'site_location': site_location,
            'report_type': report_type,
            'weather_conditions': weather_conditions
        },
        'attendees': st.session_state.attendees,
        'stakeholders': {
            'master_contractor': master_contractor,
            'main_contractor': main_contractor,
            'project_manager': project_manager,
            'supervisor': supervisor
        },
        'contract_documents': st.session_state.contract_documents,
        'planning': {
            'start_notification': start_notification,
            'contractual_delay': contractual_delay,
            'remaining_delay': remaining_delay,
            'progress_percentage': progress_percentage,
            'planning_status': planning_status
        },
        'observations': {
            'general_observations': general_observations,
            'recommendations': recommendations
        },
        'reporter': {
            'reporter_name': reporter_name,
            'reporter_function': reporter_function,
            'report_location': report_location,
            'report_completion_date': report_completion_date
        }
    }

def collect_enhanced_mission_data():
    """Interface pour collecter des données enrichies sur la mission"""
    st.markdown("### 📝 Informations détaillées sur la mission")
    
    # Contexte général de la mission
    col1, col2 = st.columns(2)
    
    with col1:
        mission_objective = st.text_area(
            "🎯 Objectif principal de la mission",
            placeholder="Ex: Audit des agences régionales, formation du personnel, prospection commerciale...",
            height=100,
            key="mission_objective"
        )
        
        mission_participants = st.text_input(
            "👥 Participants à la mission",
            placeholder="Ex: Jean Dupont (Chef de projet), Marie Martin (Analyste)...",
            key="mission_participants"
        )
    
    with col2:
        mission_budget = st.number_input(
            "💰 Budget alloué (FCFA)",
            min_value=0,
            value=0,
            step=10000,
            key="mission_budget"
        )
        
        mission_priority = st.selectbox(
            "⚡ Priorité de la mission",
            ["Faible", "Normale", "Élevée", "Critique"],
            index=1,
            key="mission_priority"
        )
    
    # Notes par site/activité
    st.markdown("### 📋 Notes détaillées par site")
    
    if st.session_state.planning_results:
        sites = st.session_state.planning_results['sites_ordered']
        
        if 'activity_details' not in st.session_state:
            st.session_state.activity_details = {}
        
        for i, site in enumerate(sites):
            # Utilisation d'un container au lieu d'un expander pour éviter l'imbrication
            st.markdown(f"### 📍 {site['Ville']} - {site['Type']} ({site['Activité']})")
            with st.container():
                col_notes, col_details = st.columns(2)
                
                with col_notes:
                    notes = st.text_area(
                        "📝 Notes et observations",
                        placeholder="Décrivez ce qui s'est passé, les résultats obtenus, les difficultés rencontrées...",
                        height=120,
                        key=f"notes_{i}"
                    )
                    
                    success_level = st.select_slider(
                        "✅ Niveau de réussite",
                        options=["Échec", "Partiel", "Satisfaisant", "Excellent"],
                        value="Satisfaisant",
                        key=f"success_{i}"
                    )
                
                with col_details:
                    contacts_met = st.text_input(
                        "🤝 Personnes rencontrées",
                        placeholder="Noms et fonctions des contacts",
                        key=f"contacts_{i}"
                    )
                    
                    outcomes = st.text_area(
                        "🎯 Résultats obtenus",
                        placeholder="Accords signés, informations collectées, problèmes identifiés...",
                        height=80,
                        key=f"outcomes_{i}"
                    )
                    
                    follow_up = st.text_input(
                        "📅 Actions de suivi",
                        placeholder="Prochaines étapes, rendez-vous programmés...",
                        key=f"follow_up_{i}"
                    )
                
                # Stocker les détails
                st.session_state.activity_details[f"site_{i}"] = {
                    'site_name': site['Ville'],
                    'site_type': site['Type'],
                    'activity': site['Activité'],
                    'notes': notes,
                    'success_level': success_level,
                    'contacts_met': contacts_met,
                    'outcomes': outcomes,
                    'follow_up': follow_up
                }
    
    # Observations générales
    st.markdown("### 🔍 Observations générales")
    
    col_obs1, col_obs2 = st.columns(2)
    
    with col_obs1:
        challenges = st.text_area(
            "⚠️ Difficultés rencontrées",
            placeholder="Problèmes logistiques, retards, obstacles imprévus...",
            height=100,
            key="challenges"
        )
        
        lessons_learned = st.text_area(
            "📚 Leçons apprises",
            placeholder="Ce qui a bien fonctionné, ce qu'il faut améliorer...",
            height=100,
            key="lessons_learned"
        )
    
    with col_obs2:
        recommendations = st.text_area(
            "💡 Recommandations",
            placeholder="Suggestions pour les prochaines missions...",
            height=100,
            key="mission_recommendations"
        )
        
        overall_satisfaction = st.select_slider(
            "😊 Satisfaction globale",
            options=["Très insatisfait", "Insatisfait", "Neutre", "Satisfait", "Très satisfait"],
            value="Satisfait",
            key="overall_satisfaction"
        )
    
    # Stocker le contexte de mission
    st.session_state.mission_context = {
        'objective': mission_objective,
        'participants': mission_participants,
        'budget': mission_budget,
        'priority': mission_priority,
        'challenges': challenges,
        'lessons_learned': lessons_learned,
        'recommendations': recommendations,
        'overall_satisfaction': overall_satisfaction
    }
    
    return True

def ask_interactive_questions():
    """Pose des questions interactives pour orienter le rapport"""
    st.markdown("### 🤖 Questions pour personnaliser votre rapport")
    
    questions_data = {}
    
    # Questions sur le type de rapport souhaité
    col1, col2 = st.columns(2)
    
    with col1:
        report_focus = st.multiselect(
            "🎯 Sur quoi souhaitez-vous que le rapport se concentre ?",
            ["Résultats obtenus", "Efficacité opérationnelle", "Aspects financiers", 
             "Relations clients", "Problèmes identifiés", "Opportunités découvertes",
             "Performance de l'équipe", "Logistique et organisation"],
            default=["Résultats obtenus", "Efficacité opérationnelle"],
            key="report_focus"
        )
        
        target_audience = st.selectbox(
            "👥 Qui va lire ce rapport ?",
            ["Direction générale", "Équipe projet", "Clients", "Partenaires", 
             "Équipe terrain", "Conseil d'administration"],
            key="target_audience"
        )
    
    with col2:
        report_length = st.selectbox(
            "📄 Longueur souhaitée du rapport",
            ["Court (1-2 pages)", "Moyen (3-5 pages)", "Détaillé (5+ pages)"],
            index=1,
            key="report_length"
        )
        
        include_metrics = st.checkbox(
            "📊 Inclure des métriques et KPIs",
            value=True,
            key="include_metrics"
        )
    
    # Questions spécifiques selon le contexte
    st.markdown("**Questions spécifiques :**")
    
    col3, col4 = st.columns(2)
    
    with col3:
        highlight_successes = st.checkbox(
            "🏆 Mettre en avant les succès",
            value=True,
            key="highlight_successes"
        )
        
        discuss_challenges = st.checkbox(
            "⚠️ Discuter des défis en détail",
            value=True,
            key="discuss_challenges"
        )
        
        future_planning = st.checkbox(
            "🔮 Inclure la planification future",
            value=True,
            key="future_planning"
        )
    
    with col4:
        cost_analysis = st.checkbox(
            "💰 Analyser les coûts en détail",
            value=False,
            key="cost_analysis"
        )
        
        time_efficiency = st.checkbox(
            "⏱️ Analyser l'efficacité temporelle",
            value=True,
            key="time_efficiency"
        )
        
        stakeholder_feedback = st.checkbox(
            "💬 Inclure les retours des parties prenantes",
            value=False,
            key="stakeholder_feedback"
        )
    
    # Question ouverte pour personnalisation
    specific_request = st.text_area(
        "✨ Y a-t-il des aspects spécifiques que vous souhaitez voir dans le rapport ?",
        placeholder="Ex: Comparaison avec la mission précédente, focus sur un site particulier, analyse d'un problème spécifique...",
        height=80,
        key="specific_request"
    )
    
    questions_data = {
        'report_focus': report_focus,
        'target_audience': target_audience,
        'report_length': report_length,
        'include_metrics': include_metrics,
        'highlight_successes': highlight_successes,
        'discuss_challenges': discuss_challenges,
        'future_planning': future_planning,
        'cost_analysis': cost_analysis,
        'time_efficiency': time_efficiency,
        'stakeholder_feedback': stakeholder_feedback,
        'specific_request': specific_request
    }
    
    return questions_data

def generate_enhanced_ai_report(mission_data, questions_data, api_key):
    """Génère un rapport de mission amélioré via l'IA Adja DeepSeek"""
    try:
        # Construction du prompt amélioré
        prompt = build_enhanced_report_prompt(mission_data, questions_data)
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Ajuster max_tokens selon la longueur demandée
        max_tokens_map = {
            "Court (1-2 pages)": 2000,
            "Moyen (3-5 pages)": 4000,
            "Détaillé (5+ pages)": 6000
        }
        
        max_tokens = max_tokens_map.get(questions_data.get('report_length', 'Moyen (3-5 pages)'), 4000)
        
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": max_tokens
        }
        
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers=headers,
            json=data,
            timeout=90
        )
        
        if response.status_code == 200:
            result = response.json()
            content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
            return content
        else:
            st.error(f"Erreur API DeepSeek: {response.status_code}")
            return None
            
    except Exception as e:
        st.error(f"Erreur lors de la génération: {str(e)}")
        return None

def build_enhanced_report_prompt(mission_data, questions_data):
    """Construit un prompt amélioré orienté activités pour la génération de rapport"""
    
    stats = mission_data['stats']
    sites = mission_data['sites']
    activities = mission_data['activities_breakdown']
    detailed_activities = mission_data.get('detailed_activities', [])
    mission_context = mission_data.get('mission_context', {})
    activity_details = mission_data.get('activity_details', {})
    
    # Construction des informations détaillées sur les activités
    activities_info = ""
    if activity_details:
        activities_info = "\nDÉTAILS DES ACTIVITÉS PAR SITE:\n"
        for site_key, details in activity_details.items():
            if details.get('notes') or details.get('outcomes'):
                activities_info += f"\n📍 {details['site_name']} ({details['site_type']}):\n"
                activities_info += f"   - Activité: {details['activity']}\n"
                if details.get('notes'):
                    activities_info += f"   - Notes: {details['notes']}\n"
                if details.get('contacts_met'):
                    activities_info += f"   - Contacts: {details['contacts_met']}\n"
                if details.get('outcomes'):
                    activities_info += f"   - Résultats: {details['outcomes']}\n"
                if details.get('success_level'):
                    activities_info += f"   - Niveau de réussite: {details['success_level']}\n"
                if details.get('follow_up'):
                    activities_info += f"   - Suivi: {details['follow_up']}\n"
    
    # Contexte de mission
    context_info = ""
    if mission_context:
        context_info = f"\nCONTEXTE DE LA MISSION:\n"
        if mission_context.get('objective'):
            context_info += f"- Objectif: {mission_context['objective']}\n"
        if mission_context.get('participants'):
            context_info += f"- Participants: {mission_context['participants']}\n"
        if mission_context.get('budget') and mission_context['budget'] > 0:
            context_info += f"- Budget: {mission_context['budget']:,} FCFA\n"
        if mission_context.get('priority'):
            context_info += f"- Priorité: {mission_context['priority']}\n"
        if mission_context.get('challenges'):
            context_info += f"- Défis: {mission_context['challenges']}\n"
        if mission_context.get('lessons_learned'):
            context_info += f"- Leçons apprises: {mission_context['lessons_learned']}\n"
        if mission_context.get('overall_satisfaction'):
            context_info += f"- Satisfaction globale: {mission_context['overall_satisfaction']}\n"
    
    # Focus du rapport selon les réponses
    focus_areas = questions_data.get('report_focus', [])
    focus_instruction = ""
    if focus_areas:
        focus_instruction = f"\nLE RAPPORT DOIT SE CONCENTRER PARTICULIÈREMENT SUR: {', '.join(focus_areas)}"
    
    # Instructions spécifiques
    specific_instructions = []
    if questions_data.get('highlight_successes'):
        specific_instructions.append("- Mettre en évidence les succès et réalisations")
    if questions_data.get('discuss_challenges'):
        specific_instructions.append("- Analyser en détail les défis rencontrés")
    if questions_data.get('future_planning'):
        specific_instructions.append("- Inclure des recommandations pour l'avenir")
    if questions_data.get('cost_analysis'):
        specific_instructions.append("- Fournir une analyse détaillée des coûts")
    if questions_data.get('time_efficiency'):
        specific_instructions.append("- Analyser l'efficacité temporelle de la mission")
    if questions_data.get('stakeholder_feedback'):
        specific_instructions.append("- Intégrer les retours des parties prenantes")
    if questions_data.get('include_metrics'):
        specific_instructions.append("- Inclure des métriques et indicateurs de performance")
    
    instructions_text = "\n".join(specific_instructions) if specific_instructions else ""
    
    prompt = f"""Tu es un expert en rédaction de rapports de mission professionnels. Génère un rapport détaillé et orienté ACTIVITÉS (pas trajets) en français.

DONNÉES DE BASE:
- Durée totale: {stats['total_days']} jour(s)
- Distance totale: {stats['total_km']:.1f} km
- Temps de visite total: {stats['total_visit_hours']:.1f} heures
- Nombre de sites visités: {len([s for s in sites if s.get('Type') != 'Base'])}
- Sites visités: {', '.join([s.get('Ville') for s in sites if s.get('Type') != 'Base'])}
- Méthode de calcul: {mission_data['calculation_method']}

RÉPARTITION DES ACTIVITÉS:
{chr(10).join([f"- {act}: {hours:.1f}h" for act, hours in activities.items()])}

{context_info}

{activities_info}

PARAMÈTRES DU RAPPORT:
- Public cible: {questions_data.get('target_audience', 'Direction générale')}
- Longueur: {questions_data.get('report_length', 'Moyen (3-5 pages)')}
{focus_instruction}

INSTRUCTIONS SPÉCIFIQUES:
{instructions_text}

DEMANDE SPÉCIALE:
{questions_data.get('specific_request', 'Aucune demande spéciale')}

STRUCTURE REQUISE:
1. 📋 RÉSUMÉ EXÉCUTIF
2. 🎯 OBJECTIFS ET CONTEXTE
3. 📍 DÉROULEMENT DES ACTIVITÉS (focus principal)
   - Détail par site avec résultats obtenus
   - Personnes rencontrées et échanges
   - Succès et difficultés par activité
4. 📊 ANALYSE DES RÉSULTATS
   - Objectifs atteints vs prévus
   - Indicateurs de performance
   - Retour sur investissement
5. 🔍 OBSERVATIONS ET ENSEIGNEMENTS
6. 💡 RECOMMANDATIONS ET ACTIONS DE SUIVI
7. 📈 CONCLUSION ET PERSPECTIVES

IMPORTANT: 
- Concentre-toi sur les ACTIVITÉS et leurs RÉSULTATS, pas sur les trajets
- Utilise les données détaillées fournies pour chaque site
- Adopte un ton professionnel adapté au public cible
- Structure clairement avec des titres et sous-titres
- Inclus des métriques concrètes quand disponibles"""

    return prompt

def build_report_prompt(mission_data, report_type, tone, include_recommendations,
                       include_risks, include_costs, include_timeline, custom_context):
    stats = mission_data['stats']
    sites = mission_data['sites']
    activities = mission_data['activities_breakdown']
    prompt = f"""Tu es un expert en rédaction de rapports de mission professionnels. 

DONNÉES DE LA MISSION:
- Durée totale: {stats['total_days']} jour(s)
- Distance totale: {stats['total_km']:.1f} km
- Temps de visite total: {stats['total_visit_hours']:.1f} heures
- Nombre de sites visités: {len([s for s in sites if s.get('Type') != 'Base'])}
- Sites visités: {', '.join([s.get('Ville') for s in sites if s.get('Type') != 'Base'])}
- Méthode de calcul: {mission_data['calculation_method']}

RÉPARTITION DES ACTIVITÉS:
{chr(10).join([f"- {act}: {hours:.1f}h" for act, hours in activities.items()])}

CONTEXTE SUPPLÉMENTAIRE:
{custom_context if custom_context else "Aucun contexte spécifique fourni"}

INSTRUCTIONS:
- Type de rapport: {report_type}
- Ton: {tone}
- Inclure recommandations: {'Oui' if include_recommendations else 'Non'}
- Inclure analyse des risques: {'Oui' if include_risks else 'Non'}
- Inclure analyse des coûts: {'Oui' if include_costs else 'Non'}
- Inclure timeline détaillée: {'Oui' if include_timeline else 'Non'}

Génère un rapport complet et structuré en français, avec:
1. Résumé exécutif
2. Objectifs et contexte
3. Déroulement de la mission
4. Résultats et observations
5. Analyse des performances (temps, distances, efficacité)
{"6. Recommandations pour l'avenir" if include_recommendations else ""}
{"7. Analyse des risques identifiés" if include_risks else ""}
{"8. Analyse des coûts et budget" if include_costs else ""}
{"9. Timeline détaillée des activités" if include_timeline else ""}
10. Conclusion

Utilise un style {tone.lower()} et structure le rapport avec des titres clairs et des sections bien organisées."""
    return prompt

def build_mission_request_prompt(mission_data):
    stats = mission_data['stats']
    sites = mission_data['sites']
    start_date = st.session_state.planning_results.get('start_date') if hasattr(st.session_state,'planning_results') else datetime.now().date()
    end_date = start_date + timedelta(days=stats.get('total_days',1)-1)
    cities = [s.get('Ville') for s in sites if s.get('Type')!='Base']
    prompt = f"""Tu es un assistant administratif. Génère une proposition de demande de mission basée sur ces données:
- Durée: {stats['total_days']} jour(s)
- Distance: {stats['total_km']:.1f} km
- Sites: {', '.join(cities)}
- Période: {start_date} → {end_date}

Réponds uniquement en JSON avec les clés:
objet, justification, participants, taches, budget_estime_fcfa, risques, approbateurs, carburant_litres, vehicule, lieu, date_depart, date_retour.

Utilise des valeurs réalistes et concises. Dates au format YYYY-MM-DD."""
    return prompt

def generate_mission_request_ai_prefill(mission_data, api_key):
    if not api_key or not mission_data:
        return None
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"model": "deepseek-chat", "messages": [{"role": "user", "content": build_mission_request_prompt(mission_data)}], "temperature": 0.2, "max_tokens": 1200}
    try:
        r = requests.post("https://api.deepseek.com/chat/completions", headers=headers, json=data, timeout=30)
        if r.status_code != 200:
            return None
        text = r.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        s = text.find("{"); e = text.rfind("}") + 1
        if s >= 0 and e > s:
            return json.loads(text[s:e])
        return None
    except Exception:
        return None

def build_mission_request_text(fields):
    def fmt_date(d):
        try:
            return d.strftime("%d/%m/%Y")
        except Exception:
            return str(d)
    lines = []
    lines.append(f"Objet: {fields.get('objet','')}")
    lines.append(f"Lieu: {fields.get('lieu','')}")
    lines.append(f"Période: {fmt_date(fields.get('date_depart'))} → {fmt_date(fields.get('date_retour'))}")
    lines.append(f"Participants: {fields.get('participants','')}")
    lines.append(f"Véhicule: {fields.get('vehicule','')} • Carburant: {fields.get('carburant_litres',0)} L")
    lines.append(f"Budget estimé: {int(fields.get('budget_estime_fcfa',0)):,} FCFA")
    lines.append("")
    lines.append("Justification:")
    lines.append(fields.get('justification',''))
    lines.append("")
    lines.append("Tâches principales:")
    for t in str(fields.get('taches','')).split("\n"):
        t=t.strip()
        if t:
            lines.append(f"- {t}")
    lines.append("")
    if int(fields.get('budget_perdiem_fcfa',0)) or int(fields.get('hotel_driver_fcfa',0)):
        lines.append("Budget détaillé:")
        if int(fields.get('budget_perdiem_fcfa',0)):
            lines.append(f"- Per diem: {int(fields.get('budget_perdiem_fcfa',0)):,} FCFA/jour")
        if int(fields.get('hotel_driver_fcfa',0)):
            lines.append(f"- Hôtel chauffeur: {int(fields.get('hotel_driver_fcfa',0)):,} FCFA/nuit")
    return "\n".join(lines)

def generate_pv_report(mission_data, questions_data, deepseek_api_key):
    """Génère un rapport au format procès-verbal professionnel avec l'IA Adja DeepSeek"""
    
    if not deepseek_api_key:
        return None, "Clé API DeepSeek manquante"
    
    try:
        # Construction du prompt spécialisé pour le procès-verbal
        prompt = f"""Tu es un expert en rédaction de procès-verbaux professionnels pour des projets d'infrastructure. 
Génère un procès-verbal de visite de chantier détaillé et professionnel au format officiel, basé sur les informations suivantes :

INFORMATIONS DE LA MISSION :
- Date : {mission_data.get('date', 'Non spécifiée')}
- Lieu/Site : {mission_data.get('location', 'Non spécifié')}
- Objectif : {mission_data.get('objective', 'Non spécifié')}
- Participants : {', '.join(mission_data.get('participants', []))}
- Durée : {mission_data.get('duration', 'Non spécifiée')}

DÉTAILS SUPPLÉMENTAIRES :
- Contexte : {questions_data.get('context', 'Non spécifié')}
- Observations : {questions_data.get('observations', 'Non spécifiées')}
- Problèmes identifiés : {questions_data.get('issues', 'Aucun')}
- Actions réalisées : {questions_data.get('actions', 'Non spécifiées')}
- Recommandations : {questions_data.get('recommendations', 'Aucune')}

STRUCTURE OBLIGATOIRE DU PROCÈS-VERBAL (respecter exactement cette numérotation) :

I. Cadre général
   1. Cadre général
      - Contexte du projet et objectifs généraux
      - Cadre contractuel et réglementaire
      - Intervenants principaux du projet

   2. Objet de la mission
      - Motif précis de la visite
      - Périmètre d'intervention
      - Objectifs spécifiques de la mission

II. Déroulement de la mission
   A. SITE DE [NOM DU SITE 1]
      - Reconnaître l'équipe présente dans le secteur concerné
      - Vérifier l'avancement des travaux (donner un pourcentage)
      - Faire un bilan, s'enquérir des éventuelles difficultés et contraintes
      - Apprécier la qualité des travaux réalisés
      - Donner des orientations pour la suite des travaux

   B. SITE DE [NOM DU SITE 2] (si applicable)
      - Mêmes points que pour le site 1
      - Spécificités du site

III. Bilan et recommandations
   A. Points positifs constatés
      - Éléments satisfaisants observés
      - Bonnes pratiques identifiées
      - Respect des délais et procédures

   B. Points d'attention et difficultés
      - Problèmes techniques identifiés
      - Contraintes rencontrées
      - Risques potentiels

   C. Recommandations et orientations
      - Actions correctives immédiates
      - Mesures préventives
      - Orientations pour la suite du projet

IV. Observations détaillées
   - Constats techniques précis
   - Mesures et données relevées
   - Documentation photographique (mentionner si applicable)
   - Respect des normes de sécurité et environnementales

CONSIGNES DE RÉDACTION STRICTES :
- Style administratif formel et professionnel
- Terminologie technique précise du BTP/infrastructure
- Phrases courtes et factuelles
- Éviter absolument les opinions personnelles
- Utiliser le passé composé pour les actions réalisées
- Utiliser le présent pour les constats
- Numérotation stricte avec chiffres romains et lettres
- Longueur : 1000-1500 mots minimum
- Inclure des données chiffrées quand possible (pourcentages, mesures, délais)
- Mentionner les normes et références techniques applicables

FORMAT DE PRÉSENTATION :
- Titres en majuscules pour les sections principales
- Sous-titres avec numérotation claire
- Paragraphes structurés avec puces pour les listes
- Conclusion avec date et lieu de rédaction

Le procès-verbal doit être conforme aux standards administratifs et prêt pour validation hiérarchique et archivage officiel."""

        # Appel à l'API DeepSeek
        headers = {
            'Authorization': f'Bearer {deepseek_api_key}',
            'Content-Type': 'application/json'
        }
        
        data = {
            'model': 'deepseek-chat',
            'messages': [
                {
                    'role': 'user',
                    'content': prompt
                }
            ],
            'temperature': 0.3,  # Plus faible pour plus de cohérence
            'max_tokens': 2000
        }
        
        response = requests.post(
            'https://api.deepseek.com/chat/completions',
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                pv_content = result['choices'][0]['message']['content']
                return pv_content, None
            else:
                return None, "Réponse invalide de l'API DeepSeek"
        else:
            return None, f"Erreur API DeepSeek: {response.status_code} - {response.text}"
            
    except requests.exceptions.Timeout:
        return None, "Timeout lors de l'appel à l'API DeepSeek"
    except requests.exceptions.RequestException as e:
        return None, f"Erreur de connexion à l'API DeepSeek: {str(e)}"
    except Exception as e:
        return None, f"Erreur lors de la génération du PV: {str(e)}"

# --------------------------
# FONCTIONS UTILITAIRES
# --------------------------

def test_graphhopper_connection(api_key):
    """Teste la connexion à GraphHopper"""
    if not api_key:
        return False, "Clé API manquante"
    
    try:
        test_points = [[-17.4441, 14.6928], [-17.2732, 14.7167]]
        url = "https://graphhopper.com/api/1/matrix"
        
        data = {
            "points": test_points,
            "profile": "car",
            "out_arrays": ["times", "distances"]
        }
        
        headers = {"Content-Type": "application/json"}
        params = {"key": api_key}
        
        response = requests.post(url, json=data, params=params, headers=headers, timeout=10)
        
        if response.status_code == 200:
            result = response.json()
            if result and "times" in result and "distances" in result:
                distance_km = result['distances'][0][1] / 1000
                # Les "times" de GraphHopper Matrix sont généralement en secondes.
                # Si une valeur est très grande (>100000), on suppose des millisecondes.
                time_val = result['times'][0][1]
                time_min = (time_val / 1000 / 60) if time_val > 100000 else (time_val / 60)
                return True, f"Connexion OK - Test: {distance_km:.1f}km en {time_min:.0f}min"
            else:
                return False, "Réponse invalide de l'API"
        elif response.status_code == 401:
            return False, "Clé API invalide"
        elif response.status_code == 429:
            return False, "Limite de requêtes atteinte"
        else:
            return False, f"Erreur HTTP {response.status_code}"
            
    except Exception as e:
        return False, f"Erreur: {str(e)}"

def _get_matrix_ttl_seconds():
    """TTL pour le cache de matrices (par défaut 24h)."""
    try:
        return int(st.secrets.get("MATRIX_TTL_SECONDS", 24 * 3600))
    except Exception:
        return 24 * 3600

@st.cache_data(ttl=_get_matrix_ttl_seconds(), show_spinner=False)
def improved_graphhopper_duration_matrix(api_key, coords):
    """Calcul de matrice via GraphHopper avec gestion d'erreurs"""
    if not api_key:
        return None, None, "Clé API manquante"
    
    try:
        if len(coords) > 25:
            return None, None, f"Trop de points ({len(coords)}), limite: 25"
        
        # Vérifier que toutes les coordonnées sont valides
        for i, coord in enumerate(coords):
            if not coord or len(coord) != 2:
                return None, None, f"Coordonnées invalides pour le point {i+1}"
            lon, lat = coord
            if not (-180 <= lon <= 180) or not (-90 <= lat <= 90):
                return None, None, f"Coordonnées hors limites pour le point {i+1}: ({lon}, {lat})"
        
        points = [[coord[0], coord[1]] for coord in coords]
        url = "https://graphhopper.com/api/1/matrix"
        data = {
            "points": points,
            "profile": "car",
            "out_arrays": ["times", "distances"]
        }
        headers = {"Content-Type": "application/json"}
        params = {"key": api_key}

        last_error = None
        for attempt in range(3):
            try:
                response = requests.post(url, json=data, params=params, headers=headers, timeout=30)
            except Exception as e:
                last_error = str(e)
                time_module.sleep(1 + attempt)
                continue
            
            if response.status_code == 200:
                result = response.json()
                times = result.get("times")
                distances = result.get("distances")
                if not times or not distances:
                    return None, None, "Données manquantes dans la réponse"
                try:
                    flat_times = [t for row in times for t in row]
                    max_time = max(flat_times) if flat_times else 0
                except Exception:
                    max_time = 0
                durations = [[t / 1000.0 for t in row] for row in times] if max_time > 100000 else times
                return durations, distances, "Succès"
            else:
                if response.status_code == 401:
                    return None, None, "Clé API invalide"
                elif response.status_code == 400:
                    try:
                        error_detail = response.json()
                        error_msg = error_detail.get('message', 'Requête invalide')
                        return None, None, f"Erreur HTTP 400: {error_msg}. Vérifiez que toutes les villes sont valides et géolocalisables."
                    except:
                        return None, None, "Erreur HTTP 400: Requête invalide. Vérifiez que toutes les villes sont valides et géolocalisables."
                elif response.status_code == 429:
                    last_error = "Limite de requêtes atteinte"
                    time_module.sleep(2 + attempt)
                    continue
                elif 500 <= response.status_code < 600:
                    last_error = f"Erreur HTTP {response.status_code}"
                    time_module.sleep(1 + attempt)
                    continue
                else:
                    return None, None, f"Erreur HTTP {response.status_code}"
        return None, None, f"Échec après retries: {last_error or 'Erreur inconnue'}"
    except Exception as e:
        return None, None, f"Erreur: {str(e)}"

@st.cache_data(ttl=_get_matrix_ttl_seconds(), show_spinner=False)
def improved_osrm_duration_matrix(base_url, coords):
    """Calcul de matrice via OSRM Table avec gestion d'erreurs et fallback distances.
    Retourne (durations_sec, distances_m, message).
    """
    if not base_url:
        return None, None, "URL de base OSRM manquante"
    try:
        if len(coords) > 100:
            return None, None, f"Trop de points ({len(coords)}), limite recommandée: 100"
        # Préparer la chaîne de coordonnées pour l'API OSRM
        try:
            coord_str = ';'.join([f"{c[0]},{c[1]}" for c in coords])
        except Exception:
            return None, None, "Coordonnées invalides"
        url = f"{base_url.rstrip('/')}/table/v1/driving/{coord_str}"
        params = {"annotations": "duration,distance"}
        headers = {"Accept": "application/json"}

        last_error = None
        for attempt in range(3):
            try:
                response = requests.get(url, params=params, headers=headers, timeout=30)
            except Exception as e:
                last_error = str(e)
                time_module.sleep(1 + attempt)
                continue
            if response.status_code == 200:
                result = response.json()
                durations = result.get("durations")
                distances = result.get("distances")
                if durations is None:
                    return None, None, "Données manquantes: durations"
                # OSRM fournit les durées en secondes; distances en mètres si activées
                if distances is None:
                    # Fallback distances via Haversine (corrigé 1.2) si non fournies
                    n = len(coords)
                    distances = [[0.0]*n for _ in range(n)]
                    for i in range(n):
                        for j in range(n):
                            if i != j:
                                km = haversine(coords[i][0], coords[i][1], coords[j][0], coords[j][1]) * 1.2
                                distances[i][j] = km * 1000.0
                return durations, distances, "Succès"
            else:
                if response.status_code == 429:
                    last_error = "Limite de requêtes atteinte (OSRM)"
                    time_module.sleep(2 + attempt)
                    continue
                elif 500 <= response.status_code < 600:
                    last_error = f"Erreur HTTP {response.status_code} (OSRM)"
                    time_module.sleep(1 + attempt)
                    continue
                elif response.status_code == 400:
                    try:
                        err = response.json()
                        msg = err.get('message') or err.get('error') or 'Requête invalide (OSRM)'
                        return None, None, f"Erreur HTTP 400: {msg}"
                    except Exception:
                        return None, None, "Erreur HTTP 400: Requête invalide (OSRM)"
                else:
                    return None, None, f"Erreur HTTP {response.status_code} (OSRM)"
        return None, None, f"Échec après retries: {last_error or 'Erreur inconnue'}"
    except Exception as e:
        return None, None, f"Erreur: {str(e)}"

def _get_deepseek_matrix_ttl_seconds():
    """TTL pour le cache de matrices DeepSeek (par défaut 6h)."""
    try:
        return int(st.secrets.get("DEEPSEEK_MATRIX_TTL_SECONDS", 6 * 3600))
    except Exception:
        return 6 * 3600

@st.cache_data(ttl=_get_deepseek_matrix_ttl_seconds(), show_spinner=False)
def improved_deepseek_estimate_matrix(cities, api_key, debug=False):
    """Estimation via DeepSeek avec distances exactes"""
    if not api_key:
        return None, "DeepSeek non disponible"
    
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        prompt = f"""Tu es un expert en transport routier au Sénégal. Calcule les durées ET distances de trajet routier entre ces {len(cities)} villes: {', '.join(cities)}

DISTANCES EXACTES PAR ROUTE (À UTILISER - BIDIRECTIONNELLES):
- Dakar ↔ Thiès: 70 km (55-65 min)
- Dakar ↔ Saint-Louis: 270 km (2h45-3h15)
- Dakar ↔ Kaolack: 190 km (2h-2h30)
- Thiès ↔ Saint-Louis: 200 km (2h-2h30)
- Thiès ↔ Kaolack: 120 km (1h15-1h30)
- Saint-Louis ↔ Kaolack: 240 km (2h30-3h)

IMPORTANT: Les distances sont identiques dans les deux sens (A→B = B→A).

Réponds uniquement en JSON:
{{
  "durations_minutes": [[matrice {len(cities)}x{len(cities)}]],
  "distances_km": [[matrice {len(cities)}x{len(cities)}]]
}}"""

        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1,
            "max_tokens": 2000
        }

        last_error = None
        for attempt in range(3):
            try:
                response = requests.post(
                    "https://api.deepseek.com/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=30
                )
            except Exception as e:
                last_error = str(e)
                time_module.sleep(1 + attempt)
                continue
            
            if response.status_code != 200:
                if response.status_code == 429:
                    last_error = "Limite de requêtes atteinte"
                    time_module.sleep(2 + attempt)
                    continue
                elif 500 <= response.status_code < 600:
                    last_error = f"Erreur HTTP {response.status_code}"
                    time_module.sleep(1 + attempt)
                    continue
                else:
                    return None, f"Erreur API: {response.status_code}"

            result = response.json()
            text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
            start = text.find("{")
            end = text.rfind("}") + 1
            if start >= 0 and end > start:
                try:
                    json_str = text[start:end]
                    parsed = json.loads(json_str)
                    minutes_matrix = parsed.get("durations_minutes", [])
                    km_matrix = parsed.get("distances_km", [])
                    seconds_matrix = [[int(m) * 60 for m in row] for row in minutes_matrix]
                    distances_matrix = [[int(km * 1000) for km in row] for row in km_matrix]
                    return (seconds_matrix, distances_matrix), "Succès DeepSeek"
                except Exception as parse_err:
                    return None, f"Format invalide: {parse_err}"
            else:
                last_error = "Réponse non JSON"
                time_module.sleep(1 + attempt)
                continue
        return None, f"Échec après retries: {last_error or 'Erreur inconnue'}"
    
    except Exception as e:
        return None, f"Erreur: {str(e)}"

def build_ics_from_itinerary(itinerary, start_date, mission_title="Mission Terrain"):
    """Construit un fichier ICS à partir du planning."""
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Mission Planner//EN"
    ]
    now_str = datetime.now().strftime("%Y%m%dT%H%M%S")
    for idx, (day, sdt, edt, desc) in enumerate(itinerary):
        dtstart = sdt.strftime("%Y%m%dT%H%M%S")
        dtend = edt.strftime("%Y%m%dT%H%M%S")
        uid = f"mission-{now_str}-{idx}@planner"
        summary = f"{mission_title} - {desc}"
        lines.extend([
            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTAMP:{now_str}",
            f"DTSTART:{dtstart}",
            f"DTEND:{dtend}",
            f"SUMMARY:{summary}",
            f"DESCRIPTION:Jour {day}",
            "END:VEVENT"
        ])
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines)

@st.cache_data(show_spinner=False)
def _get_geolocator():
    """Ressource Geopy réutilisable (évite recréation à chaque appel)."""
    # Utilise un cache ressource pour conserver l'instance et respecter le rate limiter
    @st.cache_resource(show_spinner=False)
    def _cached_geolocator():
        return Nominatim(user_agent="mission-planner-senegal/2.0", timeout=10)
    return _cached_geolocator()

def _get_rate_limited_geocode():
    """Retourne une fonction geocode rate-limitée avec retries."""
    @st.cache_resource(show_spinner=False)
    def _cached_rate_limiter():
        geolocator = _get_geolocator()
        return RateLimiter(
            geolocator.geocode,
            min_delay_seconds=1,
            max_retries=3,
            error_wait_seconds=2,
            swallow_exceptions=False
        )
    return _cached_rate_limiter()

def _normalize_city_key(name: str) -> str:
    """Normalise un nom de ville pour les correspondances hors-ligne (sans accents/espaces/ponctuations)."""
    if not isinstance(name, str):
        return ""
    s = name.strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = s.encode("ascii", "ignore").decode("ascii")
    # Unifier les variantes de 'saint', 'ste', 'st'
    s = re.sub(r"\bste\b", "saint", s)
    s = re.sub(r"\bst\b", "saint", s)
    # Supprimer tout sauf alphanumérique
    s = re.sub(r"[^a-z0-9]", "", s)
    return s

# Coordonnées approximatives de grandes villes du Sénégal (lon, lat)
SENEGAL_CITY_COORDS = {
    _normalize_city_key("Dakar"): (-17.4677, 14.7167),
    _normalize_city_key("Pikine"): (-17.3570, 14.7642),
    _normalize_city_key("Touba"): (-15.8833, 14.8667),
    _normalize_city_key("Thiès"): (-16.9359, 14.7910),
    _normalize_city_key("Thies"): (-16.9359, 14.7910),  # variante sans accent
    _normalize_city_key("Saint-Louis"): (-16.4896, 16.0179),
    _normalize_city_key("Saint Louis"): (-16.4896, 16.0179),
    _normalize_city_key("St-Louis"): (-16.4896, 16.0179),
    _normalize_city_key("Kaolack"): (-16.0726, 14.1475),
    _normalize_city_key("Ziguinchor"): (-16.2719, 12.5833),
    _normalize_city_key("Louga"): (-16.2167, 15.6167),
    _normalize_city_key("Tambacounda"): (-13.6673, 13.7703),
    _normalize_city_key("Diourbel"): (-16.2348, 14.6550),
    _normalize_city_key("Fatick"): (-16.4150, 14.3390),
    _normalize_city_key("Kolda"): (-14.9500, 12.8833),
    _normalize_city_key("Matam"): (-13.2554, 15.6559),
    _normalize_city_key("Kaffrine"): (-15.5508, 14.1059),
    _normalize_city_key("Kedougou"): (-12.1742, 12.5556),
    _normalize_city_key("Kédougou"): (-12.1742, 12.5556),
    _normalize_city_key("Sedhiou"): (-15.5569, 12.7081),
    _normalize_city_key("Sédhiou"): (-15.5569, 12.7081),
    _normalize_city_key("Rufisque"): (-17.2729, 14.7158),
    _normalize_city_key("Mbour"): (-16.9600, 14.4361),
    _normalize_city_key("Richard-Toll"): (-15.6994, 16.4611),
    _normalize_city_key("Richard Toll"): (-15.6994, 16.4611),
}

def _offline_lookup_city_coords(city: str):
    key = _normalize_city_key(city)
    return SENEGAL_CITY_COORDS.get(key)

def _graphhopper_geocode(city: str):
    """Fallback via GraphHopper Geocoding API si disponible.
    Sélectionne en priorité les lieux de type city/town/village au Sénégal.
    """
    try:
        gh_key = globals().get("graphhopper_api_key")
        if not gh_key:
            return None
        url = "https://graphhopper.com/api/1/geocode"
        params = {
            "q": f"{city}, Senegal",
            "locale": "fr",
            "limit": 8,
            "key": gh_key,
        }
        resp = requests.get(url, params=params, timeout=10)
        if resp.status_code != 200:
            return None
        hits = (resp.json().get("hits") or [])
        if not hits:
            return None
        def is_sn(h):
            country = (h.get("country") or h.get("countrycode") or "").lower()
            return country in ("senegal", "sénégal", "sn")
        def is_place(h):
            return (h.get("osm_key", "").lower() == "place" and (h.get("osm_value", "").lower() in ("city", "town", "village")))
        chosen = None
        for h in hits:
            if is_sn(h) and is_place(h):
                chosen = h
                break
        if not chosen:
            for h in hits:
                if is_sn(h):
                    chosen = h
                    break
        chosen = chosen or hits[0]
        pt = chosen.get("point") or {}
        lat = pt.get("lat")
        lng = pt.get("lng")
        if lat is None or lng is None:
            return None
        # Valide que le point est dans un bbox raisonnable pour le Sénégal
        if not (-17.8 <= float(lng) <= -11.0 and 12.0 <= float(lat) <= 16.9):
            return None
        return (float(lng), float(lat))
    except Exception:
        return None

def _geocode_city_senegal_raw(city: str):
    """Implémentation brute sans cache (avec ressource + retries). Privilégie coordonnées locales si activé."""
    if not city or not isinstance(city, str) or not city.strip():
        return None

    # Option: préférer coordonnées locales pour grandes villes (fiabilité)
    try:
        prefer_offline = st.session_state.get("prefer_offline_geocoding", True)
    except Exception:
        prefer_offline = True
    if prefer_offline:
        offline = _offline_lookup_city_coords(city)
        if offline:
            return offline

    last_error = None
    for attempt in range(3):  # 3 tentatives
        try:
            rate_limited = _get_rate_limited_geocode()
            query = f"{city}, Sénégal" if "sénégal" not in city.lower() else city
            
            # Essai principal: ciblé Sénégal
            loc = rate_limited(query, language="fr", country_codes="SN")
            
            # Fallback: requête générale
            if not loc:
                loc = rate_limited(city, language="fr")
            
            if loc:
                lon, lat = (loc.longitude, loc.latitude)
                # Vérifie le type renvoyé par Nominatim si disponible
                try:
                    raw = getattr(loc, "raw", {}) or {}
                    place_type = (raw.get("type") or "").lower()
                    if place_type and place_type not in ("city", "town", "village", "hamlet"):
                        raise ValueError(f"Résultat non-ville: {place_type}")
                except Exception:
                    pass
                # Valide que les coordonnées sont plausibles pour le Sénégal
                if not (-17.8 <= float(lon) <= -11.0 and 12.0 <= float(lat) <= 16.9):
                    raise ValueError("Coordonnées hors Sénégal")
                return (lon, lat)
        
        except ConnectionRefusedError as e:
            last_error = f"Connexion refusée au service de géocodage. Vérifiez votre connexion ou l'état du service. ({e})"
            time_module.sleep(1 + attempt) # Attente progressive
            continue
        except Exception as e:
            last_error = e
            time_module.sleep(1 + attempt) # Attente progressive
            continue
    # Fallback 1: GraphHopper (si clé dispo)
    gh_coords = _graphhopper_geocode(city)
    if gh_coords:
        st.warning(f"Géocodage Nominatim indisponible/inexact. Fallback GraphHopper utilisé pour {city}.")
        return gh_coords

    # Fallback 2: Dictionnaire hors-ligne
    offline = _offline_lookup_city_coords(city)
    if offline:
        st.info(f"Mode hors-ligne: coordonnées vérifiées utilisées pour {city}.")
        return offline

    st.error(f"Erreur de géocodage persistante pour {city} après plusieurs tentatives: {last_error}")
    return None

def _get_geocode_ttl_seconds():
    # TTL configurable via secrets, défaut 7 jours
    try:
        return int(st.secrets.get("GEOCODE_TTL_SECONDS", 7 * 24 * 3600))
    except Exception:
        return 7 * 24 * 3600

@st.cache_data(ttl=_get_geocode_ttl_seconds(), show_spinner=False)
def _geocode_city_senegal_cached(city: str):
    return _geocode_city_senegal_raw(city)

def geocode_city_senegal(city: str, use_cache: bool = True):
    """Géocode une ville au Sénégal, avec cache TTL et ressource partagée.

    Args:
        city: Nom de la ville
        use_cache: Active ou non le cache des résultats
    """
    return _geocode_city_senegal_cached(city) if use_cache else _geocode_city_senegal_raw(city)

def solve_tsp_fixed_start_end(matrix):
    """Résout le TSP avec départ et arrivée fixes"""
    n = len(matrix)
    if n <= 2:
        return list(range(n))
    
    if n > 10:
        st.warning("Plus de 10 sites: heuristique voisin + 2-opt")
        nn_path = solve_tsp_nearest_neighbor(matrix)
        improved_path = two_opt_fixed_start_end(nn_path, matrix)
        return improved_path
    
    nodes = list(range(1, n-1))
    best_order = None
    best_time = float("inf")
    
    for perm in permutations(nodes):
        total_time = matrix[0][perm[0]]
        for i in range(len(perm)-1):
            total_time += matrix[perm[i]][perm[i+1]]
        total_time += matrix[perm[-1]][n-1]
        
        if total_time < best_time:
            best_time = total_time
            best_order = perm
    
    best_path = [0] + list(best_order) + [n-1]
    # Lissage via 2-opt si des incohérences existent
    try:
        best_path = two_opt_fixed_start_end(best_path, matrix)
    except Exception:
        pass
    return best_path

def solve_tsp_nearest_neighbor(matrix):
    """Heuristique du plus proche voisin"""
    n = len(matrix)
    unvisited = set(range(1, n-1))
    path = [0]
    current = 0
    
    while unvisited:
        nearest = min(unvisited, key=lambda x: matrix[current][x])
        path.append(nearest)
        unvisited.remove(nearest)
        current = nearest
    
    path.append(n-1)
    return path

def path_cost(path, matrix):
    total = 0
    for i in range(len(path)-1):
        total += matrix[path[i]][path[i+1]]
    return total

def two_opt_fixed_start_end(path, matrix):
    """Amélioration locale 2-opt en conservant départ (0) et arrivée (n-1)"""
    if not path or len(path) < 4:
        return path
    improved = True
    while improved:
        improved = False
        for i in range(1, len(path)-2):
            for k in range(i+1, len(path)-1):
                new_path = path[:i] + path[i:k+1][::-1] + path[k+1:]
                if path_cost(new_path, matrix) < path_cost(path, matrix):
                    path = new_path
                    improved = True
                    break
            if improved:
                break
    return path

# OR-Tools integration for advanced optimization (TSP with fixed start/end)
try:
    from ortools.constraint_solver import pywrapcp, routing_enums_pb2
    ORTOOLS_AVAILABLE = True
except Exception:
    ORTOOLS_AVAILABLE = False


def solve_tsp_ortools_fixed_start_end(matrix, service_times=None, time_limit_s=5):
    """Optimise l'ordre via OR-Tools (TSP), départ 0 et arrivée n-1 fixés.
    - Utilise la matrice des durées (secondes)
    - Peut intégrer les durées de visite (service_times, en secondes) sur chaque site
    - Retourne un chemin sous forme d'indices: [0, ..., n-1]
    """
    n = len(matrix)
    if n <= 2:
        return list(range(n))

    # Si OR-Tools indisponible, fallback sur l'implémentation TSP existante
    if not ORTOOLS_AVAILABLE:
        return solve_tsp_fixed_start_end(matrix)

    try:
        manager = pywrapcp.RoutingIndexManager(n, 1, 0, n-1)
        routing = pywrapcp.RoutingModel(manager)

        def time_callback(from_index, to_index):
            from_node = manager.IndexToNode(from_index)
            to_node = manager.IndexToNode(to_index)
            travel = int(matrix[from_node][to_node] or 0)
            service = 0
            if service_times and isinstance(service_times, (list, tuple)):
                if 0 <= from_node < n and from_node not in (0, n - 1):
                    st_value = service_times[from_node]
                    try:
                        service = int(float(st_value) if st_value is not None else 0)
                    except Exception:
                        service = 0
            return travel + service

        transit_cb = routing.RegisterTransitCallback(time_callback)
        routing.SetArcCostEvaluatorOfAllVehicles(transit_cb)

        # Dimension de temps simple (horizon large)
        routing.AddDimension(
            transit_cb,
            0,            # marge
            24 * 3600,    # horizon max (24h)
            True,         # cumul de départ à 0
            "Time"
        )

        search_params = pywrapcp.DefaultRoutingSearchParameters()
        search_params.first_solution_strategy = routing_enums_pb2.FirstSolutionStrategy.PATH_CHEAPEST_ARC
        search_params.local_search_metaheuristic = routing_enums_pb2.LocalSearchMetaheuristic.GUIDED_LOCAL_SEARCH
        search_params.time_limit.seconds = int(time_limit_s)

        solution = routing.SolveWithParameters(search_params)
        if solution:
            index = routing.Start(0)
            path = []
            while not routing.IsEnd(index):
                node = manager.IndexToNode(index)
                path.append(node)
                index = solution.Value(routing.NextVar(index))
            path.append(manager.IndexToNode(index))
            return path
    except Exception:
        pass

    # Fallback en cas d'échec
    return solve_tsp_fixed_start_end(matrix)

def haversine(lon1, lat1, lon2, lat2):
    """Calcule la distance géodésique entre deux points en kilomètres"""
    from math import radians, sin, cos, sqrt, atan2
    R = 6371.0
    dlon = radians(lon2 - lon1)
    dlat = radians(lat2 - lat1)
    a = sin(dlat/2)**2 + cos(radians(lat1))*cos(radians(lat2))*sin(dlon/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    return R * c

def haversine_fallback_matrix(coords, kmh=95.0):
    """Calcule une matrice basée sur distances géodésiques"""
    
    n = len(coords)
    durations = [[0.0]*n for _ in range(n)]
    distances = [[0.0]*n for _ in range(n)]
    
    for i in range(n):
        for j in range(n):
            if i != j:
                km = haversine(coords[i][0], coords[i][1], coords[j][0], coords[j][1])
                # Facteur de correction pour tenir compte des routes réelles
                km *= 1.2
                hours = km / kmh
                # Retourner les durées en secondes (cohérent avec GraphHopper)
                durations[i][j] = hours * 3600
                # Retourner les distances en mètres (cohérent avec GraphHopper)
                distances[i][j] = km * 1000
    
    return durations, distances

def optimize_route_with_ai(sites, coords, base_location=None, api_key=None):
    """
    Optimise l'ordre des sites en utilisant l'IA Adja DeepSeek
    
    Args:
        sites: Liste des sites avec leurs informations
        coords: Liste des coordonnées correspondantes
        base_location: Point de départ/arrivée (optionnel)
        api_key: Clé API DeepSeek
    
    Returns:
        tuple: (ordre_optimal, success, message)
    """
    if not api_key:
        return list(range(len(sites))), False, "Clé API DeepSeek manquante"
    
    try:
        # Préparer les données des sites pour l'IA Adja
        sites_info = []
        for i, site in enumerate(sites):
            site_data = {
                "index": i,
                "ville": site.get("Ville", f"Site {i}"),
                "type": site.get("Type", "Non spécifié"),
                "activite": site.get("Activité", "Non spécifié"),
                "duree": site.get("Durée (h)", 1.0),
                "coordonnees": coords[i] if i < len(coords) else None
            }
            sites_info.append(site_data)
        
        # Construire le prompt pour l'IA Adja
        prompt = f"""Tu es un expert en optimisation d'itinéraires au Sénégal. 

MISSION: Optimise l'ordre de visite des sites suivants pour minimiser le temps de trajet total.

SITES À VISITER:
"""
        
        for site in sites_info:
            coord_str = f"({site['coordonnees'][0]:.4f}, {site['coordonnees'][1]:.4f})" if site['coordonnees'] else "Coordonnées inconnues"
            prompt += f"- Site {site['index']}: {site['ville']} - {site['type']} - {site['activite']} ({site['duree']}h) - {coord_str}\n"
        
        if base_location:
            prompt += f"\nPOINT DE DÉPART/ARRIVÉE: {base_location}\n"
        
        prompt += """
CONTRAINTES:
- Minimiser la distance totale de trajet
- Tenir compte de la géographie du Sénégal
- Considérer les types d'activités (regrouper les activités similaires si logique)
- Optimiser pour un trajet efficace

RÉPONSE ATTENDUE:
Fournis UNIQUEMENT la liste des indices dans l'ordre optimal, séparés par des virgules.
Exemple: 0,2,1,3,4

Ne fournis AUCUNE explication, juste la séquence d'indices."""

        # Appel à l'API DeepSeek
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 100
        }
        
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result["choices"][0]["message"]["content"].strip()
            
            # Parser la réponse de l'IA Adja
            try:
                # Extraire les indices de la réponse
                indices_str = ai_response.split('\n')[0].strip()
                indices = [int(x.strip()) for x in indices_str.split(',')]
                
                # Vérifier que tous les indices sont valides
                if len(indices) == len(sites) and set(indices) == set(range(len(sites))):
                    return indices, True, "Optimisation IA Adja réussie"
                else:
                    # Fallback: ordre original si la réponse IA Adja est invalide
                    return list(range(len(sites))), False, f"Réponse IA Adja invalide: {ai_response[:100]}..."
                    
            except (ValueError, IndexError) as e:
                return list(range(len(sites))), False, f"Erreur parsing réponse IA Adja: {str(e)}"
        
        else:
            return list(range(len(sites))), False, f"Erreur API DeepSeek: {response.status_code}"
            
    except requests.exceptions.Timeout:
        return list(range(len(sites))), False, "Timeout API DeepSeek"
    except requests.exceptions.RequestException as e:
        return list(range(len(sites))), False, f"Erreur réseau: {str(e)}"
    except Exception as e:
        return list(range(len(sites))), False, f"Erreur inattendue: {str(e)}"

def schedule_itinerary(coords, sites, order, segments_summary,
                       start_date, start_activity_time, end_activity_time,
                       start_travel_time, end_travel_time,
                       use_lunch, lunch_start_time, lunch_end_time,
                       use_prayer, prayer_start_time, prayer_duration_min,
                       max_days=0, tolerance_hours=1.0, base_location=None, 
                       stretch_days=False, end_day_early_threshold=1.5,
                       allow_weekend_travel=True, allow_weekend_activities=True,
                       lunch_duration_min=60):
    """Génère le planning détaillé avec horaires différenciés pour activités et voyages"""
    sites_ordered = [sites[i] for i in order]
    coords_ordered = [coords[i] for i in order]
    
    current_datetime = datetime.combine(start_date, start_travel_time)  # Start with travel time
    day_end_time = datetime.combine(start_date, end_travel_time)  # End with travel time
    day_count = 1
    itinerary = []
    
    # Suivi des pauses par jour pour éviter les doublons
    daily_lunch_added = {}  # {day_count: bool}
    daily_prayer_added = {}  # {day_count: bool}
    
    total_km = 0
    total_visit_hours = 0
    
    for idx, site in enumerate(sites_ordered):
        # Handle travel to this site (except for first site)
        if idx > 0:
            # Weekend skip for travel if disabled
            if not allow_weekend_travel:
                while current_datetime.weekday() >= 5:
                    itinerary.append((day_count, current_datetime, datetime.combine(current_datetime.date(), end_travel_time), "⛱️ Week-end (pas de voyage)"))
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_travel_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
            seg_idx = idx - 1
            seg_idx = idx - 1
            if seg_idx < len(segments_summary):
                seg = segments_summary[seg_idx]
                travel_sec = seg.get("duration", 0)
                travel_km = seg.get("distance", 0) / 1000.0
                
                # Debug: Afficher les valeurs reçues
                if debug_mode:
                    st.info(f"🔍 Debug Segment {seg_idx}: travel_sec={travel_sec}, travel_km={travel_km:.2f}")
                
                # Si les données sont nulles, utiliser des valeurs par défaut simples
                if travel_sec <= 0:
                    travel_sec = 3600  # 1 heure par défaut
                    if debug_mode:
                        st.warning(f"🔍 travel_sec était ≤ 0, fixé à 3600s (1h)")
                if travel_km <= 0:
                    travel_km = 50  # 50 km par défaut
                    if debug_mode:
                        st.warning(f"🔍 travel_km était ≤ 0, fixé à 50km")
                
                total_km += travel_km
                
                travel_duration = timedelta(seconds=int(travel_sec))
                travel_end = current_datetime + travel_duration
                
                from_city = sites_ordered[idx-1]['Ville']
                to_city = site['Ville']
                
                # Format travel time for display
                travel_hours = travel_sec / 3600
                if travel_hours >= 1:
                    travel_time_str = f"{travel_hours:.1f}h"
                else:
                    travel_minutes = travel_sec / 60
                    travel_time_str = f"{travel_minutes:.0f}min"
                
                travel_desc = f"🚗 {from_city} → {to_city} ({travel_km:.1f} km, {travel_time_str})"
                
                # Check if travel extends beyond travel hours
                travel_end_time = datetime.combine(current_datetime.date(), end_travel_time)
                
                if travel_end > travel_end_time:
                    # Travel extends beyond allowed hours - split across days
                    itinerary.append((day_count, current_datetime, travel_end_time, "🏁 Fin de journée"))
                    prev_site = sites_ordered[idx-1]
                    prev_city = prev_site['Ville']
                    prev_overnight_allowed = prev_site.get('Possibilité de nuitée', True)
                    if prev_overnight_allowed:
                        itinerary.append((day_count, travel_end_time, travel_end_time, f"🏨 Nuitée à {prev_city}"))
                    else:
                        # Pas d'hébergement autorisé à la ville précédente -> avertissement + nuitée de repli
                        itinerary.append((day_count, travel_end_time, travel_end_time, f"⚠️ Déplacement nécessaire - pas d'hébergement à {prev_city}"))
                        fallback_city = None
                        for j in range(idx, len(sites_ordered)):
                            if sites_ordered[j].get('Possibilité de nuitée', True):
                                fallback_city = sites_ordered[j]['Ville']
                                break
                        if not fallback_city and base_location:
                            fallback_city = base_location
                        if fallback_city:
                            itinerary.append((day_count, travel_end_time, travel_end_time, f"🏨 Nuitée à {fallback_city}"))
                    
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_travel_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                    travel_end = current_datetime + travel_duration
                
                # Handle lunch break during travel
                lunch_window_start = datetime.combine(current_datetime.date(), lunch_start_time) if use_lunch else None
                lunch_window_end = datetime.combine(current_datetime.date(), lunch_end_time) if use_lunch else None
                
                travel_added = False
                
                if use_lunch and lunch_window_start and lunch_window_end and not daily_lunch_added.get(day_count, False):
                    if current_datetime < lunch_window_end and travel_end > lunch_window_start:
                        # Si l'arrivée se situe dans la fenêtre de déjeuner, placer la pause à l'arrivée
                        if lunch_window_start <= travel_end <= lunch_window_end:
                            # Ajouter le trajet en une seule fois jusqu'à l'arrivée
                            itinerary.append((day_count, current_datetime, travel_end, travel_desc))
                            travel_added = True
                            
                            # Placer le déjeuner immédiatement à l'arrivée
                            lunch_time = max(travel_end, lunch_window_start)
                            lunch_end_time_actual = min(lunch_time + timedelta(minutes=lunch_duration_min), lunch_window_end)
                            desc_text = f"🍽️ Déjeuner (≤{lunch_duration_min} min)"
                            if use_prayer and prayer_start_time and not daily_prayer_added.get(day_count, False):
                                prayer_window_start = datetime.combine(lunch_time.date(), prayer_start_time)
                                prayer_window_end = prayer_window_start + timedelta(hours=2)
                                if lunch_time < prayer_window_end and lunch_end_time_actual > prayer_window_start:
                                    desc_text = f"🍽️ Déjeuner (≤{lunch_duration_min} min) + 🙏 Prière (≤{prayer_duration_min} min)"
                                    daily_prayer_added[day_count] = True
                            itinerary.append((day_count, lunch_time, lunch_end_time_actual, desc_text))
                            daily_lunch_added[day_count] = True
                            
                            # Mettre à jour l'heure courante à la fin du déjeuner
                            current_datetime = lunch_end_time_actual
                            # Le trajet est terminé, éviter tout ajout résiduel
                            travel_end = current_datetime
                        else:
                            # Sinon, conserver l’ancienne logique (pause pendant le trajet)
                            lunch_time = max(current_datetime, lunch_window_start)
                            lunch_end_time_actual = min(lunch_time + timedelta(minutes=lunch_duration_min), lunch_window_end)
                            
                            # Ajouter la partie de trajet avant la pause si nécessaire
                            if lunch_time > current_datetime:
                                itinerary.append((day_count, current_datetime, lunch_time, travel_desc))
                                travel_added = True
                            
                            # Ajouter la pause déjeuner
                            itinerary.append((day_count, lunch_time, lunch_end_time_actual, f"🍽️ Déjeuner (≤{lunch_duration_min} min)"))
                            daily_lunch_added[day_count] = True
                            
                            # Reprendre le trajet après la pause
                            current_datetime = lunch_end_time_actual
                            remaining_travel = travel_end - lunch_time
                            if remaining_travel.total_seconds() < 0:
                                remaining_travel = timedelta(seconds=0)
                            travel_end = current_datetime + remaining_travel
                
                # Handle prayer break during travel (only if no lunch break)
                elif use_prayer and prayer_start_time and not daily_prayer_added.get(day_count, False):
                    prayer_window_start = datetime.combine(current_datetime.date(), prayer_start_time)
                    prayer_window_end = prayer_window_start + timedelta(hours=2)
                    
                    if current_datetime < prayer_window_end and travel_end > prayer_window_start:
                        prayer_time = max(current_datetime, prayer_window_start)
                        prayer_end_time = prayer_time + timedelta(minutes=prayer_duration_min)
                        
                        if prayer_end_time > prayer_window_end:
                            prayer_end_time = prayer_window_end
                        
                        # Add travel before prayer if needed
                        if prayer_time > current_datetime:
                            itinerary.append((day_count, current_datetime, prayer_time, travel_desc))
                            travel_added = True
                        
                        # Add prayer break
                        itinerary.append((day_count, prayer_time, prayer_end_time, "🙏 Prière (≤20 min)"))
                        daily_prayer_added[day_count] = True  # Marquer la prière comme ajoutée pour ce jour
                        current_datetime = prayer_end_time
                        
                        # Recalculate remaining travel time
                        remaining_travel = travel_end - prayer_time
                        travel_end = current_datetime + remaining_travel
                
                # Add remaining travel time (include post-break remaining travel if any)
                if current_datetime < travel_end:
                    itinerary.append((day_count, current_datetime, travel_end, travel_desc))
                
                current_datetime = travel_end
        
        visit_hours = float(site.get("Durée (h)", 0)) if site.get("Durée (h)") else 0
        
        if visit_hours > 0:
            # Weekend skip for activities if disabled
            if not allow_weekend_activities:
                while current_datetime.weekday() >= 5:
                    itinerary.append((day_count, current_datetime, datetime.combine(current_datetime.date(), end_activity_time), "⛱️ Week-end (pas d'activités)"))
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
            total_visit_hours += visit_hours
            visit_duration = timedelta(hours=visit_hours)
            visit_end = current_datetime + visit_duration
            
            type_site = site.get('Type', 'Site')
            activite = site.get('Activité', 'Visite')
            city = site['Ville'].upper()
            
            visit_desc = f"{city} – {activite}"
            if type_site not in ["Base"]:
                visit_desc = f"{city} – Visite {type_site}"
            
            # Check if visit extends beyond activity hours
            activity_end_time = datetime.combine(current_datetime.date(), end_activity_time)
            tolerance_end_time = activity_end_time + timedelta(hours=tolerance_hours)
            
            # Vérifier si l'activité peut continuer (nouvelle option)
            can_continue = site.get('Peut continuer', False)  # Par défaut False
            
            # Vérifier si la nuitée est possible dans cette zone
            overnight_allowed = site.get('Possibilité de nuitée', True)  # Par défaut True
            
            # Handle visit that extends beyond activity hours
            if visit_end > activity_end_time:
                # Si l'activité se termine dans le seuil de tolérance, elle peut continuer le même jour
                if visit_end <= tolerance_end_time and can_continue:
                    # L'activité continue sur le même jour malgré le dépassement
                    pass  # Pas de division, traitement normal
                elif can_continue and overnight_allowed:
                    # L'activité dépasse le seuil de tolérance et peut être divisée, ET la nuitée est autorisée
                    if current_datetime < activity_end_time:
                        # Add partial visit for current day
                        itinerary.append((day_count, current_datetime, activity_end_time, f"{visit_desc} (à continuer)"))
                    
                    # End current day
                    itinerary.append((day_count, activity_end_time, activity_end_time, "🏁 Fin de journée"))
                    if overnight_allowed:
                        itinerary.append((day_count, activity_end_time, activity_end_time, f"🏨 Nuitée à {city}"))
                    else:
                        itinerary.append((day_count, activity_end_time, activity_end_time, f"⚠️ Déplacement nécessaire - pas d'hébergement à {city}"))
                        # Nuitée de repli vers un site prochain autorisé ou la base
                        fallback_city = None
                        for j in range(idx+1, len(sites_ordered)):
                            if sites_ordered[j].get('Possibilité de nuitée', True):
                                fallback_city = sites_ordered[j]['Ville']
                                break
                        if not fallback_city and base_location:
                            fallback_city = base_location
                        if fallback_city:
                            itinerary.append((day_count, activity_end_time, activity_end_time, f"🏨 Nuitée à {fallback_city}"))
                    
                    # Start next day
                    remaining = visit_end - activity_end_time
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                    visit_end = current_datetime + remaining
                    visit_desc = f"Suite {visit_desc}"
                elif can_continue and not overnight_allowed:
                    # L'activité peut continuer mais la nuitée n'est pas autorisée - chercher un site proche avec nuitée
                    # Pour l'instant, on force la fin de l'activité et on ajoute un avertissement
                    visit_end = activity_end_time
                    if current_datetime < activity_end_time:
                        itinerary.append((day_count, current_datetime, activity_end_time, f"{visit_desc} (interrompu - pas de nuitée possible)"))
                    
                    # End current day et chercher un hébergement ailleurs
                    itinerary.append((day_count, activity_end_time, activity_end_time, "🏁 Fin de journée"))
                    itinerary.append((day_count, activity_end_time, activity_end_time, f"⚠️ Déplacement nécessaire - pas d'hébergement à {city}"))
                    
                    # Start next day
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                    # Reprendre l'activité restante le jour suivant
                    remaining_hours = (visit_duration - (activity_end_time - current_datetime)).total_seconds() / 3600
                    if remaining_hours > 0:
                        visit_end = current_datetime + timedelta(hours=remaining_hours)
                        visit_desc = f"Suite {visit_desc}"
                    else:
                        # L'activité était déjà terminée
                        visit_end = current_datetime
                else:
                    # L'activité ne peut pas continuer - la forcer à se terminer à l'heure limite
                    visit_end = activity_end_time
                    if current_datetime >= activity_end_time:
                        # Si on est déjà en dehors des heures, terminer la journée et ajouter la nuitée (avec fallback si nécessaire)
                        itinerary.append((day_count, current_datetime, current_datetime, "🏁 Fin de journée"))
                        if overnight_allowed:
                            itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {city}"))
                        else:
                            itinerary.append((day_count, current_datetime, current_datetime, f"⚠️ Déplacement nécessaire - pas d'hébergement à {city}"))
                            # Chercher une nuitée autorisée dans les sites suivants ou la base
                            fallback_city = None
                            for j in range(idx+1, len(sites_ordered)):
                                if sites_ordered[j].get('Possibilité de nuitée', True):
                                    fallback_city = sites_ordered[j]['Ville']
                                    break
                            if not fallback_city and base_location:
                                fallback_city = base_location
                            if fallback_city:
                                itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {fallback_city}"))
                        
                        # Reporter au jour suivant
                        day_count += 1
                        current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                        day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                        visit_end = current_datetime + visit_duration
            
            # Handle breaks during visit (only if visit fits in current day)
            if visit_end <= activity_end_time:
                lunch_window_start = datetime.combine(current_datetime.date(), lunch_start_time) if use_lunch else None
                lunch_window_end = datetime.combine(current_datetime.date(), lunch_end_time) if use_lunch else None
                
                prayer_window_start = datetime.combine(current_datetime.date(), prayer_start_time) if use_prayer else None
                prayer_window_end = prayer_window_start + timedelta(hours=2) if use_prayer else None
                
                # Check for lunch break during visit — do not split, schedule lunch after visit
                place_lunch_after_visit = False
                if use_lunch and lunch_window_start and lunch_window_end and not daily_lunch_added.get(day_count, False):
                    if current_datetime < lunch_window_end and visit_end > lunch_window_start:
                        place_lunch_after_visit = True
                # If lunch will be placed after the visit, and prayer window overlaps that lunch window,
                # combine prayer with lunch instead of splitting the visit
                combine_prayer_with_lunch = False
                if place_lunch_after_visit and use_prayer and prayer_window_start and prayer_window_end and not daily_prayer_added.get(day_count, False):
                    planned_lunch_start = max(visit_end, lunch_window_start)
                    planned_lunch_end = min(planned_lunch_start + timedelta(minutes=lunch_duration_min), lunch_window_end)
                    if planned_lunch_start < prayer_window_end and planned_lunch_end > prayer_window_start:
                        combine_prayer_with_lunch = True
                
                # Check for prayer break during visit (skip if it will be combined with lunch after visit)
                if use_prayer and prayer_window_start and prayer_window_end and not daily_prayer_added.get(day_count, False) and not combine_prayer_with_lunch:
                    if current_datetime < prayer_window_end and visit_end > prayer_window_start:
                        prayer_time = max(current_datetime, prayer_window_start)
                        prayer_end_time = min(prayer_time + timedelta(minutes=prayer_duration_min), prayer_window_end)
                        
                        # Add visit part before prayer
                        if prayer_time > current_datetime:
                            itinerary.append((day_count, current_datetime, prayer_time, visit_desc))
                        
                        # Add prayer break
                        itinerary.append((day_count, prayer_time, prayer_end_time, "🙏 Prière (≤20 min)"))
                        daily_prayer_added[day_count] = True  # Marquer la prière comme ajoutée pour ce jour
                        
                        # Update timing for remaining visit
                        current_datetime = prayer_end_time
                        remaining_visit = visit_end - prayer_time
                        visit_end = current_datetime + remaining_visit
                        visit_desc = f"Suite {visit_desc}" if prayer_time > current_datetime else visit_desc
            
            # Add final visit segment
            if current_datetime < visit_end:
                itinerary.append((day_count, current_datetime, visit_end, visit_desc))
                current_datetime = visit_end
            
            # Place lunch right after the visit if the window overlapped
            if 'place_lunch_after_visit' in locals() and place_lunch_after_visit and not daily_lunch_added.get(day_count, False):
                lunch_time = max(current_datetime, lunch_window_start)
                if lunch_time < lunch_window_end:
                    lunch_end_time_actual = min(lunch_time + timedelta(minutes=lunch_duration_min), lunch_window_end)
                    desc_text = f"🍽️ Déjeuner (≤{lunch_duration_min} min)"
                    if 'combine_prayer_with_lunch' in locals() and combine_prayer_with_lunch and use_prayer and not daily_prayer_added.get(day_count, False):
                        desc_text = f"🍽️ Déjeuner (≤{lunch_duration_min} min) + 🙏 Prière (≤{prayer_duration_min} min)"
                        daily_prayer_added[day_count] = True
                    itinerary.append((day_count, lunch_time, lunch_end_time_actual, desc_text))
                    daily_lunch_added[day_count] = True
                    current_datetime = lunch_end_time_actual
            
            # Check if we need to end the day early
            time_until_end = (day_end_time - current_datetime).total_seconds() / 3600
            
            # Si on doit étaler, on termine la journée plus tôt pour répartir sur plus de jours
            if stretch_days and day_count < max_days and idx < len(sites_ordered) - 1:
                itinerary.append((day_count, current_datetime, current_datetime, f"🏁 Fin de journée"))
                # Nuitée conditionnelle selon la possibilité
                if overnight_allowed:
                    itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {city}"))
                else:
                    itinerary.append((day_count, current_datetime, current_datetime, f"⚠️ Déplacement nécessaire - pas d'hébergement à {city}"))
                    # Chercher une nuitée autorisée dans les sites suivants ou la base
                    fallback_city = None
                    for j in range(idx+1, len(sites_ordered)):
                        if sites_ordered[j].get('Possibilité de nuitée', True):
                            fallback_city = sites_ordered[j]['Ville']
                            break
                    if not fallback_city and base_location:
                        fallback_city = base_location
                    if fallback_city:
                        itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {fallback_city}"))

                # Démarrer le jour suivant
                day_count += 1
                current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)

            elif time_until_end <= end_day_early_threshold and idx < len(sites_ordered) - 1:
                # End current day and prepare for next day
                itinerary.append((day_count, current_datetime, current_datetime, f"🏁 Fin de journée"))
                # Nuitée conditionnelle selon la possibilité
                if overnight_allowed:
                    itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {city}"))
                else:
                    itinerary.append((day_count, current_datetime, current_datetime, f"⚠️ Déplacement nécessaire - pas d'hébergement à {city}"))
                    # Chercher une nuitée autorisée dans les sites suivants ou la base
                    fallback_city = None
                    for j in range(idx+1, len(sites_ordered)):
                        if sites_ordered[j].get('Possibilité de nuitée', True):
                            fallback_city = sites_ordered[j]['Ville']
                            break
                    if not fallback_city and base_location:
                        fallback_city = base_location
                    if fallback_city:
                        itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {fallback_city}"))
                
                # Start next day
                day_count += 1
                current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
    
    # Add final overnight stay for the last day
    if day_count > 0 and sites_ordered:
        last_site = sites_ordered[-1]
        last_city = last_site['Ville']
        if last_site.get('Possibilité de nuitée', True):
            itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {last_city}"))
        else:
            # Fallback to base_location if overnight is not possible at the last site
            if base_location:
                itinerary.append((day_count, current_datetime, current_datetime, f"🏨 Nuitée à {base_location}"))

    # Add final arrival marker
    if day_count > 0 and sites_ordered:
        last_city = sites_ordered[-1]['Ville'].upper()
        
        itinerary.append((day_count, current_datetime, current_datetime, f"📍 Arrivée {last_city} – Fin de mission"))
    
    # Message d'avertissement si le nombre de jours est dépassé (ne devrait plus arriver avec la nouvelle logique)
    if max_days > 0 and day_count > max_days and not stretch_days:
        st.warning(f"⚠️ L'itinéraire nécessite {day_count} jours, mais le maximum était fixé à {max_days}. Le planning est compressé.")
    
    stats = {
        "total_days": day_count,
        "total_km": total_km,
        "total_visit_hours": total_visit_hours
    }
    
    return itinerary, sites_ordered, coords_ordered, stats

def build_professional_html(itinerary, start_date, stats, sites_ordered, segments_summary=None, speed_kmh=110, mission_title="Mission Terrain", coords_ordered=None, include_map=False, lunch_start_time=None, lunch_end_time=None, lunch_duration_min=60, prayer_start_time=None, prayer_duration_min=20, include_details=True):
    """Génère un HTML professionnel"""
    def fmt_time(dt):
        return dt.strftime("%Hh%M")
    
    def extract_distance_from_desc(desc, speed_kmh_param):
        import re
        # Chercher d'abord le format avec temps réel : "(123.4 km, 2h30)"
        m_with_time = re.search(r"\(([\d\.]+)\s*km,\s*([^)]+)\)", desc)
        if m_with_time:
            km = float(m_with_time.group(1))
            time_str = m_with_time.group(2).strip()
            return f"~{int(km)} km / ~{time_str}"
        
        # Fallback : ancien format avec seulement distance "(123.4 km)"
        m = re.search(r"\(([\d\.]+)\s*km\)", desc)
        if m:
            km = float(m.group(1))
            hours = km / speed_kmh_param
            h = int(hours)
            minutes = int((hours - h) * 60)
            if h > 0:
                time_str = f"{h}h{minutes:02d}"
            else:
                time_str = f"0h{minutes:02d}"
            return f"~{int(km)} km / ~{time_str}"
        return "-"

    by_day = {}
    night_locations = {}
    
    for day, sdt, edt, desc in itinerary:
        by_day.setdefault(day, []).append((sdt, edt, desc))
        
        if "Nuitée à" in desc or "nuitée à" in desc:
            if " à " in desc:
                parts = desc.split(" à ")
                if len(parts) >= 2:
                    city = parts[1].strip().split("(")[0].strip().split(" ")[0]
                    night_locations[day] = city.upper()
        elif "🏨" in desc and ("Nuitée" in desc or "nuitée" in desc):
            # Gestion spécifique pour les descriptions avec emoji 🏨
            if " à " in desc:
                parts = desc.split(" à ")
                if len(parts) >= 2:
                    city = parts[1].strip().split("(")[0].strip().split(" ")[0]
                    night_locations[day] = city.upper()
        elif "installation" in desc.lower() and "nuitée" in desc.lower():
            words = desc.split()
            for i, word in enumerate(words):
                if "installation" in word.lower() and i + 1 < len(words):
                    city = words[i + 1].strip().split("(")[0].strip()
                    night_locations[day] = city.upper()
                    break
        elif "Fin de journée" in desc:
            for _, _, d in reversed(by_day[day]):
                if any(x in d for x in ["VISITE", "Visite", "–"]) and "→" not in d:
                    if "–" in d:
                        city = d.split("–")[0].strip()
                        night_locations[day] = city.upper()
                        break
    
    max_day = max(by_day.keys()) if by_day else 1
    if max_day in night_locations:
        last_events = by_day[max_day]
        if any("Fin de mission" in desc for _, _, desc in last_events):
            for _, _, desc in last_events:
                if "Arrivée" in desc and "Fin de mission" in desc:
                    city = desc.split("Arrivée")[1].split("–")[0].strip()
                    night_locations[max_day] = city

    first_date = start_date
    last_date = start_date + timedelta(days=stats['total_days']-1)
    
    months = ['janvier', 'février', 'mars', 'avril', 'mai', 'juin', 
              'juillet', 'août', 'septembre', 'octobre', 'novembre', 'décembre']
    date_range = f"{first_date.strftime('%d')} → {last_date.strftime('%d')} {months[last_date.month-1]} {last_date.year}"
    
    num_nights = stats['total_days'] - 1 if stats['total_days'] > 1 else 0

    # KPIs et méta
    actual_sites_count = len([s for s in sites_ordered if s.get('Type') != 'Base'])
    distance_km = stats.get('total_km', 0)
    total_visit_hours = stats.get('total_visit_hours', 0)
    route_summary = " → ".join([s.get('Ville', '').upper() for s in sites_ordered if s.get('Ville')])
    gen_date_str = datetime.now().strftime("%d/%m/%Y")
    
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Planning {mission_title} ({date_range})</title>
    <style>
        body {{ font-family: Tahoma, Calibri, 'Segoe UI', sans-serif; margin: 20px; background-color: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; border-radius: 10px; padding: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }}
        h1 {{ text-align: center; color: #2c3e50; margin-bottom: 4px; font-size: 24px; }}
        p.subtitle {{ text-align: center; color: #7f8c8d; margin: 0 0 16px; font-size: 13px; }}
        .kpi-grid {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 12px; margin: 16px 0 18px; }}
        .kpi {{ background:#f8f9fb; border:1px solid #e6e8eb; border-radius:8px; padding:10px 12px; text-align:center; }}
        .kpi-label {{ color:#6c7a89; font-size:12px; }}
        .kpi-value {{ color:#2c3e50; font-weight:bold; font-size:18px; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 10px; font-size: 14px; }}
        th {{ background-color: #34495e; color: white; padding: 12px 8px; text-align: left; font-weight: bold; }}
        td {{ padding: 10px 8px; border-bottom: 1px solid #ddd; vertical-align: top; }}
        tr:nth-child(even) {{ background-color: #f8f9fa; }}
        tr:hover {{ background-color: #e8f4f8; }}
        .jour {{ font-weight: bold; color: #2980b9; background-color: #ecf0f1 !important; }}
        .horaire {{ font-weight: bold; color: #27ae60; white-space: nowrap; }}
        .activite {{ color: #2c3e50; }}
        .mission {{ background-color: #fff3cd; font-weight: bold; }}
        .route {{ color: #7f8c8d; font-style: italic; }}
        .nuit {{ background-color: #d1ecf1; font-weight: bold; color: #0c5460; text-align: center; }}
        .distance {{ color: #e74c3c; font-weight: bold; white-space: nowrap; }}
        .note {{ font-size: 12px; color: #7f8c8d; margin-top: 8px; }}
        .map-section {{ margin-top: 16px; }}
        .section {{ margin-top: 20px; }}
        .section h2 {{ color:#2c3e50; font-size:18px; margin-bottom:8px; }}
        .section ul {{ margin:8px 0 0 18px; }}
        .signatures {{ display:flex; gap:24px; margin-top:16px; }}
        .signature-box {{ flex:1; border:1px dashed #bdc3c7; border-radius:8px; padding:12px; color:#34495e; }}
        .footer {{ margin-top:14px; font-size:12px; color:#7f8c8d; }}
    </style>
</head>
<body>
<div class="container">
    <h1>📋 {mission_title} – {date_range}</h1>
    <p class="subtitle">{stats['total_days']} jour{'s' if stats['total_days']>1 else ''} / {num_nights} nuitée{'s' if num_nights>1 else ''} • Pauses : déjeuner (13h00–14h30 ≤ 1h) & prière (14h00–15h00 ≤ 20 min)</p>

    <div class="kpi-grid">
        <div class="kpi"><div class="kpi-label">Durée</div><div class="kpi-value">{stats['total_days']} j</div></div>
        <div class="kpi"><div class="kpi-label">Distance</div><div class="kpi-value">{distance_km:.1f} km</div></div>
        <div class="kpi"><div class="kpi-label">Sites</div><div class="kpi-value">{actual_sites_count}</div></div>
        <div class="kpi"><div class="kpi-label">Visites</div><div class="kpi-value">{total_visit_hours:.1f} h</div></div>
        <div class="kpi"><div class="kpi-label">Nuitées</div><div class="kpi-value">{num_nights}</div></div>
    </div>

    <table>
        <thead>
            <tr>
                <th style="width: 15%;">JOUR</th>
                <th style="width: 15%;">HORAIRES</th>
                <th style="width: 40%;">ACTIVITÉS</th>
                <th style="width: 15%;">TRANSPORT</th>
                <th style="width: 15%;">NUIT</th>
            </tr>
        </thead>
        <tbody>"""

    for day in sorted(by_day.keys()):
        day_events = by_day[day]
        
        display_events = []
        for sdt, edt, desc in day_events:
            if "Nuitée" not in desc and "Fin de journée" not in desc:
                display_events.append((sdt, edt, desc))
        
        if not display_events:
            continue
            
        day_count = len(display_events)
        night_location = night_locations.get(day, "")
        
        html += f"""
            <!-- JOUR {day} -->"""
        
        for i, (sdt, edt, desc) in enumerate(display_events):
            if "→" in desc and "🚗" in desc:
                activity_class = "route"
                activity_text = desc.replace("🚗 ", "🚗 ")
                transport_info = extract_distance_from_desc(desc, speed_kmh)
            elif any(word in desc.upper() for word in ["VISITE", "AGENCE", "SITE", "CLIENT"]):
                activity_class = "mission"
                activity_text = desc.replace("🏢", "").replace("👥", "").replace("📍", "").replace("🏠", "").strip()
                transport_info = "-"
            elif "déjeuner" in desc.lower() and "prière" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "déjeuner" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "prière" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "installation" in desc.lower() or "arrivée" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "fin" in desc.lower() and "mission" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            else:
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            
            if i == 0:
                html += f"""
            <tr class="jour">
                <td rowspan="{day_count}"><strong>JOUR {day}</strong></td>
                <td class="horaire">{fmt_time(sdt)}–{fmt_time(edt)}</td>
                <td class="{activity_class}">{activity_text}</td>
                <td class="distance">{transport_info}</td>
                <td rowspan="{day_count}" class="nuit">{night_location}</td>
            </tr>"""
            else:
                html += f"""
            <tr>
                <td class="horaire">{fmt_time(sdt)}–{fmt_time(edt)}</td>
                <td class="{activity_class}">{activity_text}</td>
                <td class="distance">{transport_info}</td>
            </tr>"""

    html += f"""
        </tbody>
    </table>

    <p class="note">ℹ️ Distances/temps indicatifs. Les pauses déjeuner et prière sont flexibles et intégrées sans bloquer les activités.</p>
"""

    if include_details:
        html += f"""
    <div class="section">
        <h2>📋 Résumé exécutif</h2>
        <ul>
            <li>Itinéraire: {route_summary}</li>
            <li>Distance: {distance_km:.1f} km; Visites: {total_visit_hours:.1f} h; Sites: {actual_sites_count}</li>
            <li>Période: {date_range}</li>
        </ul>
    </div>
    <div class="section signatures">
        <div class="signature-box">Préparé par: __________________<br/>Fonction: __________________<br/>Date: {gen_date_str}</div>
        <div class="signature-box">Validé par: __________________<br/>Fonction: __________________<br/>Date: {gen_date_str}</div>
    </div>
    <div class="footer">App dev by Moctar TALL • Document généré le {gen_date_str}</div>
"""

    # Intégrer la carte en-dessous du tableau si coords_ordered est fourni (optionnel)
    map_embed_html = ""
    try:
        if include_map and coords_ordered and len(coords_ordered) > 0:
            center_lat = sum(c[1] for c in coords_ordered) / len(coords_ordered)
            center_lon = sum(c[0] for c in coords_ordered) / len(coords_ordered)

            m = folium.Map(location=[center_lat, center_lon], zoom_start=7)

            # Route via OSRM, fallback sur ligne droite
            try:
                coord_str = ";".join([f"{c[0]},{c[1]}" for c in coords_ordered])
                url = f"{osrm_base_url.rstrip('/')}/route/v1/driving/{coord_str}?overview=full&geometries=geojson"
                resp = requests.get(url, timeout=10)
                route_pts = None
                if resp.status_code == 200:
                    data = resp.json()
                    if data.get('routes'):
                        geom = data['routes'][0].get('geometry')
                        if isinstance(geom, dict) and geom.get('coordinates'):
                            route_pts = [[lat, lon] for lon, lat in geom['coordinates']]
                if not route_pts:
                    route_pts = [[c[1], c[0]] for c in coords_ordered]
            except Exception:
                route_pts = [[c[1], c[0]] for c in coords_ordered]
            folium.PolyLine(locations=route_pts, color="blue", weight=3, opacity=0.7).add_to(m)

            # Affichage spécial si départ et arrivée identiques
            n_steps = len(sites_ordered)
            start_end_same = False
            if n_steps >= 2:
                lat0, lon0 = coords_ordered[0][1], coords_ordered[0][0]
                latN, lonN = coords_ordered[-1][1], coords_ordered[-1][0]
                start_end_same = abs(lat0 - latN) < 1e-4 and abs(lon0 - lonN) < 1e-4

            for i, site in enumerate(sites_ordered):
                if i == 0 and start_end_same:
                    bg_color_left = '#2ecc71'  # Vert pour départ
                    bg_color_right = '#e74c3c'  # Rouge pour arrivée
                    html_num = f"""
<div style=\"display:flex; align-items:center; gap:4px;\">
  <div style=\"background-color:{bg_color_left}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">1</div>
  <div style=\"background-color:{bg_color_right}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{n_steps}</div>
</div>
"""
                    folium.Marker(
                        location=[coords_ordered[i][1], coords_ordered[i][0]],
                        popup=f"Étapes 1 et {n_steps}: {site.get('Ville','')}<br>{site.get('Type', '-')}",
                        tooltip=f"Étapes 1 et {n_steps}: {site.get('Ville','')}",
                        icon=folium.DivIcon(
                            icon_size=(36, 28),
                            icon_anchor=(18, 14),
                            html=html_num
                        )
                    ).add_to(m)
                    continue
                if start_end_same and i == n_steps - 1:
                    continue

                bg_color = '#2ecc71' if i == 0 else '#e74c3c' if i == len(sites_ordered)-1 else '#3498db'
                folium.Marker(
                    location=[coords_ordered[i][1], coords_ordered[i][0]],
                    popup=f"Étape {i+1}: {site.get('Ville','')}<br>{site.get('Type', '-')}",
                    tooltip=f"Étape {i+1}: {site.get('Ville','')}",
                    icon=folium.DivIcon(
                        icon_size=(28, 28),
                        icon_anchor=(14, 14),
                        html=f"""
<div style=\"background-color:{bg_color}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{i+1}</div>
"""
                    )
                ).add_to(m)

            map_html = m.get_root().render()
            import base64
            map_b64 = base64.b64encode(map_html.encode('utf-8')).decode('ascii')
            map_embed_html = f"""
    <div class=\"map-section\">
        <h2>🗺️ Carte de l'itinéraire</h2>
        <iframe src=\"data:text/html;base64,{map_b64}\" style=\"width:100%; height:600px; border:none;\"></iframe>
    </div>
"""
    except Exception:
        map_embed_html = ""

    if map_embed_html:
        html += map_embed_html

    html += """
</div>
</body>
</html>"""

    return html

def create_mission_excel(itinerary, start_date, stats, sites_ordered, segments_summary=None, mission_title="Mission Terrain"):
    """
    Génère un fichier Excel professionnel à partir des données de planning
    """
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    
    # Créer un workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Planning Mission"
    
    # Styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2E86AB", end_color="2E86AB", fill_type="solid")
    subheader_font = Font(bold=True, color="2E86AB")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_alignment = Alignment(horizontal='center', vertical='center')
    
    # En-tête principal
    ws.merge_cells('A1:F1')
    ws['A1'] = mission_title
    ws['A1'].font = Font(bold=True, size=16, color="2E86AB")
    ws['A1'].alignment = center_alignment
    
    # Informations générales
    current_row = 3
    ws[f'A{current_row}'] = f"📅 Période: {start_date.strftime('%d/%m/%Y')} → {(start_date + timedelta(days=len(itinerary)-1)).strftime('%d/%m/%Y')}"
    ws[f'A{current_row}'].font = subheader_font
    current_row += 1
    
    ws[f'A{current_row}'] = f"🏃 {stats['total_days']} jour{'s' if stats['total_days'] > 1 else ''} / 0 nuitée • Pauses flexibles : déjeuner (13h00-14h30 ≤ 1h) & prière (14h00-15h00 ≤ 20 min)"
    current_row += 2
    
    # En-têtes du tableau
    headers = ['JOUR', 'HORAIRES', 'ACTIVITÉS', 'TRANSPORT', 'NUIT']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=current_row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_alignment
        cell.border = border
    
    current_row += 1
    
    # Données du planning
    # L'itinéraire est une liste de tuples: (day, start_time, end_time, description)
    current_day = None
    day_start_row = current_row
    
    for event in itinerary:
        day, start_time, end_time, description = event
        
        # Nouvelle journée
        if day != current_day:
            current_day = day
            day_start_row = current_row
            
            # Colonne JOUR
            ws.cell(row=current_row, column=1, value=f"JOUR {day}")
            ws.cell(row=current_row, column=1).font = Font(bold=True)
            ws.cell(row=current_row, column=1).alignment = center_alignment
            ws.cell(row=current_row, column=1).border = border
        else:
            ws.cell(row=current_row, column=1, value="")
            ws.cell(row=current_row, column=1).border = border
        
        # Colonne HORAIRES
        if isinstance(start_time, str):
            time_str = start_time
        else:
            time_str = f"{start_time.strftime('%Hh%M')}-{end_time.strftime('%Hh%M')}"
        
        ws.cell(row=current_row, column=2, value=time_str)
        ws.cell(row=current_row, column=2).alignment = center_alignment
        ws.cell(row=current_row, column=2).border = border
        
        # Colonne ACTIVITÉS
        ws.cell(row=current_row, column=3, value=description)
        
        # Coloration selon le type d'activité
        if "🚗" in description or "→" in description:
            # Transport
            pass  # Pas de coloration spéciale
        elif "🍽️" in description or "Déjeuner" in description:
            # Déjeuner
            ws.cell(row=current_row, column=3).fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
        elif "🕌" in description or "Prière" in description:
            # Prière
            ws.cell(row=current_row, column=3).fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
        else:
            # Activité normale
            ws.cell(row=current_row, column=3).fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
        
        ws.cell(row=current_row, column=3).border = border
        
        # Colonne TRANSPORT
        import re
        # Extraire distance et durée de la description
        if "🚗" in description and "(" in description:
            # Format: "🚗 Dakar → Saint-Louis (240.9 km, 0min)"
            match = re.search(r"\(([\d\.]+)\s*km,\s*([^)]+)\)", description)
            if match:
                km = match.group(1)
                duration = match.group(2).strip()
                transport_text = f"~{km} km / ~{duration}"
                ws.cell(row=current_row, column=4, value=transport_text)
                ws.cell(row=current_row, column=4).font = Font(color="D32F2F")
            else:
                ws.cell(row=current_row, column=4, value="-")
        else:
            ws.cell(row=current_row, column=4, value="-")
        
        ws.cell(row=current_row, column=4).alignment = center_alignment
        ws.cell(row=current_row, column=4).border = border
        
        # Colonne NUIT
        ws.cell(row=current_row, column=5, value="")
        ws.cell(row=current_row, column=5).fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
        ws.cell(row=current_row, column=5).border = border
        
        current_row += 1
    
    # Note en bas
    current_row += 1
    ws[f'A{current_row}'] = "ℹ️ Distances/temps indicatifs. Déjeuner (13h00-14h30, ≤1h) et prière (14h00-15h00, ≤20 min) sont flexibles et intégrés sans bloquer les activités."
    ws[f'A{current_row}'].font = Font(size=9, italic=True)
    ws.merge_cells(f'A{current_row}:E{current_row}')
    
    # Ajuster la largeur des colonnes
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 50
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 12
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

# Test de connexion
if st.sidebar.button("🔍 Tester connexion Maps"):
    # Animation d'attente améliorée
    progress_bar = st.sidebar.progress(0)
    status_text = st.sidebar.empty()
    
    # Étape 1: Initialisation
    progress_bar.progress(25)
    status_text.text("🔄 Initialisation du test...")
    
    # Étape 2: Test de connexion
    progress_bar.progress(75)
    status_text.text("🌐 Test de connexion Maps...")
    
    success, message = test_graphhopper_connection(graphhopper_api_key)
    
    # Étape 3: Finalisation
    progress_bar.progress(100)
    status_text.text("✅ Test terminé")
    
    # Nettoyage de l'animation
    progress_bar.empty()
    status_text.empty()
    
    if success:
        st.sidebar.success(f"✅ {message}")
    else:
        st.sidebar.error(f"❌ {message}")

# Mention développeur
st.sidebar.markdown("---")
st.sidebar.caption("💻 Developed by @Moctar TAll (+221 77 639 96 12)")
st.sidebar.caption("All rights reserved")

# --------------------------
# FORMULAIRE
# --------------------------
st.header("📍 Paramètres de la mission")

# Champ pour le titre de mission personnalisé
st.subheader("📝 Titre de la mission")
mission_title = st.text_input(
    "Titre personnalisé de votre mission",
    value=st.session_state.get("mission_title", "Mission Terrain"),
    help="Ce titre apparaîtra dans la présentation professionnelle et tous les documents générés",
    placeholder="Ex: Mission d'inspection technique, Visite commerciale, Audit de site..."
)

st.divider()

tab1, tab2, tab3 = st.tabs(["Sites à visiter", "Dates et Horaires de la mission ", "Paramètrage des pauses"])

with tab1:
    st.markdown("**Configurez votre mission**")
    
    st.subheader("🏠 Point de départ et d'arrivée")
    col1, col2 = st.columns(2)
    
    with col1:
        use_base_location = st.checkbox("Utiliser un point de départ/arrivée fixe", value=st.session_state.get("use_base_location", True))
    
    with col2:
        if use_base_location:
            base_location = st.text_input("Ville de départ/arrivée", value=st.session_state.get("base_location", "Dakar"))
        else:
            base_location = ""
    
    st.divider()
    
    # En-tête optimisé avec informations contextuelles
    col_header1, col_header2 = st.columns([3, 1])
    with col_header1:
        st.subheader("📍 Sites à visiter")
    with col_header2:
        # Affichage compact du statut et compteur sur la même ligne
        if 'data_saved' in st.session_state and st.session_state.data_saved:
            col_status, col_count = st.columns([1, 1])
            with col_status:
                st.success("✅ Sauvegardé")
            with col_count:
                st.metric("Sites", len(st.session_state.sites_df) if 'sites_df' in st.session_state else 0)
    
    # Message d'aide contextuel
    if 'sites_df' not in st.session_state or len(st.session_state.sites_df) == 0:
        st.info("💡 **Commencez par ajouter vos sites à visiter** - Utilisez le tableau ci-dessous pour saisir les villes, types d'activités et durées prévues.")
    
    if 'sites_df' not in st.session_state:
        if use_base_location:
            st.session_state.sites_df = pd.DataFrame([
                {"Ville": "Thiès", "Type": "Client", "Activité": "Réunion commerciale", "Durée (h)": 2.0, "Peut continuer": False, "Possibilité de nuitée": True},
                {"Ville": "Saint-Louis", "Type": "Sites technique", "Activité": "Inspection", "Durée (h)": 3.0, "Peut continuer": False, "Possibilité de nuitée": True},
            ])
        else:
            st.session_state.sites_df = pd.DataFrame([
                {"Ville": "Dakar", "Type": "Agence", "Activité": "Brief", "Durée (h)": 0.5, "Peut continuer": False, "Possibilité de nuitée": True},
                {"Ville": "Thiès", "Type": "Sites technique", "Activité": "Visite", "Durée (h)": 2.0, "Peut continuer": False, "Possibilité de nuitée": True},
            ])
    
    # Gestion des types de sites personnalisés
    if 'custom_site_types' not in st.session_state:
        st.session_state.custom_site_types = []
    
    # Types de base + types personnalisés
    base_types = ["Agence", "Client", "Sites technique", "Site BTS", "Partenaire", "Autre"]
    all_types = base_types + st.session_state.custom_site_types
    
    # Tableau optimisé avec liste déroulante et saisie libre
    st.markdown("**📋 Tableau des sites à visiter :**")
    
    # Ajouter une option "Autre (saisir)" pour permettre la saisie libre
    dropdown_options = all_types + ["✏️ Autre (saisir)"]
    
    # Préparer un DataFrame éditable en y ajoutant une colonne de suppression
    editable_df = st.session_state.sites_df.copy()
    if 'Supprimer' not in editable_df.columns:
        try:
            editable_df['Supprimer'] = False
        except Exception:
            # En cas d'une structure inattendue, garantir l'existence de la colonne
            editable_df = pd.DataFrame(editable_df)
            editable_df['Supprimer'] = False

    sites_df = st.data_editor(
        editable_df, 
        num_rows="dynamic", 
        use_container_width=True,
        key="sites_data_editor",
        height=300,  # Hauteur fixe pour une meilleure lisibilité
        column_config={
            "Supprimer": st.column_config.CheckboxColumn(
                "🗑️",
                default=False,
                help="Cocher pour supprimer cette ligne",
                width=35
            ),
            "Ville": st.column_config.TextColumn(
                "🏙️ Ville", 
                required=True,
                help="Nom de la ville ou localité à visiter",
                width="medium"
            ),
            "Type": st.column_config.SelectboxColumn(
                "🏢 Type",
                options=dropdown_options,
                default="Sites technique",
                help="Sélectionnez un type ou choisissez 'Autre (saisir)' pour créer un nouveau type",
                width="medium"
            ),
            "Activité": st.column_config.TextColumn(
                "⚡ Activité", 
                default="Visite",
                help="Nature de l'activité prévue",
                width="medium"
            ),
            "Durée (h)": st.column_config.NumberColumn(
                "⏱️ Durée (h)",
                min_value=0.25,
                max_value=24,
                step=0.25,
                format="%.2f",
                default=1.0,
                help="Durée estimée en heures",
                width="small"
            ),
            "Peut continuer": st.column_config.CheckboxColumn(
                "🔄 Peut continuer",
                default=False,
                help="Cochez si cette activité peut être reportée au jour suivant si elle dépasse les heures d'activité",
                width="small"
            ),
            "Possibilité de nuitée": st.column_config.CheckboxColumn(
                "🏨 Nuitée possible",
                default=True,
                help="Décochez si cette zone ne dispose pas d'hébergement correct et qu'il faut éviter d'y passer la nuit",
                width="small"
            )
        },
        column_order=["Supprimer", "Ville", "Type", "Activité", "Durée (h)", "Peut continuer", "Possibilité de nuitée"]
    )
    
    # Interface pour saisir un nouveau type si "Autre (saisir)" est sélectionné
    if sites_df is not None and not sites_df.empty:
        # Vérifier s'il y a des lignes avec "✏️ Autre (saisir)"
        custom_rows = sites_df[sites_df['Type'] == "✏️ Autre (saisir)"]
        if not custom_rows.empty:
            st.info("💡 **Nouveau type détecté** - Veuillez spécifier le type personnalisé ci-dessous :")
            
            for idx in custom_rows.index:
                col1, col2, col3 = st.columns([2, 3, 1])
                with col1:
                    st.write(f"**Ligne {idx + 1}** - {sites_df.loc[idx, 'Ville']}")
                with col2:
                    new_custom_type = st.text_input(
                        f"Type personnalisé pour la ligne {idx + 1}",
                        placeholder="Ex: Site industriel, Centre de données...",
                        key=f"custom_type_{idx}",
                        label_visibility="collapsed"
                    )
                with col3:
                    if st.button("✅", key=f"apply_{idx}", help="Appliquer ce type"):
                        if new_custom_type and new_custom_type.strip():
                            # Ajouter le nouveau type à la liste des types personnalisés
                            if new_custom_type.strip() not in st.session_state.custom_site_types:
                                st.session_state.custom_site_types.append(new_custom_type.strip())
                            
                            # Mettre à jour la ligne dans le DataFrame
                            sites_df.loc[idx, 'Type'] = new_custom_type.strip()
                            st.session_state.sites_df = sites_df
                            # Pas de rerun automatique pour éviter de ralentir la saisie
    
    # Boutons d'action
    # Vérifier s'il y a des lignes cochées pour suppression
    has_checked_rows = 'Supprimer' in sites_df.columns and sites_df['Supprimer'].any()
    
    if has_checked_rows:
        col1, col2, col3 = st.columns([2, 1, 2])
        with col1:
            if st.button("🗑️ Supprimer lignes cochées", use_container_width=True):
                # Sécurité: gérer le cas où la colonne 'Supprimer' serait absente ou non booléenne
                if 'Supprimer' in sites_df.columns:
                    suppr_series = sites_df['Supprimer'].fillna(False)
                    try:
                        suppr_mask = ~suppr_series.astype(bool)
                    except Exception:
                        # Si conversion échoue, ne supprimer aucune ligne
                        suppr_mask = [True] * len(sites_df)
                    remaining_df = sites_df[suppr_mask].copy()
                    if 'Supprimer' in remaining_df.columns:
                        remaining_df = remaining_df.drop(columns=['Supprimer'])
                else:
                    remaining_df = sites_df.copy()
                st.session_state.sites_df = remaining_df.reset_index(drop=True)
                st.success("Lignes sélectionnées supprimées")
                st.rerun()
    else:
        col1, col2, col3 = st.columns([2, 1, 2])

    with col2:
        if st.button("💾 Enregistrer", use_container_width=True, type="primary"):
            # Nettoyer la colonne de suppression avant sauvegarde
            df_to_save = sites_df.drop(columns=['Supprimer']) if 'Supprimer' in sites_df.columns else sites_df
            st.session_state.sites_df = df_to_save
            st.session_state.data_saved = True  # Marquer comme sauvegardé
            st.rerun()  # Rafraîchir pour afficher le statut en haut
    
    # Pas d'enregistrement automatique - seulement lors du clic sur Enregistrer ou Planifier
    # st.session_state.sites_df = sites_df
    
    # Option d'ordre des sites
    order_mode = "🤖 Automatique (optimisé)"  # Valeur par défaut pour 0 ou 1 site
    if len(sites_df) > 1:  # Afficher seulement s'il y a plus d'un site
        st.subheader("🔄 Ordre des visites")
        order_mode = st.radio(
            "Mode d'ordonnancement",
            ["🤖 Automatique (optimisé)", "✋ Manuel (personnalisé)"],
            horizontal=True,
            help="Automatique: optimise l'ordre pour minimiser les distances. Manuel: vous choisissez l'ordre."
        )
        
        if order_mode == "✋ Manuel (personnalisé)":
            with st.container():
                st.info("💡 **Astuce :** Utilisez les flèches pour réorganiser vos sites dans l'ordre de visite souhaité")
                
                # Créer une liste ordonnée des sites pour réorganisation
                if 'manual_order' not in st.session_state or len(st.session_state.manual_order) != len(sites_df):
                    st.session_state.manual_order = list(range(len(sites_df)))
                
                # Interface de réorganisation manuelle améliorée
                st.markdown("**📋 Ordre de visite des sites :**")
                
                # Conteneur avec style pour la liste
                with st.container():
                    for i, idx in enumerate(st.session_state.manual_order):
                        if idx < len(sites_df):
                            site = sites_df.iloc[idx]
                            
                            # Créer une ligne avec un style visuel amélioré
                            col1, col2, col3, col4, col5 = st.columns([0.8, 2.5, 2, 1, 1])
                            
                            with col1:
                                st.markdown(f"**`{i+1}`**")
                            with col2:
                                st.markdown(f"📍 **{site['Ville']}**")
                            with col3:
                                st.markdown(f"🏢 {site['Type']}")
                            with col4:
                                st.markdown(f"⏱️ {site['Durée (h)']}h")
                            with col5:
                                # Boutons de réorganisation dans une ligne
                                subcol1, subcol2 = st.columns(2)
                                with subcol1:
                                    if i > 0:
                                        if st.button("⬆️", key=f"manual_up_{i}", help="Monter", use_container_width=True):
                                            st.session_state.manual_order[i], st.session_state.manual_order[i-1] = \
                                                st.session_state.manual_order[i-1], st.session_state.manual_order[i]
                                            st.rerun()
                                with subcol2:
                                    if i < len(st.session_state.manual_order) - 1:
                                        if st.button("⬇️", key=f"manual_down_{i}", help="Descendre", use_container_width=True):
                                            st.session_state.manual_order[i], st.session_state.manual_order[i+1] = \
                                                st.session_state.manual_order[i+1], st.session_state.manual_order[i]
                                            st.rerun()
                            
                            # Séparateur visuel entre les éléments
                            if i < len(st.session_state.manual_order) - 1:
                                st.markdown("---")
                
                # Boutons d'action
                col1, col2, col3 = st.columns([1, 1, 2])
                with col1:
                    if st.button("🔄 Réinitialiser", help="Remettre l'ordre original", use_container_width=True):
                        st.session_state.manual_order = list(range(len(sites_df)))
                        st.rerun()
                with col2:
                    if st.button("🔀 Mélanger", help="Ordre aléatoire", use_container_width=True):
                        import random
                        random.shuffle(st.session_state.manual_order)
                        st.rerun()
        else:
            st.success("🤖 **Mode automatique activé** - L'ordre des sites sera optimisé automatiquement pour minimiser les temps de trajet")
    else:
        st.info("ℹ️ Ajoutez au moins 1 site pour continuer. L'ordre n'est requis que s'il y a plusieurs sites.")

with tab2:
    col1, col2 = st.columns([1, 2])  # Réduire la largeur de la colonne Dates
    with col1:
        st.subheader("📅 Dates")
        start_date = st.date_input("Dat                                                                                                                                                                                                                                             e de départ de la mission", value=st.session_state.get("start_date", datetime.today().date()))
        max_days = st.number_input("Nombre de jours max (Optionel)", min_value=0, value=st.session_state.get("max_days", 0), step=1, help="Laisser zéro pour le calcul automatique. Agit comme une limite supérieure.")
        desired_days = st.number_input("Nombre de jours souhaités (Optionnel)", min_value=0, value=st.session_state.get("desired_days", 0), step=1, help="Laissez à zéro pour ignorer. Le planning sera ajusté pour correspondre à ce nombre si possible.")
        
        st.divider()
        
        # Ajouter des informations utiles dans la section Dates
        st.markdown("**📊 Informations**")
        if start_date:
            # Jour de la semaine avec date complète
            weekdays = ["Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi", "Dimanche"]
            months = ["janvier", "février", "mars", "avril", "mai", "juin", 
                     "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
            start_weekday = weekdays[start_date.weekday()]
            start_month = months[start_date.month - 1]
            formatted_date = f"{start_weekday.lower()} {start_date.day} {start_month} {start_date.year}"
            st.info(f"🗓️ Jour de début : {formatted_date}")
    
    with col2:
        st.subheader("⏰ Horaires")
        
        # Horaires d'activité
        st.markdown("**Horaires d'activité** (visites, réunions)")
        col_act1, col_act2 = st.columns(2)
        with col_act1:
            start_activity_time = st.time_input("Début activités", value=st.session_state.get("start_activity_time", time(8, 0)))
        with col_act2:
            end_activity_time = st.time_input("Fin activités", value=st.session_state.get("end_activity_time", time(16, 30)))
        
        # Horaires de voyage
        st.markdown("**Horaires de voyage** (trajets)")
        col_travel1, col_travel2 = st.columns(2)
        with col_travel1:
            start_travel_time = st.time_input("Début voyages", value=st.session_state.get("start_travel_time", time(7, 30)))
        with col_travel2:
            end_travel_time = st.time_input("Fin voyages", value=st.session_state.get("end_travel_time", time(19, 0)))
        
        # Options week-end
        st.markdown("**Options week-end**")
        allow_weekend_travel = st.checkbox(
            "Autoriser les voyages le week-end",
            value=st.session_state.get("allow_weekend_travel", True)
        )
        allow_weekend_activities = st.checkbox(
            "Autoriser les activités le week-end",
            value=st.session_state.get("allow_weekend_activities", True)
        )
        st.session_state.allow_weekend_travel = allow_weekend_travel
        st.session_state.allow_weekend_activities = allow_weekend_activities
        
        st.divider()
        
        # Gestion des activités longues
        st.markdown("**Gestion des activités longues**")
        col_tol1, col_tol2 = st.columns(2)
        with col_tol1:
            tolerance_hours = st.number_input(
                "Seuil de tolérance (heures)", 
                min_value=0.0, 
                max_value=3.0, 
                value=st.session_state.get("tolerance_hours", 1.0), 
                step=0.25,
                help="Activités se terminant dans ce délai après la fin des heures d'activité peuvent continuer le même jour"
            )
        with col_tol2:
            default_can_continue = st.checkbox(
                "Une partie d’une activité non achevée à l’heure de la descente pourra être poursuivie le lendemain", 
                value=False,
                help="Non poursuite cochée par défaut"
            )
        
        # Maintenir la compatibilité avec l'ancien code
        start_day_time = start_activity_time
        end_day_time = end_activity_time

with tab3:
    st.subheader("🍽️ Pauses flexibles")
    st.info("💡 Les pauses s'insèrent automatiquement pendant les trajets ou visites qui chevauchent les fenêtres définies")
    
    col1, col2 = st.columns(2)
    with col1:
        use_lunch = st.checkbox("Pause déjeuner", value=st.session_state.get("use_lunch", True))
        if use_lunch:
            st.markdown("**Fenêtre de déjeuner**")
            lunch_start_time = st.time_input("Début fenêtre", value=st.session_state.get("lunch_start_time", time(12, 30)), key="lunch_start")
            lunch_end_time = st.time_input("Fin fenêtre", value=st.session_state.get("lunch_end_time", time(15, 0)), key="lunch_end")
            lunch_duration_min = st.number_input(
                "Durée pause (min)",
                min_value=5,
                max_value=180,
                step=5,
                value=st.session_state.get("lunch_duration_min", 60),
                key="lunch_duration"
            )
    
    with col2:
        use_prayer = st.checkbox("Pause prière", value=st.session_state.get("use_prayer", False))
        if use_prayer:
            st.markdown("**Fenêtre de prière**")
            prayer_start_time = st.time_input("Début fenêtre", value=st.session_state.get("prayer_start_time", time(13, 0)), key="prayer_start")
            prayer_duration_min = st.number_input("Durée pause (min)", min_value=5, max_value=60, value=st.session_state.get("prayer_duration_min", 20) or 20, key="prayer_duration")

    st.divider()
    st.subheader("📦 Import/Export JSON")
    with st.expander("Sauvegarde et reprise (JSON)"):
        col_export, col_import = st.columns(2)
        with col_export:
            mission_config = {
                "mission_title": mission_title,
                "use_base_location": use_base_location,
                "base_location": base_location,
                "sites": (st.session_state.sites_df.to_dict(orient="records") if "sites_df" in st.session_state else []),
                "start_date": start_date.strftime("%Y-%m-%d") if isinstance(start_date, datetime) else str(start_date),
                "max_days": int(max_days),
                "start_activity_time": (start_activity_time.strftime("%H:%M") if start_activity_time else None),
                "end_activity_time": (end_activity_time.strftime("%H:%M") if end_activity_time else None),
                "start_travel_time": (start_travel_time.strftime("%H:%M") if start_travel_time else None),
                "end_travel_time": (end_travel_time.strftime("%H:%M") if end_travel_time else None),
                "tolerance_hours": float(tolerance_hours),
                "use_lunch": bool(use_lunch),
                "lunch_start_time": (lunch_start_time.strftime("%H:%M") if use_lunch and lunch_start_time else None),
                "lunch_end_time": (lunch_end_time.strftime("%H:%M") if use_lunch and lunch_end_time else None),
                "lunch_duration_min": (int(lunch_duration_min) if use_lunch else None),
                "use_prayer": bool(use_prayer),
                "prayer_start_time": (prayer_start_time.strftime("%H:%M") if use_prayer and prayer_start_time else None),
                "prayer_duration_min": (int(prayer_duration_min) if use_prayer and prayer_duration_min is not None else None),
                "distance_method": distance_method,
            }
            json_str = json.dumps(mission_config, ensure_ascii=False, indent=2)
            st.download_button(
                label="💾 Exporter JSON",
                data=json_str,
                file_name=f"mission_config_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True,
            )
        with col_import:
            uploaded_file = st.file_uploader("Importer configuration JSON", type=["json"], help="Chargez un fichier exporté précédemment pour reprendre la mission")
            if uploaded_file is not None:
                try:
                    imported = json.loads(uploaded_file.getvalue().decode("utf-8"))
                    def parse_time(val, fallback):
                        try:
                            if isinstance(val, str) and ":" in val:
                                hh, mm = val.split(":")
                                return time(int(hh), int(mm))
                        except Exception:
                            pass
                        return fallback

                    st.session_state.mission_title = imported.get("mission_title", mission_title)
                    st.session_state.use_base_location = imported.get("use_base_location", use_base_location)
                    st.session_state.base_location = imported.get("base_location", base_location)

                    # Sites
                    sites_records = imported.get("sites", [])
                    if isinstance(sites_records, list):
                        st.session_state.sites_df = pd.DataFrame(sites_records)

                    # Dates et horaires
                    sd = imported.get("start_date")
                    try:
                        if sd:
                            st.session_state.start_date = datetime.strptime(sd, "%Y-%m-%d").date()
                    except Exception:
                        st.session_state.start_date = st.session_state.get("start_date", datetime.today().date())

                    st.session_state.max_days = imported.get("max_days", max_days)
                    st.session_state.start_activity_time = parse_time(imported.get("start_activity_time"), start_activity_time)
                    st.session_state.end_activity_time = parse_time(imported.get("end_activity_time"), end_activity_time)
                    st.session_state.start_travel_time = parse_time(imported.get("start_travel_time"), start_travel_time)
                    st.session_state.end_travel_time = parse_time(imported.get("end_travel_time"), end_travel_time)
                    st.session_state.tolerance_hours = imported.get("tolerance_hours", tolerance_hours)

                    # Pauses
                    st.session_state.use_lunch = imported.get("use_lunch", use_lunch)
                    st.session_state.lunch_start_time = parse_time(imported.get("lunch_start_time"), lunch_start_time)
                    st.session_state.lunch_end_time = parse_time(imported.get("lunch_end_time"), lunch_end_time)
                    st.session_state.lunch_duration_min = imported.get("lunch_duration_min", st.session_state.get("lunch_duration_min", 60))
                    st.session_state.use_prayer = imported.get("use_prayer", use_prayer)
                    st.session_state.prayer_start_time = parse_time(imported.get("prayer_start_time"), prayer_start_time)
                    val_prayer_dur = imported.get("prayer_duration_min", None)
                    try:
                        st.session_state.prayer_duration_min = int(val_prayer_dur) if val_prayer_dur is not None else st.session_state.get("prayer_duration_min", 20)
                    except Exception:
                        st.session_state.prayer_duration_min = st.session_state.get("prayer_duration_min", 20)

                    st.success("✅ Configuration importée. Les paramètres et sites ont été mis à jour.")
                except Exception as e:
                    st.error(f"❌ Import JSON invalide: {e}")

# --------------------------
# PLANIFICATION
# --------------------------

col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    plan_button = st.button("🚀 Planifier la mission", type="primary", use_container_width=True)

if plan_button:
    # Sauvegarde automatique des données avant planification
    st.session_state.sites_df = sites_df
    # Persistance des paramètres saisis
    st.session_state.mission_title = mission_title
    st.session_state.use_base_location = use_base_location
    st.session_state.base_location = base_location
    st.session_state.start_date = start_date
    st.session_state.max_days = max_days
    st.session_state.start_activity_time = start_activity_time
    st.session_state.end_activity_time = end_activity_time
    st.session_state.start_travel_time = start_travel_time
    st.session_state.end_travel_time = end_travel_time
    st.session_state.tolerance_hours = tolerance_hours
    st.session_state.use_lunch = use_lunch
    st.session_state.lunch_start_time = lunch_start_time if use_lunch else None
    st.session_state.lunch_end_time = lunch_end_time if use_lunch else None
    st.session_state.lunch_duration_min = lunch_duration_min if use_lunch else None
    st.session_state.use_prayer = use_prayer
    st.session_state.prayer_start_time = prayer_start_time if use_prayer else None
    st.session_state.prayer_duration_min = prayer_duration_min if use_prayer else None

    # Validations basiques avant planification
    if sites_df is None or sites_df.empty:
        st.error("❌ Ajoutez au moins un site avant de planifier.")
        st.stop()
    issues = []
    for i, row in sites_df.iterrows():
        city = str(row.get("Ville", "")).strip()
        dur = row.get("Durée (h)", None)
        if not city:
            issues.append(f"Ligne {i + 1}: Ville manquante ou vide")
        try:
            val = float(dur) if dur is not None else 0
        except Exception:
            val = 0
        if val <= 0:
            issues.append(f"Ligne {i + 1}: Durée (h) doit être > 0")
    if use_base_location and not str(base_location).strip():
        issues.append("Point de départ/arrivée activé mais ville non renseignée")
    if issues:
        st.error("⚠️ Veuillez corriger ces points avant la planification:")
        for msg in issues[:10]:
            st.write(f"- {msg}")
        st.stop()
    
    # Animation CSS moderne pour l'attente
    st.markdown("""
    <style>
    .planning-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 30px;
        margin: 20px 0;
        text-align: center;
        color: white;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    
    .spinner-icon {
        font-size: 3em;
        animation: spin 2s linear infinite;
        margin-bottom: 20px;
        display: inline-block;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .pulse-text {
        animation: pulse 1.5s ease-in-out infinite alternate;
        font-size: 1.2em;
        font-weight: bold;
        margin: 10px 0;
    }
    
    @keyframes pulse {
        0% { opacity: 0.6; }
        100% { opacity: 1; }
    }
    
    .progress-enhanced {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        overflow: hidden;
        margin: 20px 0;
        height: 8px;
        position: relative;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }
    
    .progress-enhanced::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.4), transparent);
        animation: shimmer 2s infinite;
    }
    
    @keyframes shimmer {
        0% { left: -100%; }
        100% { left: 100%; }
    }
    
    .step-indicator {
        display: flex;
        justify-content: space-between;
        margin: 20px 0;
        font-size: 0.9em;
        position: relative;
    }
    
    .step-indicator::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 0;
        right: 0;
        height: 2px;
        background: linear-gradient(90deg, rgba(255,255,255,0.3) 0%, rgba(255,255,255,0.6) 50%, rgba(255,255,255,0.3) 100%);
        z-index: 1;
        transform: translateY(-50%);
    }
    
    .step {
        padding: 8px 15px;
        border-radius: 20px;
        background: rgba(255,255,255,0.15);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        position: relative;
        z-index: 2;
        border: 2px solid transparent;
        backdrop-filter: blur(10px);
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .step.active {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        transform: scale(1.15);
        color: white;
        border: 2px solid rgba(255,255,255,0.5);
        box-shadow: 0 8px 25px rgba(79, 172, 254, 0.4);
        animation: glow 2s ease-in-out infinite alternate;
    }
    
    .step.completed {
        background: linear-gradient(135deg, #56ab2f 0%, #a8e6cf 100%);
        color: white;
        border: 2px solid rgba(255,255,255,0.3);
        box-shadow: 0 4px 15px rgba(86, 171, 47, 0.3);
    }
    
    @keyframes glow {
        0% { box-shadow: 0 8px 25px rgba(79, 172, 254, 0.4); }
        100% { box-shadow: 0 12px 35px rgba(79, 172, 254, 0.6); }
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Container d'animation
    animation_container = st.empty()
    
    with animation_container.container():
        st.markdown("""
        <div class="planning-container">
            <div class="spinner-icon">🗺️</div>
            <div class="pulse-text">Planification intelligente en cours...</div>
            <div class="step-indicator">
                <span class="step active" id="step-1">📍 Géocodage</span>
                <span class="step" id="step-2">🗺️ Distances</span>
                <span class="step" id="step-3">🔄 Optimisation</span>
                <span class="step" id="step-4">🛣️ Itinéraire</span>
                <span class="step" id="step-5">📅 Planning</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    rows = sites_df.replace({pd.NA: None}).to_dict(orient="records")
    sites = [r for r in rows if r.get("Ville") and str(r["Ville"]).strip()]
    
    if use_base_location and base_location and base_location.strip():
        base_site = {"Ville": base_location.strip(), "Type": "Base", "Activité": "Départ", "Durée (h)": 0}
        return_site = {"Ville": base_location.strip(), "Type": "Base", "Activité": "Retour", "Durée (h)": 0}
        all_sites = [base_site] + sites + [return_site]
        
        if len(sites) < 1:
            st.error("❌ Ajoutez au moins 1 site à visiter")
            st.stop()
    else:
        all_sites = sites
        # Autoriser la planification avec un seul site (sans base)
        if len(all_sites) < 1:
            st.error("❌ Ajoutez au moins 1 site à visiter")
            st.stop()
        first_site = all_sites[0].copy()
        first_site["Activité"] = "Retour"
        all_sites = all_sites + [first_site]
    
    progress = st.progress(0)
    status = st.empty()
    
    # Fonction pour mettre à jour l'animation avec JavaScript
    def update_animation_step(step_number, icon, message, completed_steps=None):
        if completed_steps is None:
            completed_steps = []
        
        def get_step_class(step_num):
            if step_num in completed_steps:
                return 'completed'
            elif step_num == step_number:
                return 'active'
            else:
                return ''
        
        animation_container.markdown(f"""
        <div class="planning-container">
            <div class="spinner-icon">{icon}</div>
            <div class="pulse-text">{message}</div>
            <div class="step-indicator">
                <span class="step {get_step_class(1)}" id="step-1">📍 Géocodage</span>
                <span class="step {get_step_class(2)}" id="step-2">🗺️ Distances</span>
                <span class="step {get_step_class(3)}" id="step-3">🔄 Optimisation</span>
                <span class="step {get_step_class(4)}" id="step-4">🛣️ Itinéraire</span>
                <span class="step {get_step_class(5)}" id="step-5">📅 Planning</span>
            </div>
            <div class="progress-enhanced"></div>
        </div>
        """, unsafe_allow_html=True)
    
    # Messages dynamiques pour chaque étape
    geocoding_messages = [
        "🔍 Recherche des coordonnées GPS...",
        "📍 Géolocalisation des sites en cours...",
        "🌍 Validation des adresses...",
        "✅ Géocodage terminé avec succès!"
    ]
    
    # Étape 1: Géocodage
    update_animation_step(1, "📍", geocoding_messages[0], [])
    status.text("📍 Géocodage...")
    coords = []
    failed = []
    
    for i, s in enumerate(all_sites):
        progress.progress((i+1) / (len(all_sites) * 4))
        # Message dynamique pendant le géocodage
        if i < len(geocoding_messages) - 1:
            update_animation_step(1, "📍", geocoding_messages[min(i, len(geocoding_messages)-2)], [])
        city_val = str(s.get("Ville", "")).strip()
        if s.get("Type") == "Base" and city_val.lower() == "dakar":
            coord = (-17.470602, 14.711404)
        else:
            coord = geocode_city_senegal(city_val, use_cache)
        if not coord:
            failed.append(s["Ville"])
        else:
            coords.append(coord)
    
    update_animation_step(1, "✅", geocoding_messages[-1], [1])
    
    if failed:
        st.error(f"❌ Villes introuvables: {', '.join(failed)}")
        st.stop()
    
    # Étape 2: Calcul des distances
    distance_messages = [
        "🗺️ Connexion aux services de cartographie...",
        "📏 Calcul des distances entre les sites...",
        "⏱️ Estimation des temps de trajet...",
        "✅ Matrice de distances calculée!"
    ]
    
    update_animation_step(2, "🗺️", distance_messages[0], [1])
    status.text("🗺️ Calcul des distances...")
    progress.progress(0.4)
    
    durations_sec = None
    distances_m = None
    calculation_method = ""
    city_list = [s["Ville"] for s in all_sites]
    
    if distance_method.startswith("Maps uniquement"):
        update_animation_step(2, "🗺️", distance_messages[1], [1])
        durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
        calculation_method = "Maps"
        if durations_sec is None:
            st.error(f"❌ {error_msg}")
            st.stop()
        else:
            if debug_mode:
                st.info(f"🔍 Debug Maps: {len(durations_sec)} x {len(durations_sec[0]) if durations_sec else 0} matrice de durées reçue")
                if durations_sec and len(durations_sec) > 0:
                    sample = durations_sec[0][1] if len(durations_sec[0]) > 1 else 0
                    st.info(f"🔍 Debug Maps: Exemple durée [0][1] = {sample} secondes ({sample/3600:.2f}h)")

    elif distance_method == "OSRM uniquement (rapide)":
        update_animation_step(2, "🗺️", distance_messages[1], [1])
        durations_sec, distances_m, error_msg = improved_osrm_duration_matrix(osrm_base_url, coords)
        calculation_method = "OSRM"
        if durations_sec is None:
            st.error(f"❌ {error_msg}")
            st.stop()
        else:
            if debug_mode:
                st.info(f"🔍 Debug OSRM: {len(durations_sec)} x {len(durations_sec[0]) if durations_sec else 0} matrice de durées reçue")
                if durations_sec and len(durations_sec) > 0:
                    sample = durations_sec[0][1] if len(durations_sec[0]) > 1 else 0
                    st.info(f"🔍 Debug OSRM: Exemple durée [0][1] = {sample} secondes ({sample/3600:.2f}h)")

    elif distance_method == "Automatique uniquement":
        result, error_msg = improved_deepseek_estimate_matrix(city_list, deepseek_api_key, debug_mode)
        if result:
            durations_sec, distances_m = result
            calculation_method = "Automatique"
            st.info(f"📊 Méthode: {calculation_method}")
        else:
            st.error(f"❌ {error_msg}")
            st.stop()

    elif distance_method == "Géométrique uniquement":
        durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
        calculation_method = f"Géométrique ({default_speed_kmh} km/h)"
        st.warning(f"📊 Méthode: {calculation_method}")

    else:
        # Mode Auto: dynamique
        # - Site unique → OSRM → Automatique → Maps → Géométrique
        # - Plusieurs sites → OSRM → Automatique → Maps → Géométrique
        single_site = len(sites) == 1

        if single_site:
            update_animation_step(2, "🗺️", distance_messages[1], [1])
            # 1) OSRM en premier
            durations_sec, distances_m, error_msg = improved_osrm_duration_matrix(osrm_base_url, coords)
            if durations_sec is not None:
                calculation_method = "OSRM"
            else:
                # 2) Automatique (DeepSeek)
                if deepseek_api_key:
                    result, _ = improved_deepseek_estimate_matrix(city_list, deepseek_api_key, debug_mode)
                    if result:
                        durations_sec, distances_m = result
                        calculation_method = "Automatique"
                    else:
                        # 3) Maps (GraphHopper), si option activée
                        if use_deepseek_fallback and graphhopper_api_key:
                            durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
                            if durations_sec is not None:
                                calculation_method = "Maps"
                            else:
                                durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                                calculation_method = f"Géométrique ({default_speed_kmh} km/h)"
                        else:
                            durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                            calculation_method = f"Géométrique ({default_speed_kmh} km/h)"
                else:
                    # Pas de clé DeepSeek, tenter Maps si autorisé puis géométrique
                    if use_deepseek_fallback and graphhopper_api_key:
                        durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
                        if durations_sec is not None:
                            calculation_method = "Maps"
                        else:
                            durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                            calculation_method = f"Géométrique ({default_speed_kmh} km/h)"
                    else:
                        durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                        calculation_method = f"Géométrique ({default_speed_kmh} km/h)"
        else:
            # OSRM → Automatique → Maps (si activé) → Géométrique
            durations_sec, distances_m, error_msg = improved_osrm_duration_matrix(osrm_base_url, coords)
            if durations_sec is not None:
                calculation_method = "OSRM"
            else:
                result, error_msg = improved_deepseek_estimate_matrix(city_list, deepseek_api_key, debug_mode)
                if result:
                    durations_sec, distances_m = result
                    calculation_method = "Automatique"
                else:
                    if use_deepseek_fallback and graphhopper_api_key:
                        durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
                        if durations_sec is not None:
                            calculation_method = "Maps"
                        else:
                            durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                            calculation_method = f"Géométrique ({default_speed_kmh} km/h)"
                    else:
                        durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                        calculation_method = f"Géométrique ({default_speed_kmh} km/h)"

        method_color = "success" if ("Maps" in calculation_method or "OSRM" in calculation_method) else "info" if "Automatique" in calculation_method else "warning"
        getattr(st, method_color)(f"📊 Méthode: {calculation_method}")
    
    # Étape 3: Optimisation (commune à tous les modes)
    update_animation_step(3, "🔄", "Optimisation de l'itinéraire...", [1, 2])
    status.text("🔄 Optimisation de l'ordre des sites...")
    progress.progress(0.6)
    
    # Déterminer l'ordre des sites selon le mode choisi
    if order_mode == "✋ Manuel (personnalisé)":
        # Utiliser l'ordre manuel défini par l'utilisateur
        if use_base_location and base_location and base_location.strip():
            # Avec base: [base] + sites_manuels + [base]
            manual_sites_order = [0]  # Base de départ
            for manual_idx in st.session_state.manual_order:
                if manual_idx < len(sites):
                    manual_sites_order.append(manual_idx + 1)  # +1 car base est à l'index 0
            manual_sites_order.append(len(all_sites) - 1)  # Base de retour
            order = manual_sites_order
        else:
            # Sans base: sites_manuels + [premier_site]
            manual_sites_order = []
            for manual_idx in st.session_state.manual_order:
                if manual_idx < len(sites):
                    manual_sites_order.append(manual_idx)
            manual_sites_order.append(len(all_sites) - 1)  # Site de retour
            order = manual_sites_order
        
        st.success("✅ Ordre manuel appliqué")
    else:
        # Utiliser l'optimisation IA Adja au lieu du TSP traditionnel
        if len(coords) >= 3:
            # Essayer d'abord l'optimisation IA Adja
            ai_order, ai_success, ai_message = optimize_route_with_ai(
                all_sites, coords, 
                base_location if use_base_location else None, 
                deepseek_api_key
            )
            
            if ai_success:
                order = ai_order
                st.success(f"✅ Ordre optimisé par IA Adja: {ai_message}")
            else:
                # Fallback vers TSP si l'IA Adja échoue
                order = solve_tsp_fixed_start_end(durations_sec)
                st.warning(f"⚠️ IA Adja échouée ({ai_message}), utilisation TSP classique")
        else:
            order = list(range(len(coords)))
            st.success("✅ Ordre séquentiel (moins de 3 sites)")
            
        if debug_mode and durations_sec:
            # Calculer coût total pour transparence
            total_cost = sum(durations_sec[order[i]][order[i+1]] for i in range(len(order)-1))
            st.info(f"🔍 Debug Optimisation: ordre={order} | coût total={total_cost/3600:.2f}h")
        
    status.text("🛣️ Calcul de l'itinéraire détaillé...")
    # Étape 4: Génération de l'itinéraire
    update_animation_step(4, "🛣️", "Génération de l'itinéraire détaillé...", [1, 2, 3])
    progress.progress(0.8)
    
    segments = []
    zero_segments_indices = []
    
    for i in range(len(order)-1):
        from_idx = order[i]
        to_idx = order[i+1]
        
        if from_idx < len(durations_sec) and to_idx < len(durations_sec[0]):
            duration = durations_sec[from_idx][to_idx]
            distance = distances_m[from_idx][to_idx] if distances_m else 0
            segment_method = "Matrix"
            
            # Si la distance/durée est nulle, recalculer via OSRM/Maps avec cache, puis fallback géométrique
            if duration == 0 or distance == 0:
                # Cache par segment
                if 'segment_route_cache' not in st.session_state:
                    st.session_state.segment_route_cache = {}
                segment_cache = st.session_state.segment_route_cache
                SEGMENT_CACHE_TTL = int(st.session_state.get('segment_cache_ttl_sec', 43200))
                
                coord_from = coords[from_idx]
                coord_to = coords[to_idx]
                key = (coord_from[0], coord_from[1], coord_to[0], coord_to[1])
                now_ts = datetime.now().timestamp()
                
                # 1) Cache
                cached = segment_cache.get(key)
                if cached and (now_ts - cached.get('ts', 0)) < SEGMENT_CACHE_TTL:
                    distance = int(cached.get('distance', 0))
                    duration = int(cached.get('duration', 0))
                    segment_method = cached.get('method', 'Matrix')
                else:
                    # 2) OSRM route
                    try:
                        coord_str = f"{coord_from[0]},{coord_from[1]};{coord_to[0]},{coord_to[1]}"
                        url = f"{osrm_base_url.rstrip('/')}/route/v1/driving/{coord_str}?overview=false"
                        resp = requests.get(url, timeout=10)
                        if resp.status_code == 200:
                            data = resp.json()
                            if data.get('routes'):
                                r0 = data['routes'][0]
                                d_m = int(r0.get('distance', 0))
                                t_s = int(r0.get('duration', 0))
                                if d_m > 0 and t_s > 0:
                                    distance = d_m
                                    duration = t_s
                                    segment_method = "OSRM"
                    except Exception:
                        pass
                    
                    # 3) GraphHopper route
                    if (duration == 0 or distance == 0) and graphhopper_api_key:
                        try:
                            gh_url = "https://graphhopper.com/api/1/route"
                            qp = f"point={coord_from[1]},{coord_from[0]}&point={coord_to[1]},{coord_to[0]}&vehicle=car&locale=fr&points_encoded=false&calc_points=false&key={graphhopper_api_key}"
                            gh_resp = requests.get(f"{gh_url}?{qp}", timeout=10)
                            if gh_resp.status_code == 200:
                                gh_data = gh_resp.json()
                                paths = gh_data.get('paths')
                                if paths:
                                    p0 = paths[0]
                                    d_m = int(p0.get('distance', 0))
                                    t_ms = int(p0.get('time', 0))
                                    t_s = int(t_ms / 1000)
                                    if d_m > 0 and t_s > 0:
                                        distance = d_m
                                        duration = t_s
                                        segment_method = "Maps"
                        except Exception:
                            pass
                    
                    # 4) Fallback géométrique
                    if duration == 0 or distance == 0:
                        geometric_km = haversine(coord_from[0], coord_from[1], coord_to[0], coord_to[1]) * 1.2
                        if distance == 0:
                            distance = int(geometric_km * 1000)
                        if duration == 0:
                            distance_for_time_calc = distance / 1000 if distance > 0 else geometric_km
                            geometric_hours = distance_for_time_calc / default_speed_kmh
                            duration = int(geometric_hours * 3600)
                        segment_method = "Geo"
                    
                    # Mettre à jour le cache
                    segment_cache[key] = {
                        "distance": int(distance),
                        "duration": int(duration),
                        "method": segment_method,
                        "ts": now_ts
                    }
                
                zero_segments_indices.append(i)
                if debug_mode:
                    st.info(f"🔍 Segment {i} recalculé via {segment_method}: {distance/1000:.1f}km, {duration/3600:.2f}h")
            
            # Debug: Afficher les valeurs des segments
            if debug_mode:
                st.info(f"🔍 Debug Segment {i}: de {from_idx} vers {to_idx} = {duration}s ({duration/3600:.2f}h), {distance/1000:.1f}km")
            
            segments.append({
                "distance": distance,
                "duration": duration,
                "method": segment_method
            })
        else:
            segments.append({"distance": 0, "duration": 0})
    
    if not segments:
        st.error("❌ AUCUN segment créé!")
        st.stop()
    
    # Résumé des méthodes de recalcul des segments
    if zero_segments_indices:
        method_summary = {"OSRM": 0, "Maps": 0, "Geo": 0}
        for idx in zero_segments_indices:
            m = segments[idx].get("method", "")
            if m in method_summary:
                method_summary[m] += 1
        summary_str = " | ".join([
            f"OSRM: {method_summary['OSRM']}",
            f"Maps: {method_summary['Maps']}",
            f"Géo: {method_summary['Geo']}"
        ])
        st.success(f"✅ {len(zero_segments_indices)} segment(s) recalculé(s) ({summary_str})")
    
    # Vérifier s'il reste des segments à zéro après le recalcul
    remaining_zero_segments = [i for i, s in enumerate(segments) if s['duration'] == 0 or s['distance'] == 0]
    if remaining_zero_segments:
        st.warning(f"⚠️ {len(remaining_zero_segments)} segment(s) avec valeurs manquantes après recalcul")
    
    status.text("📅 Génération du planning détaillé...")
    # Étape 5: Génération du planning
    update_animation_step(5, "📅", "Finalisation du planning...", [1, 2, 3, 4])
    progress.progress(0.9)

    # Calcul préalable: nombre de jours optimal (dry-run, sans étirement)
    _, _, _, dry_stats = schedule_itinerary(
        coords=coords,
        sites=all_sites,
        order=order,
        segments_summary=segments,
        start_date=start_date,
        start_activity_time=start_activity_time,
        end_activity_time=end_activity_time,
        start_travel_time=start_travel_time,
        end_travel_time=end_travel_time,
        use_lunch=use_lunch,
        lunch_start_time=lunch_start_time if use_lunch else time(12,30),
        lunch_end_time=lunch_end_time if use_lunch else time(14,0),
        use_prayer=use_prayer,
        prayer_start_time=prayer_start_time if use_prayer else time(14,0),
        prayer_duration_min=prayer_duration_min if use_prayer else 20,
        lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
        max_days=0,
        tolerance_hours=tolerance_hours,
        base_location=base_location,
        allow_weekend_travel=allow_weekend_travel,
        allow_weekend_activities=allow_weekend_activities
    )

    optimal_days = int(dry_stats.get('total_days', 1))
    user_max = int(max_days) if isinstance(max_days, (int, float)) else 0
    user_desired = int(desired_days) if isinstance(desired_days, (int, float)) else 0

    # Logique de décision pour les jours effectifs
    if user_desired > 0:
        if user_max > 0 and user_desired > user_max:
            st.warning(f"Le nombre de jours souhaités ({user_desired}) dépasse le maximum autorisé ({user_max}). Utilisation du maximum.")
            effective_max_days = user_max
        else:
            effective_max_days = user_desired

        if effective_max_days < optimal_days:
            # Cas compression: on tente de tenir en moins de jours (journées plus chargées)
            stretch_days_flag = False
            st.warning(f"⚠️ Objectif ({effective_max_days} jours) < optimal ({optimal_days}). Compression: journées potentiellement plus chargées.")
        elif effective_max_days > optimal_days:
            # Cas étalement: on répartit sur plus de jours, fin de journée plus tôt
            stretch_days_flag = True
            st.success(f"✅ Planning étalé sur {effective_max_days} jours (optimal: {optimal_days}). Journées plus légères.")
        else:
            # Égal à l'optimal
            stretch_days_flag = False
            st.info(f"🟰 Planning sur {effective_max_days} jours, égal à l'optimal.")

    elif user_max > 0:
        if user_max < optimal_days:
            effective_max_days = user_max
            stretch_days_flag = True
            st.warning(f"⚠️ Objectif ({user_max} jours) < optimal ({optimal_days}). Compression en {user_max} jours avec journées étirées.")
        else:
            effective_max_days = user_max
            stretch_days_flag = False
            st.success(f"✅ Le planning tient en {optimal_days} jours (objectif max: {user_max} jours).")
    else:
        effective_max_days = optimal_days
        stretch_days_flag = False
        st.info(f"🧮 Jours optimaux calculés automatiquement: {optimal_days} jour(s).")

    # Planification finale avec paramètres effectifs
    itinerary, sites_ordered, coords_ordered, stats = schedule_itinerary(
        coords=coords,
        sites=all_sites,
        order=order,
        segments_summary=segments,
        start_date=start_date,
        start_activity_time=start_activity_time,
        end_activity_time=end_activity_time,
        start_travel_time=start_travel_time,
        end_travel_time=end_travel_time,
        use_lunch=use_lunch,
        lunch_start_time=lunch_start_time if use_lunch else time(12,30),
        lunch_end_time=lunch_end_time if use_lunch else time(14,0),
        use_prayer=use_prayer,
        prayer_start_time=prayer_start_time if use_prayer else time(14,0),
        prayer_duration_min=prayer_duration_min if use_prayer else 20,
        lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
        max_days=effective_max_days,
        tolerance_hours=tolerance_hours,
        base_location=base_location,
        stretch_days=stretch_days_flag,
        allow_weekend_travel=allow_weekend_travel,
        allow_weekend_activities=allow_weekend_activities
    )

    if stretch_days_flag and stats.get('total_days', 0) > effective_max_days:
        st.error(f"❌ Impossible de tenir en {effective_max_days} jour(s). Besoin de {stats.get('total_days')} jours même en étirant les journées.")
    
    progress.progress(1.0)
    status.text("✅ Terminé!")
    
    st.session_state.planning_results = {
        'itinerary': itinerary,
        'sites_ordered': sites_ordered,
        'coords_ordered': coords_ordered,
        'route_polyline': None,
        'stats': stats,
        'start_date': start_date,
        'calculation_method': calculation_method,
        'segments_summary': segments,
        'original_order': order.copy(),  # Sauvegarder l'ordre original
        'durations_matrix': durations_sec,
        'distances_matrix': distances_m,
        'all_coords': coords,
        'base_location': base_location
    }
    st.session_state.manual_itinerary = None
    st.session_state.edit_mode = False

    # Nettoyer l'animation et la barre de progression pour éviter le spinner persistant
    try:
        progress.empty()
    except Exception:
        pass
    try:
        animation_container.empty()
    except Exception:
        pass

# --------------------------
# AFFICHAGE RÉSULTATS
# --------------------------
if st.session_state.planning_results:
    results = st.session_state.planning_results
    itinerary = st.session_state.manual_itinerary if st.session_state.manual_itinerary else results['itinerary']
    sites_ordered = results['sites_ordered']
    coords_ordered = results['coords_ordered']
    stats = results['stats']
    start_date = results['start_date']
    calculation_method = results.get('calculation_method', 'Inconnu')
    segments_summary = results.get('segments_summary', [])
    
    st.header("📊 Résumé de la mission")
    
    method_color = "success" if "Maps" in calculation_method else "info" if "Automatique" in calculation_method else "warning"
    # Message supprimé pour alléger l'UI
# st.caption("📊 Distances calculées")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Durée totale", f"{stats['total_days']} jour(s)")
    with col2:
        st.metric("Distance totale", f"{stats['total_km']:.1f} km")
    with col3:
        # Compter seulement les vrais sites (exclure les sites de type "Base")
        actual_sites_count = len([site for site in sites_ordered if site.get('Type') != 'Base'])
        st.metric("Sites visités", f"{actual_sites_count}")
    with col4:
        st.metric("Temps de visite", f"{stats['total_visit_hours']:.1f} h")
    
    tab_planning, tab_map, tab_fuel, tab_edit, tab_manual, tab_request, tab_report, tab_export = st.tabs(["📅 Planning", "🗺️ Carte", "⛽ Carburant", "✏️ Éditer", "🔄 Modifier ordre", "📑 Demande de mission", "📋 Rapport", "💾 Export"])
    
    with tab_planning:
        st.subheader("Planning détaillé")
        
        view_mode = st.radio(
            "Mode d'affichage",
            ["📋 Vue interactive", "🎨 Présentation professionnelle"],
            horizontal=True,
            index=1
        )
        
        if view_mode == "🎨 Présentation professionnelle":
            include_map_prof = st.checkbox("Inclure la carte", value=st.session_state.get("include_map_prof_html", False))
            st.session_state.include_map_prof_html = include_map_prof

            include_prof_details = st.checkbox(
                "Inclure section résumé",
                value=st.session_state.get("include_prof_details", False)
            )
            st.session_state.include_prof_details = include_prof_details

            html_str = build_professional_html(
                itinerary,
                start_date,
                stats,
                sites_ordered,
                segments_summary,
                default_speed_kmh,
                mission_title,
                coords_ordered,
                include_map=include_map_prof,
                lunch_start_time=st.session_state.get("lunch_start_time"),
                lunch_end_time=st.session_state.get("lunch_end_time"),
                lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
                prayer_start_time=st.session_state.get("prayer_start_time"),
                prayer_duration_min=st.session_state.get("prayer_duration_min", 20),
                include_details=include_prof_details
            )
            st.components.v1.html(html_str, height=1100, scrolling=True)
            
            col_html, col_pdf = st.columns(2)
            
            with col_html:
                st.download_button(
                    label="📥 Télécharger HTML",
                    data=html_str,
                    file_name=f"mission_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                    mime="text/html"
                )
            
            with col_pdf:
                try:
                    excel_data = create_mission_excel(
                        itinerary=itinerary,
                        start_date=start_date,
                        stats=stats,
                        sites_ordered=sites_ordered,
                        segments_summary=segments_summary,
                        mission_title=mission_title
                    )
                    st.download_button(
                        label="📊 Télécharger Excel",
                        data=excel_data,
                        file_name=f"mission_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                except Exception as e:
                    st.error(f"❌ Erreur lors de la génération du fichier Excel: {str(e)}")
        
        else:
            total_days = max(ev[0] for ev in itinerary) if itinerary else 1
            
            if total_days > 1:
                selected_day = st.selectbox(
                    "Jour",
                    options=range(1, total_days + 1),
                    format_func=lambda x: f"Jour {x} - {(start_date + timedelta(days=x-1)).strftime('%d/%m/%Y')}"
                )
            else:
                selected_day = 1
            
            day_events = [ev for ev in itinerary if ev[0] == selected_day]
            
            if day_events:
                date_str = (start_date + timedelta(days=selected_day-1)).strftime("%A %d %B %Y")
                st.info(f"**{date_str}**")
                
                for day, sdt, edt, desc in day_events:
                    col1, col2 = st.columns([1, 3])
                    with col1:
                        st.write(f"**{sdt.strftime('%H:%M')} - {edt.strftime('%H:%M')}**")
                    with col2:
                        if "→" in desc:
                            st.write(f"🚗 {desc}")
                        elif "Visite" in desc or "Site" in desc or "Client" in desc:
                            st.success(desc)
                        elif "Pause" in desc or "Déjeuner" in desc or "Prière" in desc:
                            st.info(desc)
                        elif "Nuitée" in desc:
                            st.warning(desc)
                        else:
                            st.write(desc)
    
    with tab_fuel:
        st.subheader("⛽ Module Carburant")
        
        if st.session_state.planning_results is None:
            st.warning("⚠️ Veuillez d'abord générer un planning dans l'onglet 'Planning' pour calculer la consommation de carburant.")
        else:
            # Récupérer la distance totale du planning
            stats = st.session_state.planning_results.get('stats', {})
            total_distance_km = stats.get('total_km', 0)
            
            if total_distance_km > 0:
                st.info(f"📏 **Distance totale de la mission :** {total_distance_km:.1f} km")
                
                # Sélection du type de véhicule
                st.subheader("🚗 Sélection du véhicule")
                
                vehicle_types = get_vehicle_types()
                vehicle_names = list(vehicle_types.keys())
                
                # Station-Wagon par défaut
                default_index = vehicle_names.index("Station-Wagon") if "Station-Wagon" in vehicle_names else 0
                
                selected_vehicle = st.selectbox(
                    "Type de véhicule",
                    options=vehicle_names,
                    index=default_index,
                    help="Sélectionnez le type de véhicule pour calculer la consommation"
                )
                
                # Affichage des caractéristiques du véhicule sélectionné
                vehicle_info = vehicle_types[selected_vehicle]
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Consommation", f"{vehicle_info['consumption']} L/100km")
                with col2:
                    st.metric("Facteur CO₂", f"{vehicle_info['co2_factor']} kg CO₂/L")
                
                st.divider()
                
                # Calculs de consommation et d'empreinte carbone
                fuel_data = calculate_fuel_consumption(total_distance_km, selected_vehicle)
                carbon_data = calculate_carbon_footprint(fuel_data, total_distance_km, selected_vehicle)
                cost_data = estimate_fuel_cost(fuel_data)
                
                # Affichage des résultats
                st.subheader("📊 Résultats des calculs")
                
                # Métriques principales
                col1, col2 = st.columns(2)
                
                with col1:
                    st.metric(
                        "🛢️ Carburant nécessaire",
                        f"{fuel_data['fuel_needed_liters']:.1f} L",
                        help="Quantité de carburant nécessaire pour la mission"
                    )
                
                with col2:
                    st.metric(
                        "🌍 CO₂ émis",
                        f"{carbon_data['co2_emissions_kg']:.1f} kg",
                        help="Émissions de CO₂ pour la mission"
                    )
                
                st.divider()
                
                # Détails de l'empreinte carbone
                st.subheader("🌱 Empreinte carbone détaillée")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Émissions CO₂ :**")
                    st.write(f"• En kilogrammes : **{carbon_data['co2_emissions_kg']:.2f} kg**")
                    st.write(f"• En tonnes : **{carbon_data['co2_emissions_tons']:.3f} tonnes**")
                
                with col2:
                    st.write("**Équivalence environnementale :**")
                    st.write(f"• Arbres à planter pour compenser : **{carbon_data['trees_equivalent']:.0f} arbres**")
                    st.write("• *(1 arbre absorbe ~22 kg CO₂/an)*")

                # Message d'engagement environnemental
                try:
                    trees_to_plant = int(carbon_data['trees_equivalent'] + 0.9999)  # Arrondi à l'entier supérieur
                except Exception:
                    trees_to_plant = int(round(carbon_data.get('trees_equivalent', 0)))

                st.warning(
                    f"🌿 Pour un engagement en faveur de l'environnement, engagez-vous à planter au moins "
                    f"**{trees_to_plant} arbre(s)** lors de votre mission."
                )

                st.info(
                    "Conseils éco-responsables: privilégiez l'éco-conduite et le covoiturage lors des missions, maintenez une pression "
                    "des pneus optimale, limitez la climatisation et optimisez vos trajets "
                    "pour réduire les kilomètres à vide."
                )
                
                st.divider()
                
                # Section demande de carburant
                st.subheader("📋 Demande de véhicule ou de carburant")
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("📝 Générer demande de carburant", type="primary", use_container_width=True):
                        st.session_state.show_fuel_request_modal = True
                
                with col_btn2:
                    st.info("💡 Génère un document Word")
                
                # Modal pour la demande de carburant
                if st.session_state.get('show_fuel_request_modal', False):
                    with st.container():
                        st.markdown("---")
                        st.subheader("📋 Informations pour la demande de carburant")
                        st.info("💡 Remplissez les informations manquantes pour générer le document")
                        
                        # Afficher les informations de la mission si disponibles
                        if st.session_state.planning_results:
                            stats = st.session_state.planning_results['stats']
                            with st.expander("📊 Informations de la mission", expanded=True):
                                col_info1, col_info2, col_info3 = st.columns(3)
                                with col_info1:
                                    st.metric("🗓️ Durée", f"{stats['total_days']} jour(s)")
                                with col_info2:
                                    # Utiliser le nombre de sites configurés par l'utilisateur
                                    nb_sites = len(st.session_state.sites_df) if 'sites_df' in st.session_state else 0
                                    st.metric("📍 Sites à visiter", f"{nb_sites}")
                                with col_info3:
                                    st.metric("🛣️ Distance totale", f"{stats.get('total_km', 0):.1f} km")
                        
                        # Informations du demandeur
                        col_req1, col_req2 = st.columns(2)
                        
                        with col_req1:
                            st.markdown("**👤 Informations du demandeur**")
                            demandeur_nom = st.text_input("Nom et qualité du demandeur", 
                                                        value="",
                                                        placeholder="Ex: Moctar TALL Responsable de projets",
                                                        key="fuel_req_nom")
                            demandeur_dir = st.text_input("Direction/Département", 
                                                        value="",
                                                        placeholder="Ex: DAL/GPR",
                                                        key="fuel_req_dir")
                            demandeur_cr = st.text_input("N° C.R.", 
                                                       value="",
                                                       placeholder="Ex: L2100",
                                                       key="fuel_req_cr")
                            demandeur_tel = st.text_input("N° Téléphone", 
                                                        value="",
                                                        placeholder="Ex: 77 639 96 12",
                                                        key="fuel_req_tel")
                        
                        with col_req2:
                            st.markdown("**📋 Détails de la mission**")
                            motif_demande = st.text_area("Motif de la demande", 
                                                       value=mission_title,
                                                       key="fuel_req_motif",
                                                       height=100)
                            
                            col_nb_pers, col_carburant = st.columns(2)
                            with col_nb_pers:
                                nb_personnes = st.number_input("Nombre de personnes", 
                                                             min_value=1, max_value=20, value=2,
                                                             key="fuel_req_nb_pers")
                            
                            with col_carburant:
                                # Utiliser le même calcul que dans les résultats des calculs
                                default_fuel = 50
                                if st.session_state.planning_results:
                                    distance = st.session_state.planning_results['stats'].get('total_km', 0)
                                    if distance > 0:
                                        # Utiliser le même calcul que dans la section "Carburant nécessaire"
                                        # Par défaut, utiliser Station-Wagon (8.5 L/100km)
                                        fuel_data = calculate_fuel_consumption(distance, "Station-Wagon")
                                        default_fuel = int(fuel_data['fuel_needed_liters'])
                                    else:
                                        default_fuel = 50
                                
                                quantite_carburant = st.number_input("Quantité de carburant (litres)", 
                                                                   min_value=0, max_value=1000, 
                                                                   value=default_fuel,
                                                                   key="fuel_req_quantity",
                                                                   help="Quantité calculée automatiquement selon la distance et le type de véhicule (Station-Wagon par défaut)")
                        
                        # Dates automatiquement récupérées du planning
                        if st.session_state.planning_results:
                            # Récupérer les dates du planning
                            planning_start_date = st.session_state.planning_results['start_date']
                            itinerary = st.session_state.manual_itinerary or st.session_state.planning_results['itinerary']
                            stats = st.session_state.planning_results['stats']
                            
                            # Calculer la date de retour (date de début + nombre de jours - 1)
                            planning_end_date = planning_start_date + timedelta(days=stats['total_days'] - 1)
                            
                            # Afficher les dates récupérées du planning
                            col_date1, col_date2 = st.columns(2)
                            with col_date1:
                                st.markdown("**📅 Date de départ**")
                                st.info(f"🗓️ {planning_start_date.strftime('%d/%m/%Y')}")
                                date_depart = planning_start_date
                            with col_date2:
                                st.markdown("**📅 Date de retour**")
                                st.info(f"🗓️ {planning_end_date.strftime('%d/%m/%Y')}")
                                date_retour = planning_end_date
                        else:
                            # Si pas de planning, utiliser les champs manuels
                            col_date1, col_date2 = st.columns(2)
                            with col_date1:
                                date_depart = st.date_input("Date de départ prévue", 
                                                          value=datetime.now().date(),
                                                          key="fuel_req_date_dep")
                            with col_date2:
                                date_retour = st.date_input("Date de retour prévue", 
                                                          value=datetime.now().date(),
                                                          key="fuel_req_date_ret")
                        
                        # Boutons d'action
                        col_action1, col_action2, col_action3 = st.columns(3)
                        
                        with col_action1:
                            if st.button("📄 Générer document Word", type="primary", use_container_width=True):
                                # Générer le document Word
                                try:
                                    from docx import Document
                                    from docx.shared import Inches, Pt, Cm
                                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                                    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
                                    from docx.oxml.shared import OxmlElement, qn
                                    from docx.oxml.ns import nsdecls
                                    from docx.oxml import parse_xml
                                    import io
                                    
                                    # Créer le document
                                    doc = Document()
                                    
                                    # Définir les marges
                                    sections = doc.sections
                                    for section in sections:
                                        section.top_margin = Cm(2)
                                        section.bottom_margin = Cm(2)
                                        section.left_margin = Cm(2)
                                        section.right_margin = Cm(2)
                                    
                                    # En-tête principal
                                    header_para = doc.add_paragraph()
                                    header_run = header_para.add_run('DEMANDE DE CARBURANT')
                                    header_run.font.name = 'Tahoma'
                                    header_run.font.size = Pt(14)
                                    header_run.bold = True
                                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Sous-titre
                                    subtitle_para = doc.add_paragraph()
                                    subtitle_run = subtitle_para.add_run('A remplir et à déposer à la DAL/GPR')
                                    subtitle_run.font.name = 'Tahoma'
                                    subtitle_run.font.size = Pt(11)
                                    subtitle_run.italic = True
                                    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Espace
                                    doc.add_paragraph()
                                    
                                    # Numéro de demande avec encadrement
                                    num_table = doc.add_table(rows=1, cols=1)
                                    num_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    num_cell = num_table.cell(0, 0)
                                    num_cell.width = Cm(6)
                                    num_para = num_cell.paragraphs[0]
                                    num_run = num_para.add_run('N°')
                                    num_run.font.name = 'Tahoma'
                                    num_run.font.size = Pt(11)
                                    num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Bordures pour le numéro
                                    def set_cell_border(cell, **kwargs):
                                        tc = cell._tc
                                        tcPr = tc.get_or_add_tcPr()
                                        tcBorders = tcPr.first_child_found_in("w:tcBorders")
                                        if tcBorders is None:
                                            tcBorders = OxmlElement('w:tcBorders')
                                            tcPr.append(tcBorders)
                                        
                                        for edge in ('top', 'left', 'bottom', 'right'):
                                            edge_data = kwargs.get(edge)
                                            if edge_data:
                                                tag = 'w:{}'.format(edge)
                                                element = tcBorders.find(qn(tag))
                                                if element is None:
                                                    element = OxmlElement(tag)
                                                    tcBorders.append(element)
                                                for key, value in edge_data.items():
                                                    element.set(qn('w:{}'.format(key)), str(value))
                                    
                                    border_kwargs = {
                                        'top': {'sz': 12, 'val': 'single', 'color': '000000'},
                                        'bottom': {'sz': 12, 'val': 'single', 'color': '000000'},
                                        'left': {'sz': 12, 'val': 'single', 'color': '000000'},
                                        'right': {'sz': 12, 'val': 'single', 'color': '000000'}
                                    }
                                    set_cell_border(num_cell, **border_kwargs)
                                    
                                    # Espace
                                    doc.add_paragraph()
                                    
                                    # Tableau principal des informations
                                    main_table = doc.add_table(rows=2, cols=2)
                                    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    
                                    # Définir les largeurs des colonnes
                                    main_table.columns[0].width = Cm(8)
                                    main_table.columns[1].width = Cm(8)
                                    
                                    # Première ligne - Nom du demandeur
                                    cell_demandeur = main_table.cell(0, 0)
                                    cell_demandeur.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                    para_demandeur = cell_demandeur.paragraphs[0]
                                    run_title = para_demandeur.add_run('Nom et qualité du demandeur\n')
                                    run_title.font.name = 'Tahoma'
                                    run_title.font.size = Pt(11)
                                    run_title.bold = True
                                    
                                    run_name = para_demandeur.add_run(f'{demandeur_nom}\n\n')
                                    run_name.font.name = 'Tahoma'
                                    run_name.font.size = Pt(11)
                                    
                                    run_details = para_demandeur.add_run(f'DIR. /DEP. : {demandeur_dir}\nN° C.R.     : {demandeur_cr}\nN° Tél.     : {demandeur_tel}')
                                    run_details.font.name = 'Tahoma'
                                    run_details.font.size = Pt(11)
                                    
                                    # Première ligne - Motif de la demande
                                    cell_motif = main_table.cell(0, 1)
                                    cell_motif.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                    para_motif = cell_motif.paragraphs[0]
                                    run_motif_title = para_motif.add_run('Motif de la demande\n\n')
                                    run_motif_title.font.name = 'Tahoma'
                                    run_motif_title.font.size = Pt(11)
                                    run_motif_title.bold = True
                                    
                                    run_motif_content = para_motif.add_run(motif_demande)
                                    run_motif_content.font.name = 'Tahoma'
                                    run_motif_content.font.size = Pt(11)
                                    
                                    # Deuxième ligne - Dates
                                    cell_dates = main_table.cell(1, 0)
                                    cell_dates.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                    para_dates = cell_dates.paragraphs[0]
                                    
                                    run_depart = para_dates.add_run(f'Départ prévu : {date_depart.strftime("%d/%m/%Y")}\n\n')
                                    run_depart.font.name = 'Tahoma'
                                    run_depart.font.size = Pt(11)
                                    
                                    run_retour = para_dates.add_run(f'Retour prévu : {date_retour.strftime("%d/%m/%Y")}')
                                    run_retour.font.name = 'Tahoma'
                                    run_retour.font.size = Pt(11)
                                    
                                    # Deuxième ligne - Nombre de personnes et quantité de carburant
                                    cell_nb = main_table.cell(1, 1)
                                    cell_nb.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                    para_nb = cell_nb.paragraphs[0]
                                    
                                    run_nb = para_nb.add_run(f'Nombre de personnes : {nb_personnes:02d}\n\n')
                                    run_nb.font.name = 'Tahoma'
                                    run_nb.font.size = Pt(11)
                                    
                                    # Nouveau champ pour la quantité de carburant
                                    run_carburant_title = para_nb.add_run('Quantité de carburant demandée :\n\n')
                                    run_carburant_title.font.name = 'Tahoma'
                                    run_carburant_title.font.size = Pt(11)
                                    run_carburant_title.bold = True
                                    
                                    run_carburant_value = para_nb.add_run(f'{quantite_carburant} litres')
                                    run_carburant_value.font.name = 'Tahoma'
                                    run_carburant_value.font.size = Pt(11)
                                    
                                    # Appliquer les bordures au tableau principal
                                    for row in main_table.rows:
                                        for cell in row.cells:
                                            set_cell_border(cell, **border_kwargs)
                                    
                                    # Espace
                                    doc.add_paragraph()
                                    
                                    # Tableau itinéraire
                                    itinerary_table = doc.add_table(rows=1, cols=2)
                                    itinerary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    itinerary_table.columns[0].width = Cm(12)
                                    itinerary_table.columns[1].width = Cm(4)
                                    
                                    # En-têtes du tableau itinéraire
                                    hdr_cells = itinerary_table.rows[0].cells
                                    
                                    hdr_para1 = hdr_cells[0].paragraphs[0]
                                    hdr_run1 = hdr_para1.add_run('Itinéraire à suivre')
                                    hdr_run1.font.name = 'Tahoma'
                                    hdr_run1.font.size = Pt(11)
                                    hdr_run1.bold = True
                                    hdr_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    hdr_para2 = hdr_cells[1].paragraphs[0]
                                    hdr_run2 = hdr_para2.add_run('KM')
                                    hdr_run2.font.name = 'Tahoma'
                                    hdr_run2.font.size = Pt(11)
                                    hdr_run2.bold = True
                                    hdr_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Ajouter les sites de la mission ou lignes vides
                                    if st.session_state.planning_results:
                                        sites = st.session_state.planning_results.get('sites_ordered', [])
                                        segments = st.session_state.planning_results.get('segments_summary', [])
                                        base_location = st.session_state.planning_results.get('base_location', '')
                                        
                                        # Commencer à partir du deuxième site pour éviter d'afficher juste "Dakar"
                                        for i in range(1, len(sites)):
                                            row_cells = itinerary_table.add_row().cells
                                            
                                            para_site = row_cells[0].paragraphs[0]
                                            prev_site = sites[i-1]
                                            current_site = sites[i]
                                            site_text = f"{prev_site['Ville']} → {current_site['Ville']}"
                                            
                                            run_site = para_site.add_run(site_text)
                                            run_site.font.name = 'Tahoma'
                                            run_site.font.size = Pt(11)
                                            
                                            para_km = row_cells[1].paragraphs[0]
                                            # Utiliser l'index i-1 pour les segments car on commence à i=1
                                            if (i-1) < len(segments):
                                                distance_km = segments[i-1]['distance'] / 1000
                                                km_text = f"{distance_km:.1f}"
                                            else:
                                                km_text = "___"
                                            
                                            run_km = para_km.add_run(km_text)
                                            run_km.font.name = 'Tahoma'
                                            run_km.font.size = Pt(11)
                                            para_km.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        
                                        # Distance totale
                                        total_row = itinerary_table.add_row().cells
                                        para_total = total_row[0].paragraphs[0]
                                        run_total = para_total.add_run('Distance totale :')
                                        run_total.font.name = 'Tahoma'
                                        run_total.font.size = Pt(11)
                                        run_total.bold = True
                                        
                                        para_total_km = total_row[1].paragraphs[0]
                                        run_total_km = para_total_km.add_run(f"{stats.get('total_km', 0):.1f}")
                                        run_total_km.font.name = 'Tahoma'
                                        run_total_km.font.size = Pt(11)
                                        run_total_km.bold = True
                                        para_total_km.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        # Lignes vides si pas de planning
                                        for _ in range(8):
                                            row_cells = itinerary_table.add_row().cells
                                            para_empty = row_cells[0].paragraphs[0]
                                            run_empty = para_empty.add_run("")
                                            run_empty.font.name = 'Tahoma'
                                            run_empty.font.size = Pt(11)
                                            
                                            para_empty_km = row_cells[1].paragraphs[0]
                                            run_empty_km = para_empty_km.add_run("")
                                            run_empty_km.font.name = 'Tahoma'
                                            run_empty_km.font.size = Pt(11)
                                        
                                        # Distance totale vide
                                        total_row = itinerary_table.add_row().cells
                                        para_total = total_row[0].paragraphs[0]
                                        run_total = para_total.add_run('Distance totale :')
                                        run_total.font.name = 'Tahoma'
                                        run_total.font.size = Pt(11)
                                        run_total.bold = True
                                        
                                        para_total_km = total_row[1].paragraphs[0]
                                        run_total_km = para_total_km.add_run("")
                                        run_total_km.font.name = 'Tahoma'
                                        run_total_km.font.size = Pt(11)
                                    
                                    # Appliquer les bordures au tableau itinéraire
                                    for row in itinerary_table.rows:
                                        for cell in row.cells:
                                            set_cell_border(cell, **border_kwargs)
                                    
                                    # Espace réduit
                                    doc.add_paragraph()
                                    
                                    # Date
                                    date_para = doc.add_paragraph()
                                    date_run = date_para.add_run(f'Date : Le {datetime.now().strftime("%d/%m/%Y")}')
                                    date_run.font.name = 'Tahoma'
                                    date_run.font.size = Pt(11)
                                    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    
                                    # Tableau des signatures
                                    sig_table = doc.add_table(rows=1, cols=3)
                                    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    
                                    # Définir les largeurs des colonnes de signature
                                    sig_table.columns[0].width = Cm(5.3)
                                    sig_table.columns[1].width = Cm(5.3)
                                    sig_table.columns[2].width = Cm(5.3)
                                    
                                    # Contenu des cellules de signature
                                    sig_cells = sig_table.rows[0].cells
                                    
                                    # Première signature
                                    para_sig1 = sig_cells[0].paragraphs[0]
                                    run_sig1_title = para_sig1.add_run('Signature et cachet\n')
                                    run_sig1_title.font.name = 'Tahoma'
                                    run_sig1_title.font.size = Pt(11)
                                    run_sig1_title.bold = True
                                    para_sig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    run_sig1_subtitle = para_sig1.add_run('Chef de service Demandeur')
                                    run_sig1_subtitle.font.name = 'Tahoma'
                                    run_sig1_subtitle.font.size = Pt(11)
                                    
                                    # Deuxième signature
                                    para_sig2 = sig_cells[1].paragraphs[0]
                                    run_sig2 = para_sig2.add_run('Responsable POOL')
                                    run_sig2.font.name = 'Tahoma'
                                    run_sig2.font.size = Pt(11)
                                    run_sig2.bold = True
                                    para_sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Troisième signature
                                    para_sig3 = sig_cells[2].paragraphs[0]
                                    run_sig3 = para_sig3.add_run('DAL/GPR')
                                    run_sig3.font.name = 'Tahoma'
                                    run_sig3.font.size = Pt(11)
                                    run_sig3.bold = True
                                    para_sig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Définir la hauteur des cellules de signature
                                    for cell in sig_cells:
                                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                        # Ajouter de l'espace pour les signatures
                                        for _ in range(4):
                                            cell.add_paragraph()
                                    
                                    # Appliquer les bordures au tableau de signatures
                                    for row in sig_table.rows:
                                        for cell in row.cells:
                                            set_cell_border(cell, **border_kwargs)
                                    
                                    # Sauvegarder dans un buffer
                                    buffer = io.BytesIO()
                                    doc.save(buffer)
                                    buffer.seek(0)
                                    
                                    # Bouton de téléchargement
                                    st.success("✅ Document généré avec succès!")
                                    st.download_button(
                                        label="📥 Télécharger la demande de carburant (Word)",
                                        data=buffer.getvalue(),
                                        file_name=f"Demande_carburant_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                    
                                except ImportError:
                                    st.error("❌ Le module python-docx n'est pas installé. Veuillez l'installer avec: pip install python-docx")
                                except Exception as e:
                                    st.error(f"❌ Erreur lors de la génération du document: {str(e)}")
                        
                        with col_action2:
                            if st.button("❌ Annuler", use_container_width=True):
                                st.session_state.show_fuel_request_modal = False
                                st.rerun()


                
            else:
                st.error("❌ Aucune distance calculée. Vérifiez votre planning.")
    
    with tab_edit:
        st.subheader("✏️ Édition manuelle du planning")
        
        st.info("💡 Modifiez les horaires, ajoutez ou supprimez des événements. Les modifications sont automatiquement sauvegardées.")
        
        # Initialiser manual_itinerary si nécessaire
        if st.session_state.manual_itinerary is None:
            st.session_state.manual_itinerary = list(itinerary)
        
        # Sélection du jour
        total_days = max(ev[0] for ev in st.session_state.manual_itinerary) if st.session_state.manual_itinerary else 1
        
        selected_edit_day = st.selectbox(
            "Sélectionnez le jour à éditer",
            options=range(1, total_days + 1),
            format_func=lambda x: f"Jour {x} - {(start_date + timedelta(days=x-1)).strftime('%d/%m/%Y')}",
            key="edit_day_select"
        )
        
        # Filtrer les événements du jour
        day_events_edit = [(i, ev) for i, ev in enumerate(st.session_state.manual_itinerary) if ev[0] == selected_edit_day]
        
        st.markdown("---")
        
        # Afficher chaque événement avec possibilité d'édition
        for idx, (global_idx, (day, sdt, edt, desc)) in enumerate(day_events_edit):
            with st.expander(f"**Événement {idx+1}** : {desc[:50]}...", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    new_start = st.time_input(
                        "Heure de début",
                        value=sdt.time(),
                        key=f"start_{global_idx}"
                    )
                
                with col2:
                    new_end = st.time_input(
                        "Heure de fin",
                        value=edt.time(),
                        key=f"end_{global_idx}"
                    )
                
                new_desc = st.text_area(
                    "Description",
                    value=desc,
                    height=100,
                    key=f"desc_{global_idx}"
                )

                # Saisie manuelle des distances pour les trajets
                override_km = None
                override_h = None
                override_m = None
                if "→" in desc:
                    import re
                    km_val = 0.0
                    h_val = 0
                    m_val = 0
                    # Essayer d'extraire (123.4 km, 2h30) ou (123.4 km)
                    m_with_time = re.search(r"\((['\d\.']+)\s*km,\s*([^)]+)\)", desc)
                    if m_with_time:
                        try:
                            km_val = float(m_with_time.group(1))
                            time_str = m_with_time.group(2).strip()
                            m_time = re.match(r"(\d+)\s*h\s*(\d{1,2})", time_str)
                            if m_time:
                                h_val = int(m_time.group(1))
                                m_val = int(m_time.group(2))
                        except Exception:
                            pass
                    else:
                        m_km = re.search(r"\(([\d\.]+)\s*km\)", desc)
                        if m_km:
                            try:
                                km_val = float(m_km.group(1))
                            except Exception:
                                pass
                    st.markdown("**Distance du trajet (si différente)**")
                    col_d1, col_d2, col_d3 = st.columns([1, 1, 1])
                    with col_d1:
                        override_km = st.number_input("Distance (km)", min_value=0.0, value=float(km_val), key=f"km_{global_idx}")
                    with col_d2:
                        override_h = st.number_input("Heures", min_value=0, value=int(h_val), key=f"kh_{global_idx}")
                    with col_d3:
                        override_m = st.number_input("Minutes", min_value=0, max_value=59, value=int(m_val), key=f"kmn_{global_idx}")
                
                col_btn1, col_btn2, col_btn3 = st.columns(3)
                
                with col_btn1:
                    if st.button("💾 Sauvegarder", key=f"save_{global_idx}", use_container_width=True):
                        new_sdt = datetime.combine(sdt.date(), new_start)
                        new_edt = datetime.combine(edt.date(), new_end)
                        # Appliquer éventuelles distances sur la description (pour les trajets)
                        updated_desc = new_desc
                        if override_km is not None:
                            import re
                            # Retirer ancien motif distance s'il existe
                            updated_desc = re.sub(r"\s*\([\d\.\s]+km(?:,\s*[^)]*)?\)\s*$", "", updated_desc).strip()
                            if override_km > 0:
                                if (override_h or 0) > 0 or (override_m or 0) > 0:
                                    updated_desc = f"{updated_desc} ({override_km} km, {int(override_h or 0)}h{int(override_m or 0):02d})"
                                else:
                                    updated_desc = f"{updated_desc} ({override_km} km)"
                        st.session_state.manual_itinerary[global_idx] = (day, new_sdt, new_edt, updated_desc)
                        st.success("Modifications sauvegardées!")
                        st.rerun()
                
                with col_btn2:
                    if st.button("🗑️ Supprimer", key=f"delete_{global_idx}", use_container_width=True):
                        st.session_state.manual_itinerary.pop(global_idx)
                        st.success("Événement supprimé!")
                        st.rerun()
                
                with col_btn3:
                    if st.button("↕️ Déplacer", key=f"move_{global_idx}", use_container_width=True):
                        st.session_state.editing_event = global_idx
        
        # Ajouter un nouvel événement
        st.markdown("---")
        st.subheader("➕ Ajouter un événement")
        
        with st.form("add_event_form"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                new_event_start = st.time_input("Début", value=time(8, 0))
            
            with col2:
                new_event_end = st.time_input("Fin", value=time(9, 0))
            
            with col3:
                event_type = st.selectbox(
                    "Type",
                    ["Visite", "Trajet", "Pause", "Autre"]
                )
            
            new_event_desc = st.text_input("Description", value="Nouvel événement")
            
            if st.form_submit_button("Ajouter l'événement"):
                event_date = start_date + timedelta(days=selected_edit_day-1)
                new_sdt = datetime.combine(event_date, new_event_start)
                new_edt = datetime.combine(event_date, new_event_end)
                
                prefix = ""
                if event_type == "Trajet":
                    prefix = "🚗 "
                elif event_type == "Pause":
                    prefix = "⏸️ "
                elif event_type == "Visite":
                    prefix = ""
                
                new_event = (selected_edit_day, new_sdt, new_edt, f"{prefix}{new_event_desc}")
                st.session_state.manual_itinerary.append(new_event)
                st.session_state.manual_itinerary.sort(key=lambda x: (x[0], x[1]))
                st.success("Événement ajouté!")
                st.rerun()
        
        # Boutons d'action globaux
        st.markdown("---")
        col_reset, col_recalc = st.columns(2)
        
        with col_reset:
            if st.button("🔄 Réinitialiser les modifications", use_container_width=True):
                st.session_state.manual_itinerary = None
                st.success("Planning réinitialisé!")
                st.rerun()
        
        with col_recalc:
            if st.button("🔢 Recalculer les statistiques", use_container_width=True):
                # Recalculer les stats basées sur manual_itinerary
                total_km = 0
                total_visit_hours = 0
                
                for day, sdt, edt, desc in st.session_state.manual_itinerary:
                    import re
                    # Compter les kilomètres UNIQUEMENT pour les trajets (flèche ou emoji voiture)
                    if ("→" in desc) or ("🚗" in desc):
                        # Nouveau format avec temps réel: "(... km, HhMM)"
                        m_with_time = re.search(r"\(([\d\.]+)\s*km,\s*([^)]+)\)", desc)
                        if m_with_time:
                            try:
                                total_km += float(m_with_time.group(1))
                            except Exception:
                                pass
                        else:
                            # Ancien format: "(... km)"
                            m = re.search(r"\(([\d\.]+)\s*km\)", desc)
                            if m:
                                try:
                                    total_km += float(m.group(1))
                                except Exception:
                                    pass
                    
                    # Cumuler les heures de visite (éviter de compter les trajets)
                    if any(x in desc for x in ["Visite", "–"]) and "→" not in desc:
                        duration = (edt - sdt).total_seconds() / 3600
                        total_visit_hours += duration
                
                stats['total_km'] = total_km
                stats['total_visit_hours'] = total_visit_hours
                
                st.success("Statistiques recalculées!")
                st.rerun()
    
    with tab_manual:
        st.subheader("🔄 Modification manuelle de l'ordre des sites")
        
        st.info("💡 Réorganisez l'ordre des sites en les faisant glisser. L'itinéraire sera automatiquement recalculé.")
        
        # Vérifier que nous avons les données nécessaires
        if 'original_order' not in results or 'durations_matrix' not in results:
            st.warning("⚠️ Données insuffisantes pour la modification manuelle. Veuillez relancer le calcul.")
        else:
            # Récupérer les données
            original_order = results['original_order']
            durations_matrix = results['durations_matrix']
            distances_matrix = results['distances_matrix']
            all_coords = results['all_coords']
            
            # Éditeur de table (optionnel)
            with st.expander("🧮 Mode tableau (numéro d'ordre)", expanded=False):
                import pandas as pd
                rows = []
                for i, site_idx in enumerate(st.session_state.get('manual_order', original_order)):
                    if isinstance(site_idx, int) and 0 <= site_idx < len(sites_ordered):
                        s = sites_ordered[site_idx]
                        rows.append({
                            "Index": site_idx,
                            "Ville": s['Ville'],
                            "Type": s.get('Type', 'Site'),
                            "Activité": s.get('Activité', 'Activité'),
                            "Ordre": i+1
                        })
                df_order = pd.DataFrame(rows)
                edited_df = st.data_editor(
                    df_order,
                    column_config={
                        "Ordre": st.column_config.NumberColumn("Ordre", min_value=1, max_value=len(rows), step=1),
                        "Index": st.column_config.TextColumn("Index", disabled=True),
                        "Ville": st.column_config.TextColumn("Ville", disabled=True),
                        "Type": st.column_config.TextColumn("Type", disabled=True),
                        "Activité": st.column_config.TextColumn("Activité", disabled=True),
                    },
                    column_order=["Ville", "Type", "Activité", "Ordre"],
                    use_container_width=True,
                    num_rows="fixed"
                )
                colA, colB = st.columns([2,1])
                with colA:
                    if st.button("✅ Appliquer l'ordre (table)", type="primary", use_container_width=True):
                        try:
                            ords = edited_df["Ordre"].tolist()
                            if sorted(ords) != list(range(1, len(rows)+1)):
                                st.error("Veuillez attribuer des numéros d'ordre uniques de 1 à N.")
                            else:
                                new_order = [int(row["Index"]) for _, row in edited_df.sort_values("Ordre").iterrows()]
                                st.session_state.manual_order = new_order
                                st.success("Ordre mis à jour depuis la table!")
                                st.rerun()
                        except Exception as e:
                            st.error(f"Erreur lors de l'application: {str(e)[:120]}...")
                with colB:
                    if st.button("↩️ Restaurer", use_container_width=True):
                        st.session_state.manual_order = original_order.copy()
                        st.success("Ordre original restauré!")
                        st.rerun()
            
            # Créer une liste des sites avec leur ordre actuel
            if 'manual_order' not in st.session_state:
                st.session_state.manual_order = original_order.copy()
            
            # Afficher l'ordre actuel des sites
            st.markdown("**Ordre actuel des sites :**")
            st.info(f"📊 **{len(st.session_state.manual_order)} sites** dans l'ordre actuel")
            
            # Interface pour réorganiser les sites avec conteneur scrollable
            with st.container():
                # Utiliser des boutons pour déplacer les sites
                for i, site_idx in enumerate(st.session_state.manual_order):
                    # Vérifier que l'index est valide et de type entier
                    if isinstance(site_idx, int) and 0 <= site_idx < len(sites_ordered):
                        site = sites_ordered[site_idx]
                        
                        col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                        
                        with col1:
                            st.write(f"**{i+1}.** {site['Ville']} - {site.get('Type', 'Site')} - {site.get('Activité', 'Activité')}")
                        
                        with col2:
                            if i > 0 and st.button("⬆️", key=f"enhanced_up_{i}", help="Monter"):
                                # Échanger avec l'élément précédent
                                st.session_state.manual_order[i], st.session_state.manual_order[i-1] = \
                                    st.session_state.manual_order[i-1], st.session_state.manual_order[i]
                                st.rerun()
                        
                        with col3:
                            if i < len(st.session_state.manual_order) - 1 and st.button("⬇️", key=f"enhanced_down_{i}", help="Descendre"):
                                # Échanger avec l'élément suivant
                                st.session_state.manual_order[i], st.session_state.manual_order[i+1] = \
                                    st.session_state.manual_order[i+1], st.session_state.manual_order[i]
                                st.rerun()
                        
                        with col4:
                            if i != 0 and i != len(st.session_state.manual_order) - 1:  # Ne pas permettre de supprimer le départ et l'arrivée
                                if st.button("🗑️", key=f"remove_{i}", help="Supprimer"):
                                    st.session_state.manual_order.pop(i)
                                    st.rerun()
                    else:
                        # Index invalide - nettoyer
                        st.warning(f"⚠️ Index invalide détecté ({site_idx}), nettoyage en cours...")
                        st.session_state.manual_order = [idx for idx in st.session_state.manual_order if idx < len(sites_ordered)]
                        st.rerun()
            
            st.markdown("---")
            
            # Boutons d'action
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("🔄 Recalculer l'itinéraire", use_container_width=True):
                    # Recalculer l'itinéraire avec le nouvel ordre
                    new_order = st.session_state.manual_order
                    
                    # Recalculer les segments
                    new_segments = []
                    for i in range(len(new_order)-1):
                        from_idx = new_order[i]
                        to_idx = new_order[i+1]
                        
                        if from_idx < len(durations_matrix) and to_idx < len(durations_matrix[0]):
                            duration = durations_matrix[from_idx][to_idx]
                            distance = distances_matrix[from_idx][to_idx] if distances_matrix else 0
                            
                            new_segments.append({
                                "distance": distance,
                                "duration": duration
                            })
                        else:
                            new_segments.append({"distance": 0, "duration": 0})
                    
                    # Recalculer l'itinéraire complet
                    new_sites = [sites_ordered[i] for i in new_order]
                    new_coords = [coords_ordered[i] for i in new_order]
                    # Calcul du nombre de jours optimal (dry-run)
                    _, _, _, new_dry_stats = schedule_itinerary(
                        coords=new_coords,
                        sites=new_sites,
                        order=list(range(len(new_order))),
                        segments_summary=new_segments,
                        start_date=start_date,
                        start_activity_time=time(8, 0),
                        end_activity_time=time(17, 0),
                        start_travel_time=time(7, 0),
                        end_travel_time=time(19, 0),
                        use_lunch=True,
                        lunch_start_time=time(12, 30),
                        lunch_end_time=time(14, 0),
                        use_prayer=False,
                        prayer_start_time=time(14, 0),
                        prayer_duration_min=20,
                        lunch_duration_min=st.session_state.get('lunch_duration_min', 60),
                        max_days=0,
                        tolerance_hours=1.0,
                        base_location=results.get('base_location', ''),
                        allow_weekend_travel=st.session_state.get('allow_weekend_travel', True),
                        allow_weekend_activities=st.session_state.get('allow_weekend_activities', True)
                    )

                    new_optimal_days = int(new_dry_stats.get('total_days', 1))
                    user_max_recalc = int(st.session_state.get('max_days', 0))
                    user_desired_recalc = int(st.session_state.get('desired_days', 0))

                    # Logique de décision pour les jours effectifs lors du recalcul
                    if user_desired_recalc > 0:
                        if user_max_recalc > 0 and user_desired_recalc > user_max_recalc:
                            st.warning(f"Souhait ({user_desired_recalc} jours) > max ({user_max_recalc}). Utilisation du max.")
                            new_effective_days = user_max_recalc
                        else:
                            new_effective_days = user_desired_recalc
                        
                        if new_effective_days < new_optimal_days:
                            new_stretch = True
                            st.warning(f"⚠️ Objectif ({new_effective_days}) < optimal ({new_optimal_days}). Compression.")
                        else:
                            new_stretch = False
                            st.success(f"✅ Ajusté à {new_effective_days} jours.")

                    elif user_max_recalc > 0:
                        if user_max_recalc < new_optimal_days:
                            new_effective_days = user_max_recalc
                            new_stretch = True
                            st.warning(f"⚠️ Objectif ({user_max_recalc}) < optimal ({new_optimal_days}). Compression.")
                        else:
                            new_effective_days = user_max_recalc
                            new_stretch = False
                            st.success(f"✅ Tient en {new_optimal_days} jours (max: {user_max_recalc}).")
                    else:
                        new_effective_days = new_optimal_days
                        new_stretch = False
                        st.info(f"🧮 Jours optimaux recalculés: {new_optimal_days}.")

                    # Planification finale recalculée
                    new_itinerary, new_sites_ordered, new_coords_ordered, new_stats = schedule_itinerary(
                        coords=new_coords,
                        sites=new_sites,
                        order=list(range(len(new_order))),  # Ordre séquentiel car sites déjà réorganisés
                        segments_summary=new_segments,
                        start_date=start_date,
                        start_activity_time=time(8, 0),  # Utiliser les valeurs par défaut ou récupérer depuis session_state
                        end_activity_time=time(17, 0),
                        start_travel_time=time(7, 0),
                        end_travel_time=time(19, 0),
                        use_lunch=True,
                        lunch_start_time=time(12, 30),
                        lunch_end_time=time(14, 0),
                        use_prayer=False,
                        prayer_start_time=time(14, 0),
                        prayer_duration_min=20,
                        lunch_duration_min=st.session_state.get('lunch_duration_min', 60),
                        max_days=new_effective_days,
                        tolerance_hours=1.0,
                        base_location=results.get('base_location', ''),
                        stretch_days=new_stretch,
                        allow_weekend_travel=st.session_state.get('allow_weekend_travel', True),
                        allow_weekend_activities=st.session_state.get('allow_weekend_activities', True)
                    )

                    if new_stretch and new_stats.get('total_days', 0) > new_effective_days:
                        st.error(f"❌ Impossible de tenir en {new_effective_days} jour(s). Besoin de {new_stats.get('total_days')} jours même en étirant les journées.")
                    
                    # Mettre à jour les résultats
                    st.session_state.manual_itinerary = new_itinerary
                    st.session_state.planning_results.update({
                        'sites_ordered': new_sites_ordered,
                        'coords_ordered': new_coords_ordered,
                        'stats': new_stats,
                        'segments_summary': new_segments
                    })
                    
                    st.success("✅ Itinéraire recalculé avec le nouvel ordre!")
                    st.rerun()
            
            with col2:
                if st.button("↩️ Restaurer l'ordre original", use_container_width=True):
                    st.session_state.manual_order = original_order.copy()
                    st.session_state.manual_itinerary = None
                    st.success("Ordre original restauré!")
                    st.rerun()
            
            with col3:
                if st.button("🎯 Optimiser automatiquement", use_container_width=True):
                    # Réoptimiser avec IA Adja
                    try:
                        ai_order, ai_success, ai_message = optimize_route_with_ai(sites_ordered, coords_ordered, base_location, deepseek_api_key)
                        if ai_success and isinstance(ai_order, list):
                            st.session_state.manual_order = ai_order
                            st.success(f"Ordre optimisé automatiquement par IA Adja! {ai_message}")
                        else:
                            # Fallback vers TSP si l'IA Adja échoue ou réponse invalide
                            optimized_order = solve_tsp_fixed_start_end(durations_matrix)
                            st.session_state.manual_order = optimized_order
                            st.warning(f"IA Adja indisponible ou réponse invalide, optimisation TSP utilisée. {ai_message if not ai_success else ''}")
                    except Exception as e:
                        # Fallback vers TSP en cas d'erreur
                        optimized_order = solve_tsp_fixed_start_end(durations_matrix)
                        st.session_state.manual_order = optimized_order
                        st.warning(f"Erreur IA Adja ({str(e)[:50]}...), optimisation TSP utilisée.")
                    st.rerun()

            with col4:
                if st.button("⚙️ Optimiser (OR-Tools)", use_container_width=True):
                    try:
                        # Construire les durées de service à partir des données des sites (Durée (h))
                        service_times_sec = []
                        for i in range(len(sites_ordered)):
                            dur_h = sites_ordered[i].get("Durée (h)", 0)
                            try:
                                service_times_sec.append(int(float(dur_h or 0) * 3600))
                            except Exception:
                                service_times_sec.append(0)

                        optimized_path = solve_tsp_ortools_fixed_start_end(durations_matrix, service_times_sec, time_limit_s=5)
                        st.session_state.manual_order = optimized_path
                        st.success("Ordre optimisé par OR-Tools (matrice OSRM/Maps).")
                    except Exception as e:
                        st.warning(f"OR-Tools indisponible ou erreur: {str(e)[:80]}... Fallback TSP.")
                        st.session_state.manual_order = solve_tsp_fixed_start_end(durations_matrix)
                    st.rerun()
    
    with tab_map:
        st.subheader("Carte de l'itinéraire")
        
        if coords_ordered:
            center_lat = sum(c[1] for c in coords_ordered) / len(coords_ordered)
            center_lon = sum(c[0] for c in coords_ordered) / len(coords_ordered)
            
            m = folium.Map(location=[center_lat, center_lon], zoom_start=7)
            
            # Tracé de l'itinéraire : tentative de récupération de la route réelle via OSRM, sinon fallback sur la ligne droite
            try:
                coord_str = ";".join([f"{c[0]},{c[1]}" for c in coords_ordered])
                url = f"{osrm_base_url.rstrip('/')}/route/v1/driving/{coord_str}?overview=full&geometries=geojson"
                resp = requests.get(url, timeout=10)
                route_pts = None
                if resp.status_code == 200:
                    data = resp.json()
                    if data.get('routes'):
                        geom = data['routes'][0].get('geometry')
                        if isinstance(geom, dict) and geom.get('coordinates'):
                            route_pts = [[lat, lon] for lon, lat in geom['coordinates']]
                if not route_pts:
                    route_pts = [[c[1], c[0]] for c in coords_ordered]
            except Exception:
                route_pts = [[c[1], c[0]] for c in coords_ordered]
            folium.PolyLine(locations=route_pts, color="blue", weight=3, opacity=0.7).add_to(m)
            
            # Export Google Maps (ouvrir et copier)
            try:
                gmaps_pairs = [f"{c[1]},{c[0]}" for c in coords_ordered]
                gmaps_url = "https://www.google.com/maps/dir/" + "/".join(gmaps_pairs) + "/?hl=fr"
                st.markdown(f"[📍 Ouvrir dans Google Maps]({gmaps_url})")
                st.text_input("Lien Google Maps", value=gmaps_url, help="Copiez ce lien pour partager ou ouvrir l'itinéraire dans Google Maps.", label_visibility="collapsed")
            except Exception:
                st.info("Lien Google Maps non disponible")
            
            # Préparer affichage spécial si départ et arrivée sont au même endroit
            n_steps = len(sites_ordered)
            start_end_same = False
            if n_steps >= 2:
                lat0, lon0 = coords_ordered[0][1], coords_ordered[0][0]
                latN, lonN = coords_ordered[-1][1], coords_ordered[-1][0]
                start_end_same = abs(lat0 - latN) < 1e-4 and abs(lon0 - lonN) < 1e-4

            for i, site in enumerate(sites_ordered):
                # Si le départ et l'arrivée sont identiques, afficher un double numéro sur le point de départ et ne pas dupliquer le dernier point
                if i == 0 and start_end_same:
                    bg_color_left = '#2ecc71'  # Vert pour départ
                    bg_color_right = '#e74c3c'  # Rouge pour arrivée
                    html = f"""
<div style=\"display:flex; align-items:center; gap:4px;\">
  <div style=\"background-color:{bg_color_left}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">1</div>
  <div style=\"background-color:{bg_color_right}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{n_steps}</div>
</div>
"""
                    folium.Marker(
                        location=[coords_ordered[i][1], coords_ordered[i][0]],
                        popup=f"Étapes 1 et {n_steps}: {site['Ville']}<br>{site.get('Type', '-')}",
                        tooltip=f"Étapes 1 et {n_steps}: {site['Ville']}",
                        icon=folium.DivIcon(
                            icon_size=(36, 28),
                            icon_anchor=(18, 14),
                            html=html
                        )
                    ).add_to(m)
                    continue
                if start_end_same and i == n_steps - 1:
                    # Ne pas dupliquer l'arrivée si elle est au même endroit que le départ
                    continue

                color = 'green' if i == 0 else 'red' if i == len(sites_ordered)-1 else 'blue'
                icon = 'play' if i == 0 else 'stop' if i == len(sites_ordered)-1 else 'info-sign'
                
                # Icône numérotée via DivIcon (couleur selon étape)
                bg_color = '#2ecc71' if i == 0 else '#e74c3c' if i == len(sites_ordered)-1 else '#3498db'
                folium.Marker(
                    location=[coords_ordered[i][1], coords_ordered[i][0]],
                    popup=f"Étape {i+1}: {site['Ville']}<br>{site.get('Type', '-')}",
                    tooltip=f"Étape {i+1}: {site['Ville']}",
                    icon=folium.DivIcon(
                        icon_size=(28, 28),
                        icon_anchor=(14, 14),
                        html=f"""
<div style=\"background-color:{bg_color}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{i+1}</div>
"""
                    )
                ).add_to(m)
            
            st_folium(m, width=None, height=500, use_container_width=True)
            
            # Téléchargements KML/KMZ de l'itinéraire (points + trace)
            try:
                import zipfile
                from io import BytesIO
                
                # Construire le KML: Placemarks pour chaque étape + LineString pour la route
                doc_name = mission_title if mission_title else "Itinéraire"
                kml_parts = [
                    "<?xml version=\"1.0\" encoding=\"UTF-8\"?>",
                    "<kml xmlns=\"http://www.opengis.net/kml/2.2\">",
                    "  <Document>",
                    f"    <name>{doc_name}</name>"
                ]
                
                # Placemarks pour les points (éviter la duplication si départ=arrivée)
                for i, site in enumerate(sites_ordered):
                    if start_end_same and i == n_steps - 1:
                        continue
                    lon, lat = coords_ordered[i][0], coords_ordered[i][1]
                    site_name = site.get('Ville', '')
                    site_type = site.get('Type', '-')
                    placemark = f"""
    <Placemark>
      <name>Étape {i+1}: {site_name}</name>
      <description>{site_type}</description>
      <Point>
        <coordinates>{lon},{lat},0</coordinates>
      </Point>
    </Placemark>"""
                    kml_parts.append(placemark)
                
                # LineString pour la trace (OSRM si disponible, sinon ligne droite)
                line_coords = "\n".join([f"{pt[1]},{pt[0]},0" for pt in route_pts])
                linestring = f"""
    <Placemark>
      <name>Route</name>
      <Style>
        <LineStyle>
          <color>ff0000ff</color>
          <width>3</width>
        </LineStyle>
      </Style>
      <LineString>
        <tessellate>1</tessellate>
        <coordinates>
{line_coords}
        </coordinates>
      </LineString>
    </Placemark>"""
                kml_parts.append(linestring)
                kml_parts.append("  </Document>")
                kml_parts.append("</kml>")
                kml_content = "\n".join(kml_parts)
                
                # Préparer téléchargement KML
                kml_bytes = kml_content.encode('utf-8')
                
                # Préparer téléchargement KMZ (doc.kml dans une archive ZIP)
                kmz_buffer = BytesIO()
                with zipfile.ZipFile(kmz_buffer, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr('doc.kml', kml_content)
                kmz_bytes = kmz_buffer.getvalue()
                
                col_kml, col_kmz = st.columns(2)
                with col_kml:
                    st.download_button(
                        label="📥 Télécharger KML",
                        data=kml_bytes,
                        file_name=f"itineraire_{datetime.now().strftime('%Y%m%d')}.kml",
                        mime="application/vnd.google-earth.kml+xml",
                        use_container_width=True
                    )
                with col_kmz:
                    st.download_button(
                        label="📥 Télécharger KMZ (Google Earth)",
                        data=kmz_bytes,
                        file_name=f"itineraire_{datetime.now().strftime('%Y%m%d')}.kmz",
                        mime="application/vnd.google-earth.kmz",
                        use_container_width=True
                    )
            except Exception as e:
                st.warning(f"Impossible de préparer l'export KML/KMZ: {str(e)[:80]}…")
    
    with tab_export:
        st.subheader("Export")
        
        current_itinerary = st.session_state.manual_itinerary if st.session_state.manual_itinerary else itinerary
        
        excel_data = []
        for day, sdt, edt, desc in current_itinerary:
            excel_data.append({
                "Jour": day,
                "Date": (start_date + timedelta(days=day-1)).strftime("%d/%m/%Y"),
                "Début": sdt.strftime("%H:%M"),
                "Fin": edt.strftime("%H:%M"),
                "Durée (min)": int((edt - sdt).total_seconds() / 60),
                "Activité": desc
            })
        
        df_export = pd.DataFrame(excel_data)
        
        from io import BytesIO
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df_export.to_excel(writer, sheet_name='Planning', index=False)
            pd.DataFrame(sites_ordered).to_excel(writer, sheet_name='Sites', index=False)
        
        col_excel, col_html, col_ics = st.columns(3)
        
        with col_excel:
            st.download_button(
                label="📥 Télécharger Excel",
                data=output.getvalue(),
                file_name=f"mission_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        
        with col_html:
            html_export = build_professional_html(
                current_itinerary,
                start_date,
                stats,
                sites_ordered,
                segments_summary,
                default_speed_kmh,
                mission_title,
                coords_ordered,
                include_map=st.session_state.get("include_map_prof_html", False),
                lunch_start_time=st.session_state.get("lunch_start_time"),
                lunch_end_time=st.session_state.get("lunch_end_time"),
                lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
                prayer_start_time=st.session_state.get("prayer_start_time"),
                prayer_duration_min=st.session_state.get("prayer_duration_min", 20),
                include_details=st.session_state.get("include_prof_details", True)
            )
            st.download_button(
                label="📥 Télécharger HTML",
                data=html_export,
                file_name=f"mission_{datetime.now().strftime('%Y%m%d')}.html",
                mime="text/html",
                use_container_width=True
            )
        with col_ics:
            ics_export = build_ics_from_itinerary(current_itinerary, start_date, mission_title)
            st.download_button(
                label="📥 Télécharger ICS",
                data=ics_export,
                file_name=f"mission_{datetime.now().strftime('%Y%m%d')}.ics",
                mime="text/calendar",
                use_container_width=True
            )

    with tab_request:
        st.subheader("📑 Demande de mission")
        cols = st.columns(2)
        with cols[0]:
            objet = st.text_input("Objet", value=st.session_state.get("req_objet",""))
            justification = st.text_area("Justification", value=st.session_state.get("req_justification",""), height=100)
            demandeur = st.text_input("Demandeur", value=st.session_state.get("req_demandeur",""))
            email_demandeur = st.text_input("Email du demandeur", value=st.session_state.get("req_email_demandeur",""))
            telephone_demandeur = st.text_input("Téléphone du demandeur", value=st.session_state.get("req_telephone_demandeur",""))
            lieu_default = ', '.join([s['Ville'] for s in sites_ordered if s.get('Type')!='Base'])
            lieu = st.text_input("Lieu", value=st.session_state.get("req_lieu", lieu_default))
            date_depart = st.date_input("Date de départ", value=st.session_state.get("req_date_depart", start_date))
            date_retour = st.date_input("Date de retour", value=st.session_state.get("req_date_retour", start_date + timedelta(days=stats['total_days']-1)))
        with cols[1]:
            taches = st.text_area("Tâches principales", value=st.session_state.get("req_taches",""), height=100)
            perdiem_fcfa = st.number_input("Per diem (FCFA/jour)", min_value=0, value=int(st.session_state.get("req_perdiem_fcfa",8000)), step=1000)
            hotel_driver_fcfa = st.number_input("Hôtel chauffeur (FCFA/nuit)", min_value=0, value=int(st.session_state.get("req_hotel_driver_fcfa",60000)), step=1000)
            vehicule = st.selectbox("Véhicule", options=list(get_vehicle_types().keys()), index=0)
            fuel_data_req = calculate_fuel_consumption(stats.get('total_km',0), vehicule)
            carburant_l = st.number_input("Carburant (L)", min_value=0, value=int(fuel_data_req['fuel_needed_liters']) if fuel_data_req else 0)
            budget = st.number_input("Budget estimé (FCFA)", min_value=0, value=st.session_state.get("req_budget",0), step=10000)
        colA, colB, colC = st.columns(3)
        with colA:
            if st.button("✨ Compléter automatiquement (IA Adja)", use_container_width=True):
                md = collect_mission_data_for_ai()
                ai = generate_mission_request_ai_prefill(md, deepseek_api_key)
                if ai:
                    st.session_state.req_objet = ai.get("objet","")
                    st.session_state.req_justification = ai.get("justification","")
                    st.session_state.req_taches = "\n".join(ai.get("taches", []))
                    st.session_state.req_lieu = ai.get("lieu", st.session_state.get("req_lieu",""))
                    if ai.get("budget_perdiem_fcfa") is not None:
                        try:
                            st.session_state.req_perdiem_fcfa = int(ai.get("budget_perdiem_fcfa", 0))
                        except Exception:
                            pass
                    if ai.get("hotel_driver_fcfa") is not None:
                        try:
                            st.session_state.req_hotel_driver_fcfa = int(ai.get("hotel_driver_fcfa", 0))
                        except Exception:
                            pass
                    st.session_state.req_budget = int(ai.get("budget_estime_fcfa", budget or 0))
                    if ai.get("carburant_litres"):
                        st.session_state.req_carburant = int(ai["carburant_litres"])
                    try:
                        jours = (date_retour - date_depart).days + 1
                        nuits = max(0, jours - 1)
                        total_perdiem = int(jours * (st.session_state.get("req_perdiem_fcfa", perdiem_fcfa) or 0))
                        total_hotel = int(nuits * (st.session_state.get("req_hotel_driver_fcfa", hotel_driver_fcfa) or 0))
                        st.session_state.req_budget = total_perdiem + total_hotel
                        modern_alert(f"Budget estimé mis à jour: {st.session_state.req_budget:,} FCFA (per diem {jours}j, hôtel {nuits}n)", "info")
                    except Exception:
                        pass
                    st.success("Champs complétés")
                    st.rerun()
        with colB:
            if st.button("✉️ Générer email de demande", use_container_width=True):
                email_text = f"Objet: {objet}\n\nBonjour,\n\nMerci de valider la mission {mission_title} du {date_depart.strftime('%d/%m/%Y')} au {date_retour.strftime('%d/%m/%Y')} à {lieu}.\nDemandeur: {demandeur}.\nVéhicule: {vehicule}, carburant: {carburant_l} L.\nBudget estimé: {budget:,} FCFA.\nPer diem: {int(perdiem_fcfa):,} FCFA/jour • Hôtel chauffeur: {int(hotel_driver_fcfa):,} FCFA/nuit\n\nJustification:\n{justification}\n\nTâches:\n{taches}\n\nCordialement."
                st.code(email_text)

        with colC:
            if st.button("📤 Soumettre la demande", type="primary", use_container_width=True):
                required_fields = {
                    "Objet": objet,
                    "Justification": justification,
                    "Demandeur": demandeur,
                    "Email": email_demandeur,
                    "Téléphone": telephone_demandeur,
                    "Lieu": lieu,
                }
                missing = [k for k, v in required_fields.items() if not str(v).strip()]
                if missing:
                    modern_alert(f"Champs obligatoires manquants: {', '.join(missing)}", "error")
                else:
                    try:
                        from firebase_config import MissionRequestManager
                        mgr = MissionRequestManager()
                        dt_depart = datetime.combine(date_depart, time(8, 0)) if isinstance(date_depart, datetime) == False else date_depart
                        dt_retour = datetime.combine(date_retour, time(23, 59)) if isinstance(date_retour, datetime) == False else date_retour
                        req = {
                            "motif_mission": objet.strip(),
                            "destination": lieu.strip(),
                            "date_depart": dt_depart,
                            "date_retour": dt_retour,
                            "nb_passagers": 1,
                            "type_vehicule": vehicule,
                            "avec_chauffeur": True,
                            "notes": justification.strip(),
                            "budget_estime_fcfa": int(budget or 0),
                            "budget_perdiem_fcfa": int(perdiem_fcfa or 0),
                            "hotel_driver_fcfa": int(hotel_driver_fcfa or 0),
                            "carburant_litres": int(carburant_l or 0),
                            "participants": demandeur,
                            "service_demandeur": st.session_state.get("service_demandeur",""),
                            "nom_demandeur": demandeur.strip(),
                            "email_demandeur": email_demandeur.strip().lower(),
                            "telephone_demandeur": telephone_demandeur.strip(),
                        }
                        rid = mgr.create_request(req)
                        modern_alert(f"Demande soumise avec succès • Numéro: {rid}", "success")
                    except Exception as e:
                        modern_alert(f"Erreur lors de la soumission: {e}", "error")

    with tab_report:
        st.subheader("📋 Génération de rapport de mission")
        
        with st.expander("🤖 Générer un rapport complet", expanded=False):
            st.markdown("**Utilisez l'IA Adja pour générer un rapport professionnel orienté activités**")
            
            # Onglets pour organiser l'interface
            tab_basic, tab_details, tab_questions, tab_construction, tab_generate = st.tabs([
                "📝 Rapport basique", "📋 Détails mission", "🤖 Questions IA Adja", "🏗️ Procès-verbal", "🚀 Génération"
            ])
            
            with tab_basic:
                st.markdown("### 📄 Rapport rapide (version simplifiée)")
                
                # Options de rapport basique
                col1, col2 = st.columns(2)
                
                with col1:
                    report_type = st.selectbox(
                    "Type de rapport",
                    ["Rapport complet", "Résumé exécutif", "Rapport technique", "Rapport financier", "Procès-verbal professionnel"],
                    help="Choisissez le type de rapport à générer"
                )
            
                with col2:
                    report_tone = st.selectbox(
                    "Ton du rapport",
                    ["Professionnel", "Formel", "Décontracté", "Technique"],
                    help="Définissez le ton du rapport"
                )
            
                # Options avancées (sans expander imbriqué)
                st.markdown("**Options avancées**")
                
                col_opt1, col_opt2 = st.columns(2)
                
                with col_opt1:
                    include_recommendations = st.checkbox("Inclure des recommandations", value=True)
                    include_risks = st.checkbox("Analyser les risques", value=False)
            
                with col_opt2:
                    include_costs = st.checkbox("Estimation des coûts", value=False)
                    include_timeline = st.checkbox("Planning détaillé", value=True)
            
                custom_context = st.text_area(
                "Contexte supplémentaire (optionnel)",
                placeholder="Ajoutez des informations spécifiques à votre mission...",
                height=100
                )
            
                if st.button("🚀 Générer rapport basique", type="primary", use_container_width=True):
                    if st.session_state.planning_results:
                        # Animation d'attente améliorée avec barre de progression
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        try:
                            status_text.text("🔄 Collecte des données de mission...")
                            progress_bar.progress(20)
                            mission_data = collect_mission_data_for_ai()
                            status_text.text("📝 Construction du prompt...")
                            progress_bar.progress(40)
                            prompt = build_report_prompt(
                                mission_data, report_type, report_tone,
                                include_recommendations, include_risks, 
                                include_costs, include_timeline, custom_context
                            )
                            status_text.text("🤖 Génération du rapport par l'IA Adja...")
                            progress_bar.progress(60)
                            response = requests.post(
                                "https://api.deepseek.com/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {deepseek_api_key}",
                                    "Content-Type": "application/json"
                                },
                                json={
                                    "model": "deepseek-chat",
                                    "messages": [{"role": "user", "content": prompt}],
                                    "temperature": 0.7,
                                    "max_tokens": 4000
                                }
                            )
                            status_text.text("✅ Finalisation du rapport...")
                            progress_bar.progress(100)
                            if response.status_code == 200:
                                report_content = response.json()["choices"][0]["message"]["content"]
                                
                            progress_bar.empty()
                            status_text.empty()
                             
                            modern_alert("Rapport généré avec succès!", "success")
                             
                            # Affichage du rapport
                            st.markdown("### 📄 Votre rapport")
                            st.markdown(report_content)
                            
                            # Boutons de téléchargement
                            col_txt, col_md, col_html = st.columns(3)
                            
                            with col_txt:
                                st.download_button(
                                    label="📄 TXT",
                                    data=report_content,
                                    file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_md:
                                st.download_button(
                                    label="📝 MD",
                                    data=report_content,
                                    file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                    mime="text/markdown",
                                    use_container_width=True
                                )
                            
                            with col_html:
                                html_content = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>Rapport de Mission</title>
                                    <style>
                                        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                        h1, h2, h3 {{ color: #2c3e50; }}
                                        h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                                        h2 {{ border-left: 4px solid #3498db; padding-left: 15px; }}
                                        .header {{ text-align: center; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 10px; }}
                                        .footer {{ margin-top: 30px; text-align: center; font-size: 0.9em; color: #666; }}
                                        ul, ol {{ margin-left: 20px; }}
                                        strong {{ color: #2c3e50; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>Rapport de Mission</h1>
                                        <p><strong>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</strong></p>
                                        <p>Type: {report_type} | Ton: {report_tone}</p>
                                    </div>
                                {report_content.replace(chr(10), '<br>')}
                                    <div class="footer">
                                        <p>Rapport généré automatiquement par l'IA Adja DeepSeek</p>
                                    </div>
                                </body>
                                </html>
                                """
                                st.download_button(
                                    label="🌐 HTML",
                                    data=html_content,
                                    file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            if response.status_code != 200:
                                st.error(f"❌ Erreur API: {response.status_code}")
                                st.error(response.text)
                        except Exception as e:
                            status_text.error(f"Une erreur est survenue lors de la génération du rapport : {e}")
                            progress_bar.empty()
                else:
                    st.warning("⚠️ Aucun planning disponible. Veuillez d'abord optimiser votre itinéraire.")
            
            with tab_details:
                st.markdown("### 📋 Informations détaillées de la mission")
                st.info("💡 Remplissez ces informations pour enrichir votre rapport")
                
                mission_data = collect_enhanced_mission_data()
            
            with tab_questions:
                st.markdown("### 🤖 Questions interactives pour personnaliser le rapport")
                st.info("💡 Répondez à ces questions pour obtenir un rapport sur mesure")
                
                questions_data = ask_interactive_questions()
            
            with tab_construction:
                st.markdown("### 🏗️ Génération de procès-verbal de chantier")
                st.info("💡 Créez un procès-verbal professionnel pour vos visites de chantier")
                
                construction_data = collect_construction_report_data()
                
                if st.button("📋 Générer le procès-verbal", type="primary", use_container_width=True):
                    if st.session_state.planning_results:
                        # Animation d'attente améliorée
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        try:
                            # Étape 1: Collecte des données
                            status_text.text("📋 Collecte des données de mission...")
                            progress_bar.progress(20)
                            time.sleep(0.3)
                            
                            mission_data = collect_mission_data_for_ai()
                            questions_data = construction_data
                            
                            # Étape 2: Préparation du rapport
                            status_text.text("🔧 Préparation du procès-verbal...")
                            progress_bar.progress(40)
                            time.sleep(0.3)
                            
                            # Étape 3: Génération IA Adja Adja
                            status_text.text("🤖 Génération par l'IA Adja...")
                            progress_bar.progress(70)
                            
                            pv_result = generate_pv_report(mission_data, questions_data, deepseek_api_key)
                            
                            # Étape 4: Finalisation
                            status_text.text("✨ Finalisation du procès-verbal...")
                            progress_bar.progress(100)
                            time.sleep(0.3)
                            
                            # Nettoyage de l'animation
                            progress_bar.empty()
                            status_text.empty()
                            
                            if pv_result["success"]:
                                st.success("✅ Procès-verbal généré avec succès!")
                                
                                # Affichage du PV
                                st.markdown("### 📋 Votre procès-verbal")
                                st.markdown(pv_result["content"])
                                
                                # Informations du PV pour les téléchargements
                                pv_structure = construction_data.get('pv_structure', 'Structure non spécifiée')
                                pv_date = construction_data.get('pv_date', datetime.now().date())
                                pv_site = construction_data.get('pv_site', 'Site non spécifié')
                                pv_zone = construction_data.get('pv_zone', 'Zone non spécifiée')
                                pv_mission_type = construction_data.get('pv_mission_type', 'Mission non spécifiée')
                                pv_responsable = construction_data.get('pv_responsable', 'Responsable non spécifié')
                                pv_fonction = construction_data.get('pv_fonction', 'Fonction non spécifiée')
                                pv_content = pv_result["content"]
                                
                                # Boutons de téléchargement
                                col_pv_txt, col_pv_html, col_pv_pdf, col_pv_rtf = st.columns(4)
                                
                                with col_pv_txt:
                                    pv_txt_content = f"""
PROCÈS-VERBAL DE VISITE DE CHANTIER

Structure: {pv_structure}
Date: {pv_date.strftime('%d/%m/%Y')}
Site: {pv_site}
Zone: {pv_zone}

{pv_content}

Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}

{pv_responsable}
{pv_fonction}
                                    """
                                    st.download_button(
                                        label="📄 TXT",
                                        data=pv_txt_content.strip(),
                                        file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_pv_html:
                                    pv_html_content = f"""
                                        <!DOCTYPE html>
                                        <html>
                                        <head>
                                            <meta charset="UTF-8">
                                            <title>Procès-verbal de visite de chantier</title>
                                            <style>
                                                body {{ font-family: 'Arial', sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                                .header {{ text-align: center; margin-bottom: 30px; }}
                                                .header h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                                                .info-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                                                .info-table td {{ padding: 8px; border: 1px solid #ddd; }}
                                                .info-table .label {{ background-color: #f8f9fa; font-weight: bold; width: 120px; }}
                                                .signature {{ margin-top: 50px; }}
                                                .signature-line {{ border-top: 1px solid #333; width: 200px; margin: 20px 0; }}
                                            </style>
                                        </head>
                                        <body>
                                            <div class="header">
                                                <h1>Procès-verbal de visite de chantier</h1>
                                                <p><strong>{pv_structure}</strong></p>
                                                <p>Travaux d'extension PA DAL zone {pv_zone}</p>
                                            </div>

                                            <table class="info-table">
                                                <tr>
                                                    <td class="label">DATE:</td>
                                                    <td>{pv_date.strftime('%d/%m/%Y')}</td>
                                                    <td class="label">SITE:</td>
                                                    <td>{pv_site}</td>
                                                </tr>
                                                <tr>
                                                    <td class="label">MISSION:</td>
                                                    <td>{pv_mission_type}</td>
                                                    <td class="label">ZONE:</td>
                                                    <td>{pv_zone}</td>
                                                </tr>
                                                <tr>
                                                    <td class="label">RESPONSABLE:</td>
                                                    <td>{pv_responsable}</td>
                                                    <td class="label">FONCTION:</td>
                                                    <td>{pv_fonction}</td>
                                                </tr>
                                            </table>

                                            {pv_content.replace(chr(10), '<br>')}

                                        <div class="signature">
                                            <p>Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}</p>
                                            <div class="signature-line"></div>
                                            <p><strong>{pv_responsable}</strong></p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    st.download_button(
                                        label="🌐 HTML",
                                        data=pv_html_content,
                                        file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                                
                                with col_pv_pdf:
                                    if PDF_AVAILABLE:
                                        pv_full_content = f"""
PROCÈS-VERBAL DE VISITE DE CHANTIER

Structure: {pv_structure}
Date: {pv_date.strftime('%d/%m/%Y')}
Site: {pv_site}
Zone: {pv_zone}
Mission: {pv_mission_type}
Responsable: {pv_responsable}
Fonction: {pv_fonction}

{pv_content}"""
                                        pdf_data = create_pv_pdf(
                                            content=pv_full_content,
                                            title="Procès-verbal de visite de chantier",
                                            author=pv_responsable
                                        )
                                        st.download_button(
                                            label="📄 PDF",
                                            data=pdf_data,
                                            file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                    else:
                                        st.info("PDF non disponible")
                                
                                # Suppression de l'export RTF au profit du DOCX
                            
                            else:
                                st.error(f"❌ Erreur: {pv_result['error']}")
                        
                        except Exception as e:
                            st.error("❌ Erreur lors de la génération du procès-verbal")

            with tab_generate:
                st.markdown("### 🚀 Génération du rapport amélioré")
                st.info("💡 Utilisez cette section après avoir rempli les détails et répondu aux questions")
                
                # Vérification des prérequis
                has_details = hasattr(st.session_state, 'mission_context') and st.session_state.mission_context.get('objective')
                has_questions = 'report_focus' in st.session_state

                if has_details:
                    st.success("✅ Données détaillées collectées")
                else:
                    st.warning("⚠️ Aucune donnée détaillée - Allez dans l'onglet 'Détails mission'")

                if has_questions:
                    st.success("✅ Questions répondues")
                else:
                    st.warning("⚠️ Questions non répondues - Allez dans l'onglet 'Questions IA Adja'")
                
                # Aperçu des paramètres
                if has_questions:
                    st.markdown("**Paramètres du rapport :**")
                    col_preview1, col_preview2 = st.columns(2)

                    with col_preview1:
                        if 'report_focus' in st.session_state:
                            st.write(f"🎯 **Focus :** {', '.join(st.session_state.report_focus)}")
                        if 'target_audience' in st.session_state:
                            st.write(f"👥 **Public :** {st.session_state.target_audience}")

                    with col_preview2:
                        if 'report_length' in st.session_state:
                            st.write(f"📄 **Longueur :** {st.session_state.report_length}")
                        if 'specific_request' in st.session_state and st.session_state.specific_request:
                            st.write(f"✨ **Demande spéciale :** Oui")
                
                # Boutons d'action
                col_gen1, col_gen2 = st.columns([2, 1])

                with col_gen1:
                    generate_enhanced = st.button(
                        "🚀 Générer le rapport amélioré", 
                        type="primary", 
                        use_container_width=True,
                        disabled=not (has_details or has_questions)
                    )

                with col_gen2:
                    if st.button("🔄 Réinitialiser", use_container_width=True):
                        # Supprimer toutes les données de session liées au rapport
                        for key in list(st.session_state.keys()):
                            if key.startswith(('mission_', 'activity_', 'report_', 'target_', 'specific_', 'notes_', 'success_', 'contacts_', 'outcomes_', 'follow_up_', 'challenges', 'lessons_', 'recommendations', 'overall_', 'highlight_', 'discuss_', 'future_', 'cost_', 'time_', 'stakeholder_', 'include_')):
                                del st.session_state[key]
                        st.rerun()

                if generate_enhanced:
                    if st.session_state.planning_results:
                        # Animation d'attente améliorée avec barre de progression
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        try:
                            status_text.text("🔄 Collecte des données de mission...")
                            progress_bar.progress(15)
                            mission_data = collect_mission_data_for_ai()
                            
                            status_text.text("📋 Préparation des questions...")
                            progress_bar.progress(30)
                            time.sleep(0.5)
                            
                            questions_data = {
                                'report_focus': st.session_state.get('report_focus', []),
                                'target_audience': st.session_state.get('target_audience', 'Direction générale'),
                                'report_length': st.session_state.get('report_length', 'Moyen (3-5 pages)'),
                                'include_metrics': st.session_state.get('include_metrics', True),
                                'highlight_successes': st.session_state.get('highlight_successes', True),
                                'discuss_challenges': st.session_state.get('discuss_challenges', True),
                                'future_planning': st.session_state.get('future_planning', True),
                                'cost_analysis': st.session_state.get('cost_analysis', False),
                                'time_efficiency': st.session_state.get('time_efficiency', True),
                                'stakeholder_feedback': st.session_state.get('stakeholder_feedback', False),
                                'specific_request': st.session_state.get('specific_request', '')
                            }
                            
                            # Étape 3: Construction du prompt
                            status_text.text("🔧 Construction du prompt personnalisé...")
                            progress_bar.progress(50)
                            time.sleep(0.5)
                            
                            # Étape 4: Génération IA Adja
                            status_text.text("🤖 Génération du rapport par l'IA Adja...")
                            progress_bar.progress(70)
                            
                            # Génération du rapport
                            report_result = generate_enhanced_ai_report(
                                mission_data, 
                                questions_data,
                                deepseek_api_key
                            )
                            
                            # Étape 5: Finalisation
                            status_text.text("✨ Finalisation du rapport...")
                            progress_bar.progress(100)
                            time.sleep(0.3)
                            
                            # Nettoyage de l'animation
                            progress_bar.empty()
                            status_text.empty()
                            
                            if report_result["success"]:
                                st.success("✅ Rapport amélioré généré avec succès!")
                                
                                # Affichage du rapport
                                st.markdown("### 📄 Votre rapport amélioré")
                                report_content = report_result["content"]
                                st.markdown(report_content)
                                
                                # Boutons de téléchargement
                                col_txt, col_md, col_html, col_copy = st.columns(4)
                                
                                with col_txt:
                                    st.download_button(
                                        label="📄 TXT",
                                        data=report_content,
                                        file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_md:
                                    st.download_button(
                                        label="📝 MD",
                                        data=report_content,
                                        file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                        mime="text/markdown",
                                        use_container_width=True
                                    )
                                
                                with col_html:
                                    html_content = f"""
                                    <!DOCTYPE html>
                                    <html>
                                    <head>
                                        <meta charset="UTF-8">
                                        <title>Rapport de Mission Amélioré</title>
                                        <style>
                                            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                            h1, h2, h3 {{ color: #2c3e50; }}
                                            h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                                            h2 {{ border-left: 4px solid #3498db; padding-left: 15px; }}
                                            .header {{ text-align: center; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 10px; }}
                                            .footer {{ margin-top: 30px; text-align: center; font-size: 0.9em; color: #666; }}
                                            ul, ol {{ margin-left: 20px; }}
                                            strong {{ color: #2c3e50; }}
                                        </style>
                                    </head>
                                    <body>
                                        <div class="header">
                                            <h1>Rapport de Mission Amélioré</h1>
                                            <p><strong>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</strong></p>
                                            <p>Public cible: {questions_data.get('target_audience', 'Non spécifié')}</p>
                                        </div>
                                        {report_content.replace(chr(10), '<br>')}
                                        <div class="footer">
                                            <p>Rapport généré automatiquement par l'IA Adja DeepSeek</p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    st.download_button(
                                        label="🌐 HTML",
                                        data=html_content,
                                        file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                                
                                with col_copy:
                                    st.code(report_content, language=None)
                            
                            else:
                                st.error(f"❌ Erreur: {report_result['error']}")
                        
                        except Exception as e:
                            st.error(f"❌ Erreur lors de la génération: {str(e)}")
                    else:
                        st.warning("⚠️ Aucun planning disponible. Veuillez d'abord optimiser votre itinéraire.")

# --------------------------
# MODULE RAPPORT IA ADJA AMÉLIORÉ (ANCIEN - À SUPPRIMER)
# --------------------------
if False and st.session_state.planning_results:
    st.markdown("---")
    st.header("📋 Génération de rapport de mission")
    
    with st.expander("🤖 Générer un rapport complet", expanded=False):
        st.markdown("**Utilisez l'IA Adja pour générer un rapport professionnel orienté activités**")
        
        # Onglets pour organiser l'interface
        tab_basic, tab_details, tab_questions, tab_construction, tab_generate = st.tabs([
            "📝 Rapport basique", "📋 Détails mission", "🤖 Questions IA Adja", "🏗️ Procès-verbal", "🚀 Génération"
        ])
        
        with tab_basic:
            st.markdown("### 📄 Rapport rapide (version simplifiée)")
            
            # Options de rapport basique
            col1, col2 = st.columns(2)
            
            with col1:
                report_type = st.selectbox(
                    "Type de rapport",
                    ["Rapport complet", "Résumé exécutif", "Rapport technique", "Rapport financier", "Procès-verbal professionnel"],
                    help="Choisissez le type de rapport à générer"
                )
            
            with col2:
                report_tone = st.selectbox(
                    "Ton du rapport",
                    ["Professionnel", "Formel", "Décontracté", "Technique"],
                    help="Définissez le ton du rapport"
                )
            
            # Options avancées (sans expander imbriqué)
            st.markdown("**Options avancées**")
            col3, col4 = st.columns(2)
            
            with col3:
                include_recommendations = st.checkbox("Inclure des recommandations", value=True)
                include_risks = st.checkbox("Inclure l'analyse des risques", value=True)
            
            with col4:
                include_costs = st.checkbox("Inclure l'analyse des coûts", value=True)
                include_timeline = st.checkbox("Inclure la timeline détaillée", value=True)
            
            custom_context = st.text_area(
                "Contexte supplémentaire (optionnel)",
                placeholder="Ajoutez des informations spécifiques sur votre mission, objectifs, contraintes...",
                height=100
            )
            
            # Bouton de génération basique
            if st.button("🚀 Générer le rapport basique", type="secondary", use_container_width=True):
                if not deepseek_api_key:
                    st.error("❌ Clé API DeepSeek manquante")
                else:
                    # Animation améliorée avec barre de progression
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Étape 1: Collecte des données
                    status_text.text("📋 Collecte des données de mission...")
                    progress_bar.progress(20)
                    mission_data = collect_mission_data_for_ai()
                    
                    # Étape 2: Préparation du prompt
                    status_text.text("🤖 Construction du prompt IA Adja...")
                    progress_bar.progress(40)
                    time.sleep(0.5)
                    
                    # Étape 3: Génération IA Adja
                    status_text.text("🤖 Génération du rapport par l'IA Adja...")
                    progress_bar.progress(70)
                    time.sleep(0.3)
                    
                    # Génération selon le type de rapport sélectionné
                    if report_type == "Procès-verbal professionnel":
                            # Génération du procès-verbal avec l'IA Adja
                            questions_data_pv = {
                                'context': custom_context,
                                'observations': 'Observations détaillées de la mission',
                                'issues': 'Problèmes identifiés lors de la mission',
                                'actions': 'Actions réalisées pendant la mission',
                                'recommendations': 'Recommandations pour la suite'
                            }
                            
                            report_content, error = generate_pv_report(
                                mission_data, 
                                questions_data_pv,
                                deepseek_api_key
                            )
                            
                            if error:
                                st.error(f"❌ Erreur lors de la génération du PV: {error}")
                            else:
                                st.success("✅ Procès-verbal généré avec succès!")
                                
                                # Affichage du PV
                                st.markdown("### 📋 Procès-verbal généré")
                                st.markdown(report_content)
                                
                                # Options d'export spécialisées pour le PV
                                st.markdown("### 💾 Export du procès-verbal")
                                col_txt, col_html, col_pdf = st.columns(3)
                                
                                with col_txt:
                                    st.download_button(
                                        label="📄 Télécharger TXT",
                                        data=report_content,
                                        file_name=f"pv_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_html:
                                    # HTML formaté pour le PV
                                    html_pv = f"""
                                    <!DOCTYPE html>
                                    <html>
                                    <head>
                                        <meta charset="UTF-8">
                                        <title>Procès-verbal de Mission</title>
                                        <style>
                                            @page {{ margin: 2cm; }}
                                            body {{ 
                                                font-family: 'Times New Roman', serif; 
                                                font-size: 12pt; 
                                                line-height: 1.4; 
                                                color: #000; 
                                                margin: 0;
                                            }}
                                            .header {{ 
                                                text-align: center; 
                                                margin-bottom: 30px; 
                                                border-bottom: 2px solid #000;
                                                padding-bottom: 15px;
                                            }}
                                            .header h1 {{ 
                                                font-size: 18pt; 
                                                margin: 0; 
                                                text-transform: uppercase;
                                                font-weight: bold;
                                            }}
                                            h2 {{ 
                                                font-size: 14pt; 
                                                margin: 25px 0 10px 0; 
                                                text-decoration: underline;
                                                font-weight: bold;
                                            }}
                                            h3 {{ 
                                                font-size: 12pt; 
                                                margin: 20px 0 8px 0; 
                                                font-weight: bold;
                                            }}
                                            .signature {{ 
                                                margin-top: 40px; 
                                                text-align: right;
                                            }}
                                            .signature-line {{ 
                                                border-top: 1px solid #000; 
                                                width: 200px; 
                                                margin: 30px 0 5px auto;
                                            }}
                                            ul {{ margin-left: 20px; }}
                                            li {{ margin-bottom: 5px; }}
                                        </style>
                                    </head>
                                    <body>
                                        <div class="header">
                                            <h1>Procès-verbal de Mission</h1>
                                            <p><strong>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</strong></p>
                                        </div>
                                        {report_content.replace(chr(10), '<br>')}
                                        <div class="signature">
                                            <p>Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}</p>
                                            <div class="signature-line"></div>
                                            <p><strong>Responsable Mission</strong></p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    
                                    st.download_button(
                                        label="🌐 Télécharger HTML",
                                        data=html_pv,
                                        file_name=f"pv_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                                
                                with col_pdf:
                                    st.info("💡 Ouvrez le fichier HTML dans votre navigateur et utilisez 'Imprimer > Enregistrer au format PDF' pour obtenir un PDF professionnel.")
                    else:
                        # Génération du rapport basique (utilisation de l'ancienne fonction)
                        # Pour le rapport basique, on utilise une version simplifiée
                        questions_data_simple = {
                                'report_focus': report_type,
                                'target_audience': 'Équipe',
                                'report_length': 'Moyen',
                                'include_successes': include_recommendations,
                                'include_challenges': include_risks,
                                'include_costs': include_costs,
                                'include_planning': include_timeline,
                                'custom_requests': custom_context
                            }
                        
                        report_content = generate_enhanced_ai_report(
                            mission_data_simple, 
                            questions_data_simple,
                            deepseek_api_key
                        )
                        
                        # Étape 4: Finalisation
                        status_text.text("✅ Finalisation du rapport...")
                        progress_bar.progress(100)
                        time.sleep(0.3)
                        
                        # Nettoyage des éléments d'animation
                        progress_bar.empty()
                        status_text.empty()
                        
                        if report_content:
                            st.success("✅ Rapport généré avec succès!")
                            
                            # Affichage du rapport
                            st.markdown("### 📄 Rapport généré")
                            st.markdown(report_content)
                            
                            # Options d'export
                            st.markdown("### 💾 Export du rapport")
                            
                            # Première ligne : formats de base
                            col_txt, col_md, col_html = st.columns(3)
                            
                            with col_txt:
                                st.download_button(
                                    label="📄 Télécharger TXT",
                                    data=report_content,
                                    file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_md:
                                st.download_button(
                                    label="📝 Télécharger MD",
                                    data=report_content,
                                    file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                    mime="text/markdown",
                                    use_container_width=True
                                )
                            
                            with col_html:
                                # Conversion HTML pour PDF
                                html_report = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>Rapport de Mission</title>
                                    <style>
                                        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                                        h1, h2, h3 {{ color: #2c3e50; }}
                                        .header {{ text-align: center; margin-bottom: 30px; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>Rapport de Mission</h1>
                                        <p>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</p>
                                    </div>
                                    {report_content.replace(chr(10), '<br>')}
                                </body>
                                </html>
                                """
                                
                                st.download_button(
                                    label="🌐 Télécharger HTML",
                                    data=html_report,
                                    file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            
                            # Deuxième ligne : formats professionnels (PDF et Word)
                            if PDF_AVAILABLE:
                                st.markdown("#### 📋 Formats professionnels")
                                col_pdf, col_word_rtf, col_word_docx = st.columns(3)
                                
                                with col_pdf:
                                    try:
                                        pdf_data = create_pv_pdf(
                                            content=report_content,
                                            title="Rapport de Mission",
                                            author="Responsable Mission"
                                        )
                                        st.download_button(
                                            label="📄 Télécharger PDF",
                                            data=pdf_data,
                                            file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur génération PDF: {str(e)}")
                                
                                # Suppression de l'export Word RTF, on conserve uniquement DOCX
                                
                                with col_word_docx:
                                    try:
                                        docx_data = create_docx_document(
                                            content=report_content,
                                            title="Rapport de Mission"
                                        )
                                        st.download_button(
                                            label="📄 Word (.docx)",
                                            data=docx_data,
                                            file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur génération Word DOCX: {str(e)}")
                            if not PDF_AVAILABLE:
                                st.info("💡 Installez reportlab pour activer l'export PDF et Word professionnel.")
                        else:
                            st.error("❌ Erreur lors de la génération du rapport")
        
        with tab_details:
            st.markdown("### 📋 Collecte de données détaillées")
            st.info("💡 Remplissez ces informations pour obtenir un rapport plus riche et personnalisé")
            
            # Interface de collecte de données enrichies
            collect_enhanced_mission_data()
        
        with tab_questions:
            st.markdown("### 🤖 Questions pour personnaliser le rapport")
            st.info("💡 Répondez à ces questions pour que l'IA Adja génère un rapport adapté à vos besoins")
            
            # Interface de questions interactives
            questions_data = ask_interactive_questions()
        
        with tab_construction:
            st.markdown("### 🏗️ Procès-verbal de visite de chantier")
            st.info("💡 Générez un procès-verbal professionnel au format officiel")
            
            # Formulaire pour procès-verbal de chantier
            st.markdown("#### 📋 Informations générales")
            
            col_pv1, col_pv2 = st.columns(2)
            
            with col_pv1:
                pv_date = st.date_input("📅 Date de visite", value=datetime.now().date())
                pv_site = st.text_input("🏗️ Site/Chantier", placeholder="Ex: Villengara et Kolda")
                pv_structure = st.text_input("🏢 Structure", placeholder="Ex: DAL/GPR/ESP")
                pv_zone = st.text_input("🗺️ Titre projet", placeholder="Ex: PA DAL zone SUD")
            
            with col_pv2:
                pv_mission_type = st.selectbox(
                    "📝 Type de mission",
                    ["Visite de chantier", "Inspection technique", "Suivi de travaux", "Réception de travaux", "Autre"]
                )
                pv_responsable = st.text_input("👤 Responsable mission", placeholder="Ex: Moctar TALL")
                pv_fonction = st.text_input("💼 Fonction", placeholder="Ex: Ingénieur")
                pv_contact = st.text_input("📞 Contact", placeholder="Ex: +221 XX XXX XX XX")
            
            st.markdown("#### 🎯 Objectifs de la mission")
            pv_objectifs = st.text_area(
                "Décrivez les objectifs principaux",
                placeholder="Ex: Contrôler l'avancement des travaux, vérifier la conformité, identifier les problèmes...",
                height=100
            )
            
            st.markdown("#### 📊 Observations et constats")
            
            # Sections d'observations
            col_obs1, col_obs2 = st.columns(2)
            
            with col_obs1:
                st.markdown("**🔍 Constats positifs**")
                pv_positifs = st.text_area(
                    "Points positifs observés",
                    placeholder="Ex: Respect des délais, qualité des matériaux, sécurité...",
                    height=120,
                    key="pv_positifs"
                )
                
                st.markdown("**⚠️ Points d'attention**")
                pv_attention = st.text_area(
                    "Points nécessitant une attention",
                    placeholder="Ex: Retards mineurs, ajustements nécessaires...",
                    height=120,
                    key="pv_attention"
                )
            
            with col_obs2:
                st.markdown("**❌ Problèmes identifiés**")
                pv_problemes = st.text_area(
                    "Problèmes et non-conformités",
                    placeholder="Ex: Défauts de construction, non-respect des normes...",
                    height=120,
                    key="pv_problemes"
                )
                
                st.markdown("**💡 Recommandations**")
                pv_recommandations = st.text_area(
                    "Actions recommandées",
                    placeholder="Ex: Corrections à apporter, améliorations suggérées...",
                    height=120,
                    key="pv_recommandations"
                )
            
            st.markdown("#### 📈 Avancement et planning")
            col_plan1, col_plan2 = st.columns(2)
            
            with col_plan1:
                pv_avancement = st.slider("📊 Avancement global (%)", 0, 100, 50)
                pv_respect_delais = st.selectbox("⏰ Respect des délais", ["Conforme", "Léger retard", "Retard important"])
            
            with col_plan2:
                pv_prochaine_visite = st.date_input("📅 Prochaine visite prévue", value=datetime.now().date() + timedelta(days=30))
                pv_urgence = st.selectbox("🚨 Niveau d'urgence", ["Faible", "Moyen", "Élevé", "Critique"])
            
            st.markdown("#### 👥 Participants et contacts")
            pv_participants = st.text_area(
                "Liste des participants à la visite",
                placeholder="Ex: Moctar TALL (Ingénieur), Jean DUPONT (Chef de chantier), Marie MARTIN (Architecte)...",
                height=80
            )
            
            # Génération du procès-verbal
            if st.button("📋 Générer le procès-verbal", type="primary", use_container_width=True):
                if not deepseek_api_key:
                    st.error("❌ Clé API DeepSeek manquante")
                elif not pv_site or not pv_objectifs:
                    st.error("❌ Veuillez remplir au minimum le site et les objectifs")
                else:
                    # Animation d'attente améliorée
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    try:
                        # Étape 1: Collecte des données
                        status_text.text("📋 Collecte des informations du chantier...")
                        progress_bar.progress(15)
                        time.sleep(0.3)
                        
                        # Données pour le procès-verbal
                        pv_data = {
                            'date': pv_date.strftime('%d/%m/%Y'),
                            'site': pv_site,
                            'structure': pv_structure,
                            'zone': pv_zone,
                            'mission_type': pv_mission_type,
                            'responsable': pv_responsable,
                            'fonction': pv_fonction,
                            'contact': pv_contact,
                            'objectifs': pv_objectifs,
                            'positifs': pv_positifs,
                            'attention': pv_attention,
                            'problemes': pv_problemes,
                            'recommandations': pv_recommandations,
                            'avancement': pv_avancement,
                            'respect_delais': pv_respect_delais,
                            'prochaine_visite': pv_prochaine_visite.strftime('%d/%m/%Y'),
                            'urgence': pv_urgence,
                            'participants': pv_participants
                        }
                        
                        # Mise à jour de l'animation - Préparation du rapport
                        progress_bar.progress(45)
                        status_text.text("📝 Préparation du rapport de chantier...")
                        time.sleep(0.5)
                        
                        # Mise à jour de l'animation - Génération IA Adja
                        progress_bar.progress(70)
                        status_text.text("🤖 Génération du rapport avec l'IA Adja...")
                        
                        # Génération avec l'IA Adja
                        pv_content = generate_construction_report(pv_data, deepseek_api_key)
                        
                        # Mise à jour de l'animation - Finalisation
                        progress_bar.progress(100)
                        status_text.text("✅ Rapport généré avec succès!")
                        time.sleep(0.5)
                        
                        # Nettoyage de l'animation
                        progress_bar.empty()
                        status_text.empty()
                        
                        if pv_content:
                            st.success("✅ Procès-verbal généré avec succès!")
                            
                            # Affichage du procès-verbal
                            st.markdown("### 📄 Procès-verbal généré")
                            st.markdown(pv_content)
                            
                            # Options d'export spécialisées
                            st.markdown("### 💾 Export du procès-verbal")
                            col_pv_txt, col_pv_pdf, col_pv_word = st.columns(3)
                            
                            with col_pv_txt:
                                st.download_button(
                                    label="📄 Format TXT",
                                    data=pv_content,
                                    file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_pv_pdf:
                                # HTML formaté pour impression PDF
                                html_pv = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>Procès-verbal de visite de chantier</title>
                                    <style>
                                        @page {{ margin: 2cm; }}
                                        body {{ 
                                            font-family: 'Times New Roman', serif; 
                                            font-size: 12pt; 
                                            line-height: 1.4; 
                                            color: #000; 
                                            margin: 0;
                                        }}
                                        .header {{ 
                                            text-align: center; 
                                            margin-bottom: 30px; 
                                            border-bottom: 2px solid #000;
                                            padding-bottom: 15px;
                                        }}
                                        .header h1 {{ 
                                            font-size: 18pt; 
                                            margin: 0; 
                                            text-transform: uppercase;
                                            font-weight: bold;
                                        }}
                                        .info-table {{ 
                                            width: 100%; 
                                            border-collapse: collapse; 
                                            margin: 20px 0;
                                        }}
                                        .info-table td {{ 
                                            border: 1px solid #000; 
                                            padding: 8px; 
                                            vertical-align: top;
                                        }}
                                        .info-table .label {{ 
                                            background-color: #f0f0f0; 
                                            font-weight: bold; 
                                            width: 30%;
                                        }}
                                        h2 {{ 
                                            font-size: 14pt; 
                                            margin: 25px 0 10px 0; 
                                            text-decoration: underline;
                                            font-weight: bold;
                                        }}
                                        h3 {{ 
                                            font-size: 12pt; 
                                            margin: 20px 0 8px 0; 
                                            font-weight: bold;
                                        }}
                                        .signature {{ 
                                            margin-top: 40px; 
                                            text-align: right;
                                        }}
                                        .signature-line {{ 
                                            border-top: 1px solid #000; 
                                            width: 200px; 
                                            margin: 30px 0 5px auto;
                                        }}
                                        ul {{ margin-left: 20px; }}
                                        li {{ margin-bottom: 5px; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>Procès-verbal de visite de chantier</h1>
                                        <p><strong>{pv_structure}</strong></p>
                                        <p>Travaux d'extension PA DAL zone {pv_zone}</p>
                                    </div>
                                    
                                    <table class="info-table">
                                        <tr>
                                            <td class="label">DATE:</td>
                                            <td>{pv_date.strftime('%d/%m/%Y')}</td>
                                            <td class="label">SITE:</td>
                                            <td>{pv_site}</td>
                                        </tr>
                                        <tr>
                                            <td class="label">MISSION:</td>
                                            <td>{pv_mission_type}</td>
                                            <td class="label">ZONE:</td>
                                            <td>{pv_zone}</td>
                                        </tr>
                                        <tr>
                                            <td class="label">RESPONSABLE:</td>
                                            <td>{pv_responsable}</td>
                                            <td class="label">FONCTION:</td>
                                            <td>{pv_fonction}</td>
                                        </tr>
                                    </table>
                                    
                                    {pv_content.replace(chr(10), '<br>')}
                                    
                                    <div class="signature">
                                        <p>Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}</p>
                                        <div class="signature-line"></div>
                                        <p><strong>{pv_responsable}</strong></p>
                                    </div>
                                </body>
                                </html>
                                """
                                
                                st.download_button(
                                    label="📋 Format HTML",
                                    data=html_pv,
                                    file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            
                            with col_pv_word:
                                # Format Word-compatible
                                word_content = f"""
                                PROCÈS-VERBAL DE VISITE DE CHANTIER
                                
                                Structure: {pv_structure}
                                Date: {pv_date.strftime('%d/%m/%Y')}
                                Site: {pv_site}
                                Zone: {pv_zone}
                                
                                {pv_content}
                                
                                Fait à Dakar, le {datetime.now().strftime('%d/%m/%Y')}
                                
                                {pv_responsable}
                                {pv_fonction}
                                """
                                
                                st.download_button(
                                    label="📝 Format TXT",
                                    data=word_content,
                                    file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            # Deuxième ligne : formats professionnels (PDF et Word)
                            if PDF_AVAILABLE:
                                st.markdown("#### 📋 Formats professionnels")
                                col_pv_pdf, col_pv_rtf, col_pv_docx = st.columns(3)
                                
                                with col_pv_pdf:
                                    try:
                                        # Contenu formaté pour le PV
                                        pv_full_content = f"""Structure: {pv_structure}
Date: {pv_date.strftime('%d/%m/%Y')}
Site: {pv_site}
Zone: {pv_zone}
Mission: {pv_mission_type}
Responsable: {pv_responsable}
Fonction: {pv_fonction}

{pv_content}"""
                                        
                                        pdf_data = create_pv_pdf(
                                            content=pv_full_content,
                                            title="Procès-verbal de visite de chantier",
                                            author=pv_responsable
                                        )
                                        st.download_button(
                                            label="📄 PDF",
                                            data=pdf_data,
                                            file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur génération PDF: {str(e)}")
                                
                                # Suppression de l'export Word RTF, on conserve uniquement DOCX
                                
                                with col_pv_docx:
                                    try:
                                        docx_data = create_docx_document(
                                            content=pv_full_content,
                                            title="Procès-verbal de visite de chantier"
                                        )
                                        st.download_button(
                                            label="📄 Word (.docx)",
                                            data=docx_data,
                                            file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur génération Word DOCX: {str(e)}")
                            else:
                                st.info("💡 Installez reportlab pour activer l'export PDF et Word professionnel.")

                    except Exception as e:
                        st.error(f"❌ Erreur lors de la génération du procès-verbal: {str(e)}")
                        try:
                            progress_bar.empty()
                            status_text.empty()
                        except Exception:
                            pass

        with tab_generate:
            st.markdown("### 🚀 Génération du rapport amélioré")
            st.info("💡 Utilisez cette section après avoir rempli les détails et répondu aux questions")
            
            # Vérification des données disponibles
            has_details = hasattr(st.session_state, 'mission_context') and st.session_state.mission_context.get('objective')
            has_questions = 'report_focus' in st.session_state
            
            if has_details:
                st.success("✅ Données détaillées collectées")
            else:
                st.warning("⚠️ Aucune donnée détaillée - Allez dans l'onglet 'Détails mission'")
            
            if has_questions:
                st.success("✅ Questions répondues")
            else:
                st.warning("⚠️ Questions non répondues - Allez dans l'onglet 'Questions IA Adja'")
            
            # Aperçu des paramètres
            if has_questions:
                st.markdown("**Paramètres du rapport :**")
                col_preview1, col_preview2 = st.columns(2)
                
                with col_preview1:
                    if 'report_focus' in st.session_state:
                        st.write(f"🎯 **Focus :** {', '.join(st.session_state.report_focus)}")
                    if 'target_audience' in st.session_state:
                        st.write(f"👥 **Public :** {st.session_state.target_audience}")
                
                with col_preview2:
                    if 'report_length' in st.session_state:
                        st.write(f"📄 **Longueur :** {st.session_state.report_length}")
                    if 'specific_request' in st.session_state and st.session_state.specific_request:
                        st.write(f"✨ **Demande spéciale :** Oui")
            
            # Bouton de génération améliorée
            col_gen1, col_gen2 = st.columns([2, 1])
            
            with col_gen1:
                generate_enhanced = st.button(
                    "🚀 Générer le rapport amélioré", 
                    type="primary", 
                    use_container_width=True,
                    disabled=not (has_details or has_questions)
                )
            
            with col_gen2:
                if st.button("🔄 Réinitialiser", use_container_width=True):
                    # Réinitialiser les données
                    for key in list(st.session_state.keys()):
                        if key.startswith(('mission_', 'activity_', 'report_', 'target_', 'specific_', 'notes_', 'success_', 'contacts_', 'outcomes_', 'follow_up_', 'challenges', 'lessons_', 'recommendations', 'overall_', 'highlight_', 'discuss_', 'future_', 'cost_', 'time_', 'stakeholder_', 'include_')):
                            del st.session_state[key]
                    st.rerun()
            
            if generate_enhanced:
                if not deepseek_api_key:
                    st.error("❌ Clé API DeepSeek manquante")
                else:
                    # Initialisation de l'animation d'attente
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Étape 1: Collecte des données de mission
                    status_text.text("📋 Collecte des données de mission...")
                    progress_bar.progress(20)
                    time.sleep(0.5)
                    mission_data = collect_mission_data_for_ai()
                    
                    # Étape 2: Collecte des réponses aux questions
                    status_text.text("❓ Collecte des réponses aux questions...")
                    progress_bar.progress(40)
                    time.sleep(0.5)
                    questions_data = {
                        'report_focus': st.session_state.get('report_focus', []),
                        'target_audience': st.session_state.get('target_audience', 'Direction générale'),
                        'report_length': st.session_state.get('report_length', 'Moyen (3-5 pages)'),
                        'include_metrics': st.session_state.get('include_metrics', True),
                        'highlight_successes': st.session_state.get('highlight_successes', True),
                        'discuss_challenges': st.session_state.get('discuss_challenges', True),
                        'future_planning': st.session_state.get('future_planning', True),
                        'cost_analysis': st.session_state.get('cost_analysis', False),
                        'time_efficiency': st.session_state.get('time_efficiency', True),
                        'stakeholder_feedback': st.session_state.get('stakeholder_feedback', False),
                        'specific_request': st.session_state.get('specific_request', '')
                    }
                    
                    # Étape 3: Construction du prompt
                    status_text.text("🔧 Construction du prompt personnalisé...")
                    progress_bar.progress(60)
                    time.sleep(0.5)
                    
                    # Étape 4: Génération du rapport amélioré
                    status_text.text("🤖 Génération du rapport amélioré par l'IA Adja...")
                    progress_bar.progress(80)
                    time.sleep(0.5)
                    report_content = generate_enhanced_ai_report(
                        mission_data, 
                        questions_data,
                        deepseek_api_key
                        )
                        
                    # Étape 5: Finalisation
                    status_text.text("✅ Finalisation du rapport...")
                    progress_bar.progress(100)
                    time.sleep(0.5)
                    
                    # Nettoyage de l'animation
                    progress_bar.empty()
                    status_text.empty()
                        
                    if report_content:
                            modern_alert("Rapport amélioré généré avec succès!", "success")
                            
                            # Affichage du rapport
                            st.markdown("### 📄 Rapport généré")
                            st.markdown(report_content)
                            
                            # Options d'export améliorées
                            st.markdown("### 💾 Export du rapport")
                            col_txt, col_md, col_html, col_copy = st.columns(4)
                            
                            with col_txt:
                                st.download_button(
                                    label="📄 TXT",
                                    data=report_content,
                                    file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_md:
                                st.download_button(
                                    label="📝 MD",
                                    data=report_content,
                                    file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                    mime="text/markdown",
                                    use_container_width=True
                                )
                            
                            with col_html:
                                # Conversion HTML améliorée
                                html_report = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>Rapport de Mission Amélioré</title>
                                    <style>
                                        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                        h1, h2, h3 {{ color: #2c3e50; }}
                                        h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                                        h2 {{ border-left: 4px solid #3498db; padding-left: 15px; }}
                                        .header {{ text-align: center; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 10px; }}
                                        .footer {{ margin-top: 30px; text-align: center; font-size: 0.9em; color: #666; }}
                                        ul, ol {{ margin-left: 20px; }}
                                        strong {{ color: #2c3e50; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>Rapport de Mission Amélioré</h1>
                                        <p><strong>Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}</strong></p>
                                        <p>Public cible: {questions_data.get('target_audience', 'Non spécifié')}</p>
                                    </div>
                                    {report_content.replace(chr(10), '<br>')}
                                    <div class="footer">
                                        <p>Rapport généré automatiquement par l'IA Adja DeepSeek</p>
                                    </div>
                                </body>
                                </html>
                                """
                                
                                st.download_button(
                                    label="🌐 HTML",
                                    data=html_report,
                                    file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            
                            with col_copy:
                                if st.button("📋 Copier", use_container_width=True):
                                    st.write("📋 Contenu copié dans le presse-papiers!")
                                    st.code(report_content, language=None)
                    else:
                        st.error("❌ Erreur lors de la génération du rapport")

st.markdown("---")
st.caption("🚀 Planificateur de Mission v2.4")
st.caption("💻 Developed by @Moctar All rights reserved")

